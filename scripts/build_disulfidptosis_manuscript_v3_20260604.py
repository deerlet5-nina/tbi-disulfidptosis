from __future__ import annotations

import gzip
import math
import re
import shutil
import subprocess
import textwrap
import urllib.request
from io import StringIO
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from scipy import stats


DATE = "2026-06-04"
PRIORITY_GENES = ["SLC3A2", "SLC7A11", "WASF2", "TLN1", "ACTB", "MYH9", "MYL6", "FLNA"]
TRANSPORTER_GENES = ["SLC3A2", "SLC7A11"]
ACTIN_GENES = ["WASF2", "TLN1", "ACTB", "MYH9", "MYL6", "FLNA"]
CELL_MARKERS = {
    "Neuron": ["RBFOX3", "SNAP25", "SYT1", "MAP2"],
    "Astrocyte": ["GFAP", "AQP4", "ALDH1L1", "SLC1A3"],
    "Microglia": ["P2RY12", "CX3CR1", "AIF1", "CSF1R"],
    "Oligodendrocyte": ["MBP", "MOG", "PLP1", "MAG"],
    "OPC": ["PDGFRA", "CSPG4", "VCAN", "SOX10"],
    "Endothelial": ["PECAM1", "VWF", "CLDN5", "FLT1"],
}
OKABE = {
    "blue": "#0072B2",
    "orange": "#E69F00",
    "green": "#009E73",
    "red": "#D55E00",
    "purple": "#CC79A7",
    "sky": "#56B4E9",
    "yellow": "#F0E442",
    "black": "#000000",
    "gray": "#6B7280",
}


def find_file(name: str, contains: str | None = None) -> Path:
    matches = sorted(Path.cwd().rglob(name))
    if contains:
        matches = [p for p in matches if contains in str(p)]
    if not matches:
        raise FileNotFoundError(name)
    return matches[0]


WORKDIR = find_file("TBI_disulfidptosis_focused_design_tables_20260604.xlsx").parents[1]
TABLEDIR = WORKDIR / "tables"
FIGDIR = WORKDIR / "figures"
REPORTDIR = WORKDIR / "reports"
RAWDIR = WORKDIR / "raw_data_v3_20260604"
for d in (TABLEDIR, FIGDIR, REPORTDIR, RAWDIR):
    d.mkdir(parents=True, exist_ok=True)


def configure_style() -> None:
    mpl.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "font.size": 8.5,
            "axes.labelsize": 8.5,
            "axes.titlesize": 9.5,
            "xtick.labelsize": 7.5,
            "ytick.labelsize": 7.5,
            "legend.fontsize": 7.5,
            "figure.dpi": 160,
            "savefig.dpi": 320,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
        }
    )
    sns.set_theme(style="ticks", font="Microsoft YaHei", font_scale=0.9)


def savefig(fig: plt.Figure, name: str) -> tuple[Path, Path]:
    png = FIGDIR / f"{name}.png"
    pdf = FIGDIR / f"{name}.pdf"
    fig.savefig(png, bbox_inches="tight", dpi=320)
    fig.savefig(pdf, bbox_inches="tight")
    plt.close(fig)
    return png, pdf


def bh_fdr(values: pd.Series | np.ndarray) -> np.ndarray:
    p = pd.to_numeric(pd.Series(values), errors="coerce").to_numpy(dtype=float)
    out = np.full(len(p), np.nan)
    ok = np.isfinite(p)
    if not ok.any():
        return out
    vals = p[ok]
    order = np.argsort(vals)
    ranked = vals[order]
    n = len(ranked)
    adj = ranked * n / np.arange(1, n + 1)
    adj = np.minimum.accumulate(adj[::-1])[::-1]
    restored = np.empty(n)
    restored[order] = np.clip(adj, 0, 1)
    out[ok] = restored
    return out


def clean_gene(x: object) -> str:
    if pd.isna(x):
        return ""
    s = str(x).strip().split(";")[0].split("///")[0].split(",")[0].strip()
    return s.upper()


def zscore_rows(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy().astype(float)
    mean = out.mean(axis=1)
    sd = out.std(axis=1, ddof=1).replace(0, np.nan)
    return out.sub(mean, axis=0).div(sd, axis=0)


def module_score(expr: pd.DataFrame, genes: list[str]) -> pd.Series:
    present = [g for g in genes if g in expr.index]
    if not present:
        return pd.Series(np.nan, index=expr.columns)
    z = zscore_rows(expr.loc[present])
    return z.mean(axis=0)


def add_panel_label(ax: plt.Axes, label: str) -> None:
    ax.text(-0.12, 1.08, label, transform=ax.transAxes, fontsize=11, fontweight="bold", va="top", ha="left")


def p_to_star(p: float | int | None) -> str:
    if p is None or not np.isfinite(p):
        return ""
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    if p < 0.10:
        return "†"
    return ""


def download_file(url: str, dest: Path) -> Path:
    if dest.exists() and dest.stat().st_size > 0:
        return dest
    if dest.exists() and dest.stat().st_size == 0:
        dest.unlink()
    curl = shutil.which("curl.exe") or shutil.which("curl")
    if curl:
        completed = subprocess.run(
            [curl, "-L", "--max-time", "180", "--retry", "2", url, "-o", str(dest)],
            capture_output=True,
            text=True,
            check=False,
        )
        if completed.returncode == 0 and dest.exists() and dest.stat().st_size > 0:
            return dest
    try:
        with urllib.request.urlopen(url, timeout=120) as response, dest.open("wb") as handle:
            shutil.copyfileobj(response, handle)
    except Exception:
        raise
    if not dest.exists() or dest.stat().st_size == 0:
        raise RuntimeError(f"Failed to download {url}")
    return dest


def parse_series_matrix(path: Path) -> tuple[pd.DataFrame, dict[str, list[list[str]]]]:
    sample_meta: dict[str, list[list[str]]] = {}
    table_lines: list[str] = []
    reading_table = False
    with gzip.open(path, "rt", encoding="utf-8", errors="ignore") as handle:
        for raw_line in handle:
            line = raw_line.rstrip("\n")
            if line == "!series_matrix_table_begin":
                reading_table = True
                continue
            if line == "!series_matrix_table_end":
                reading_table = False
                continue
            if reading_table:
                table_lines.append(line)
                continue
            if line.startswith("!Sample_"):
                parts = line.split("\t")
                key = parts[0][1:]
                values = [part.strip().strip('"') for part in parts[1:]]
                sample_meta.setdefault(key, []).append(values)
    expr = pd.read_csv(StringIO("\n".join(table_lines)), sep="\t")
    expr = expr.rename(columns={"ID_REF": "probe_id"}).set_index("probe_id")
    expr = expr.apply(pd.to_numeric, errors="coerce")
    return expr, sample_meta


def parse_platform_table(path: Path) -> pd.DataFrame:
    table_lines: list[str] = []
    reading = False
    with gzip.open(path, "rt", encoding="utf-8", errors="ignore") as handle:
        for raw_line in handle:
            line = raw_line.rstrip("\n")
            if line == "!platform_table_begin":
                reading = True
                continue
            if line == "!platform_table_end" and reading:
                break
            if reading:
                table_lines.append(line)
    platform = pd.read_csv(StringIO("\n".join(table_lines)), sep="\t")
    platform = platform.rename(columns={"ID": "probe_id", "ORF": "gene_symbol", "TargetID": "target_id"})
    if "gene_symbol" not in platform.columns:
        for col in platform.columns:
            if "symbol" in col.lower() or "orf" in col.lower():
                platform = platform.rename(columns={col: "gene_symbol"})
                break
    return platform


def meta_values(meta: dict[str, list[list[str]]], key: str, idx: int = 0) -> list[str]:
    rows = meta.get(key, [])
    return rows[idx] if rows else []


def parse_characteristics(values: list[str]) -> dict[str, str]:
    out: dict[str, str] = {}
    for value in values:
        if ": " in value:
            k, v = value.split(": ", 1)
            out[k.strip().lower()] = v.strip()
    return out


def build_gse223245_metadata(meta: dict[str, list[list[str]]]) -> pd.DataFrame:
    accessions = meta_values(meta, "Sample_geo_accession")
    titles = meta_values(meta, "Sample_title")
    sources = meta_values(meta, "Sample_source_name_ch1")
    char_rows = meta.get("Sample_characteristics_ch1", [])
    records: list[dict[str, object]] = []
    for i, gsm in enumerate(accessions):
        title = titles[i] if i < len(titles) else ""
        low = title.lower()
        group = "Unknown"
        if "control" in low:
            group = "Control"
        elif "mild" in low:
            group = "Mild"
        elif "moderate" in low:
            group = "Moderate"
        elif "severe" in low:
            group = "Severe"
        chars = parse_characteristics([row[i] for row in char_rows if i < len(row)])
        records.append(
            {
                "sample_id": gsm,
                "title": title,
                "group": group,
                "severity_score": {"Control": 0, "Mild": 1, "Moderate": 2, "Severe": 3}.get(group, np.nan),
                "age": pd.to_numeric(chars.get("age"), errors="coerce"),
                "sex": chars.get("gender", "").upper(),
                "source": sources[i] if i < len(sources) else "",
            }
        )
    metadata = pd.DataFrame(records)
    metadata["condition"] = np.where(metadata["group"].eq("Control"), "Control", "TBI")
    return metadata


def collapse_gse223245(expr: pd.DataFrame, platform: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    annot = platform.copy()
    annot["gene_symbol"] = annot["gene_symbol"].map(clean_gene)
    annot = annot[annot["gene_symbol"].ne("")]
    merged = expr.join(annot.set_index("probe_id"), how="inner")
    sample_cols = expr.columns.tolist()
    merged["probe_mean"] = merged[sample_cols].mean(axis=1)
    merged = merged.reset_index().rename(columns={"index": "probe_id"})
    dedup = merged.sort_values(["gene_symbol", "probe_mean"], ascending=[True, False]).drop_duplicates("gene_symbol")
    gene_expr = dedup.set_index("gene_symbol")[sample_cols].copy()
    probe_map = dedup[["gene_symbol", "probe_id", "probe_mean"]].copy()
    return gene_expr, probe_map


def analyze_gse223245() -> dict[str, pd.DataFrame]:
    matrix_url = "https://ftp.ncbi.nlm.nih.gov/geo/series/GSE223nnn/GSE223245/matrix/GSE223245_series_matrix.txt.gz"
    soft_url = "https://ftp.ncbi.nlm.nih.gov/geo/series/GSE223nnn/GSE223245/soft/GSE223245_family.soft.gz"
    matrix_path = download_file(matrix_url, RAWDIR / "GSE223245_series_matrix.txt.gz")
    soft_path = download_file(soft_url, RAWDIR / "GSE223245_family.soft.gz")
    probe_expr, sample_meta = parse_series_matrix(matrix_path)
    platform = parse_platform_table(soft_path)
    metadata = build_gse223245_metadata(sample_meta)
    expr, probe_map = collapse_gse223245(probe_expr, platform)
    expr = expr.loc[:, metadata["sample_id"].tolist()]
    focus = expr.loc[[g for g in PRIORITY_GENES if g in expr.index]].copy()

    rows = []
    for gene in PRIORITY_GENES:
        if gene not in expr.index:
            rows.append({"gene_symbol": gene, "available": False})
            continue
        vals = expr.loc[gene, metadata["sample_id"]].astype(float)
        all_r, all_p = stats.spearmanr(metadata["severity_score"], vals, nan_policy="omit")
        tbi_meta = metadata[metadata["group"].isin(["Mild", "Moderate", "Severe"])].copy()
        tbi_vals = vals[tbi_meta["sample_id"]]
        tbi_r, tbi_p = stats.spearmanr(tbi_meta["severity_score"], tbi_vals, nan_policy="omit")
        row = {
            "gene_symbol": gene,
            "available": True,
            "spearman_r_severity_all_groups": all_r,
            "p_value_severity_all_groups": all_p,
            "spearman_r_TBI_only": tbi_r,
            "p_value_TBI_only": tbi_p,
        }
        for group in ["Control", "Mild", "Moderate", "Severe"]:
            group_ids = metadata.loc[metadata["group"].eq(group), "sample_id"].tolist()
            row[f"mean_{group}"] = float(vals[group_ids].mean()) if group_ids else np.nan
        for group in ["Mild", "Moderate", "Severe"]:
            case = vals[metadata.loc[metadata["group"].eq(group), "sample_id"]].astype(float)
            ctrl = vals[metadata.loc[metadata["group"].eq("Control"), "sample_id"]].astype(float)
            if len(case) > 1 and len(ctrl) > 1:
                t, p = stats.ttest_ind(case, ctrl, equal_var=False)
            else:
                p = np.nan
            row[f"logFC_{group}_vs_Control"] = float(case.mean() - ctrl.mean())
            row[f"p_value_{group}_vs_Control"] = float(p) if np.isfinite(p) else np.nan
        rows.append(row)
    gene_stats = pd.DataFrame(rows)
    gene_stats["FDR_severity_all_groups"] = bh_fdr(gene_stats["p_value_severity_all_groups"])
    gene_stats["FDR_TBI_only"] = bh_fdr(gene_stats["p_value_TBI_only"])

    modules = pd.DataFrame(
        {
            "sample_id": metadata["sample_id"],
            "disulfidptosis_8gene": module_score(focus, [g for g in PRIORITY_GENES if g in focus.index]).reindex(metadata["sample_id"]).values,
            "transporter_2gene": module_score(focus, [g for g in TRANSPORTER_GENES if g in focus.index]).reindex(metadata["sample_id"]).values,
            "actin_6gene": module_score(focus, [g for g in ACTIN_GENES if g in focus.index]).reindex(metadata["sample_id"]).values,
        }
    ).merge(metadata, on="sample_id", how="left")
    module_rows = []
    for mod in ["disulfidptosis_8gene", "transporter_2gene", "actin_6gene"]:
        all_r, all_p = stats.spearmanr(modules["severity_score"], modules[mod], nan_policy="omit")
        tbi = modules[modules["group"].isin(["Mild", "Moderate", "Severe"])]
        tbi_r, tbi_p = stats.spearmanr(tbi["severity_score"], tbi[mod], nan_policy="omit")
        row = {
            "module": mod,
            "spearman_r_severity_all_groups": all_r,
            "p_value_severity_all_groups": all_p,
            "spearman_r_TBI_only": tbi_r,
            "p_value_TBI_only": tbi_p,
        }
        for group in ["Control", "Mild", "Moderate", "Severe"]:
            row[f"mean_{group}"] = modules.loc[modules["group"].eq(group), mod].mean()
        for group in ["Mild", "Moderate", "Severe"]:
            case = modules.loc[modules["group"].eq(group), mod]
            ctrl = modules.loc[modules["group"].eq("Control"), mod]
            t, p = stats.ttest_ind(case, ctrl, equal_var=False)
            row[f"delta_{group}_vs_Control"] = case.mean() - ctrl.mean()
            row[f"p_value_{group}_vs_Control"] = p
        module_rows.append(row)
    module_stats = pd.DataFrame(module_rows)
    module_stats["FDR_severity_all_groups"] = bh_fdr(module_stats["p_value_severity_all_groups"])
    module_stats["FDR_TBI_only"] = bh_fdr(module_stats["p_value_TBI_only"])

    focus_long = focus.T.reset_index().rename(columns={"index": "sample_id"}).melt(
        id_vars="sample_id", var_name="gene_symbol", value_name="expression"
    )
    focus_long = focus_long.merge(metadata, on="sample_id", how="left")
    pca_mat = zscore_rows(focus).T.dropna(axis=1, how="all").fillna(0)
    x = pca_mat.to_numpy(dtype=float)
    x = x - x.mean(axis=0, keepdims=True)
    u, s, vt = np.linalg.svd(x, full_matrices=False)
    var = s**2 / np.sum(s**2)
    pca = pd.DataFrame({"sample_id": pca_mat.index, "PC1": u[:, 0] * s[0], "PC2": u[:, 1] * s[1]})
    pca["PC1_var"] = var[0]
    pca["PC2_var"] = var[1] if len(var) > 1 else np.nan
    pca = pca.merge(metadata, on="sample_id", how="left")

    metadata.to_csv(TABLEDIR / "v3_GSE223245_sample_metadata_20260604.csv", index=False, encoding="utf-8-sig")
    probe_map.to_csv(TABLEDIR / "v3_GSE223245_probe_gene_map_20260604.csv", index=False, encoding="utf-8-sig")
    gene_stats.to_csv(TABLEDIR / "v3_GSE223245_8gene_severity_results_20260604.csv", index=False, encoding="utf-8-sig")
    module_stats.to_csv(TABLEDIR / "v3_GSE223245_module_severity_results_20260604.csv", index=False, encoding="utf-8-sig")
    modules.to_csv(TABLEDIR / "v3_GSE223245_module_scores_long_20260604.csv", index=False, encoding="utf-8-sig")
    focus_long.to_csv(TABLEDIR / "v3_GSE223245_8gene_expression_long_20260604.csv", index=False, encoding="utf-8-sig")
    pca.to_csv(TABLEDIR / "v3_GSE223245_8gene_PCA_20260604.csv", index=False, encoding="utf-8-sig")
    return {
        "metadata": metadata,
        "gene_stats": gene_stats,
        "module_stats": module_stats,
        "modules": modules,
        "focus_long": focus_long,
        "pca": pca,
    }


def load_human_focus() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    focus = pd.read_csv(find_file("focused_human_brain_gene_results.csv"))
    focus8 = focus[focus["gene_symbol"].isin(PRIORITY_GENES)].copy()
    focus8["effect"] = np.where(focus8["stage_correlation"].notna(), focus8["stage_correlation"], focus8["logFC"])
    focus8["comparison_short"] = (
        focus8["comparison"]
        .str.replace("GSE104687_human_brain_", "104687 ", regex=False)
        .str.replace("_TBI_vs_Control", "", regex=False)
        .str.replace("GSE209552_human_acute_severe_TBI_vs_Control", "209552 acute severe TBI", regex=False)
        .str.replace("GSE193407_human_prefrontal_BA9_", "193407 ", regex=False)
        .str.replace("GSE319253_human_superior_frontal_cortex_", "319253 ", regex=False)
        .str.replace("_", " ", regex=False)
    )
    gene_summary = (
        focus8.groupby("gene_symbol", as_index=False)
        .agg(
            human_nominal_supported=("significant_nominal_0_05", "sum"),
            human_fdr_supported=("significant_FDR_0_05", "sum"),
            n_human_comparisons=("comparison", "nunique"),
            mean_effect=("effect", "mean"),
            max_abs_effect=("effect", lambda x: float(np.nanmax(np.abs(x)))),
        )
        .sort_values(["human_fdr_supported", "human_nominal_supported", "max_abs_effect"], ascending=False)
    )
    cte = focus8[focus8["comparison"].str.contains("GSE193407", na=False)].copy()
    focus8.to_csv(TABLEDIR / "v3_human_brain_8gene_focused_results_20260604.csv", index=False, encoding="utf-8-sig")
    gene_summary.to_csv(TABLEDIR / "v3_human_brain_8gene_evidence_summary_20260604.csv", index=False, encoding="utf-8-sig")
    return focus8, gene_summary, cte


def load_gse209552() -> dict[str, pd.DataFrame]:
    summary = pd.read_csv(find_file("GSE209552_target_gene_log2CPM_summary.csv", contains="06_BHMT"))
    long = pd.read_csv(find_file("GSE209552_target_gene_log2CPM_long.csv", contains="06_BHMT"))
    module_stats = pd.read_csv(find_file("GSE209552_module_score_statistics.csv", contains="06_BHMT"))
    module_long_candidates = sorted(Path.cwd().rglob("GSE209552_module_scores_long.csv"))
    module_long = pd.read_csv(module_long_candidates[0]) if module_long_candidates else pd.DataFrame()
    running = pd.read_csv(find_file("GSE209552_gene_set_running_enrichment.csv", contains="06_BHMT"))
    focus_summary = summary[summary["gene_symbol"].isin(PRIORITY_GENES)].copy()

    raw = find_file("GSE209552_TBI_gene_count_matrix_2.csv.gz")
    counts = pd.read_csv(raw, sep="\t", compression="gzip", comment="#")
    sample_cols = [c for c in counts.columns if c.startswith("DA")]
    counts = counts.rename(columns={"Geneid": "gene_symbol"})
    counts["gene_symbol"] = counts["gene_symbol"].map(clean_gene)
    counts = counts[counts["gene_symbol"].ne("")]
    counts[sample_cols] = counts[sample_cols].apply(pd.to_numeric, errors="coerce")
    counts["mean_count"] = counts[sample_cols].mean(axis=1)
    counts = counts.sort_values("mean_count", ascending=False).drop_duplicates("gene_symbol").set_index("gene_symbol")[sample_cols]
    lib = counts.sum(axis=0).replace(0, np.nan)
    expr = np.log2(counts.div(lib, axis=1) * 1_000_000 + 1)
    sample_meta = pd.DataFrame({"sample_id": sample_cols})
    sample_meta["condition"] = np.where(sample_meta["sample_id"].str.contains("_ctrl_", case=False), "Control", "TBI")
    sample_scores = pd.DataFrame({"sample_id": sample_cols, "condition": sample_meta["condition"]})
    sample_scores["score_disulfidptosis_8gene"] = module_score(expr, PRIORITY_GENES).reindex(sample_cols).values
    sample_scores["score_transporter_2gene"] = module_score(expr, TRANSPORTER_GENES).reindex(sample_cols).values
    sample_scores["score_actin_6gene"] = module_score(expr, ACTIN_GENES).reindex(sample_cols).values
    for ctype, markers in CELL_MARKERS.items():
        sample_scores[f"marker_{ctype}"] = module_score(expr, markers).reindex(sample_cols).values
    corr_rows = []
    for module in ["score_disulfidptosis_8gene", "score_transporter_2gene", "score_actin_6gene"]:
        for ctype in CELL_MARKERS:
            marker_col = f"marker_{ctype}"
            valid = sample_scores[[module, marker_col]].dropna()
            if len(valid) >= 5:
                r, p = stats.spearmanr(valid[module], valid[marker_col])
            else:
                r, p = np.nan, np.nan
            corr_rows.append({"module": module, "celltype_proxy": ctype, "spearman_r": r, "p_value": p, "n_samples": len(valid)})
    marker_corr = pd.DataFrame(corr_rows)
    marker_corr["FDR"] = bh_fdr(marker_corr["p_value"])
    focus_summary.to_csv(TABLEDIR / "v3_GSE209552_8gene_acute_severe_results_20260604.csv", index=False, encoding="utf-8-sig")
    sample_scores.to_csv(TABLEDIR / "v3_GSE209552_bulk_marker_proxy_scores_20260604.csv", index=False, encoding="utf-8-sig")
    marker_corr.to_csv(TABLEDIR / "v3_GSE209552_bulk_marker_proxy_correlations_20260604.csv", index=False, encoding="utf-8-sig")
    return {
        "summary": focus_summary,
        "long": long[long["gene_symbol"].isin(PRIORITY_GENES)].copy(),
        "module_stats": module_stats,
        "module_long": module_long,
        "running": running,
        "marker_scores": sample_scores,
        "marker_corr": marker_corr,
    }


def load_gse163415() -> dict[str, pd.DataFrame]:
    focused = pd.read_csv(find_file("GSE163415_disulfidptosis_focused_DE_results_20260517.csv"))
    expr_long = pd.read_csv(find_file("GSE163415_disulfidptosis_expression_long_20260517.csv"))
    module_summary = pd.read_csv(find_file("GSE163415_disulfidptosis_module_score_summary_20260517.csv"))
    pathway = pd.read_csv(find_file("GSE163415_disulfidptosis_pathway_overlap_20260517.csv"))
    priority = pd.read_csv(find_file("GSE163415_disulfidptosis_priority_genes_detailed_20260518.csv"))
    focused8 = focused[focused["gene_symbol"].str.upper().isin(PRIORITY_GENES)].copy()
    expr8 = expr_long[expr_long["gene_symbol"].str.upper().isin(PRIORITY_GENES)].copy()
    focused8.to_csv(TABLEDIR / "v3_GSE163415_8gene_focused_DE_results_20260604.csv", index=False, encoding="utf-8-sig")
    return {"focused": focused8, "expr_long": expr8, "module": module_summary, "pathway": pathway, "priority": priority}


def integrated_priority(human_summary: pd.DataFrame, gse209: dict[str, pd.DataFrame], gse163: dict[str, pd.DataFrame], gse223: dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    g209 = gse209["summary"].set_index("gene_symbol")
    mouse = gse163["focused"].copy()
    sev = gse223["gene_stats"].set_index("gene_symbol")
    for gene in PRIORITY_GENES:
        h = human_summary[human_summary["gene_symbol"].eq(gene)]
        h_fdr = int(h["human_fdr_supported"].iloc[0]) if not h.empty else 0
        h_nom = int(h["human_nominal_supported"].iloc[0]) if not h.empty else 0
        acute_logfc = float(g209.loc[gene, "GSE209552_logFC_TBI_vs_Control"]) if gene in g209.index else np.nan
        acute_p = float(g209.loc[gene, "GSE209552_p_value_ttest"]) if gene in g209.index else np.nan
        m = mouse[mouse["gene_symbol"].str.upper().eq(gene)]
        mouse_nom = int((m["p_value"] < 0.05).sum()) if "p_value" in m.columns and not m.empty else 0
        mouse_fdr = int((m["FDR"] < 0.05).sum()) if "FDR" in m.columns and not m.empty else 0
        best_mouse = ""
        if not m.empty and "p_value" in m.columns:
            idx = pd.to_numeric(m["p_value"], errors="coerce").idxmin()
            best = m.loc[idx]
            best_mouse = f"{best.get('time', '')}/{best.get('region', '')}/{best.get('treatment', '')}"
        severity_r = float(sev.loc[gene, "spearman_r_severity_all_groups"]) if gene in sev.index else np.nan
        severity_p = float(sev.loc[gene, "p_value_severity_all_groups"]) if gene in sev.index else np.nan
        score = h_fdr * 3 + h_nom * 1 + mouse_fdr * 2 + mouse_nom * 0.5
        if np.isfinite(acute_p) and acute_p < 0.05:
            score += 1
        if np.isfinite(severity_p) and severity_p < 0.05:
            score += 0.5
        rows.append(
            {
                "gene_symbol": gene,
                "human_FDR_comparisons": h_fdr,
                "human_nominal_comparisons": h_nom,
                "GSE209552_acute_logFC": acute_logfc,
                "GSE209552_acute_p": acute_p,
                "mouse_FDR_units": mouse_fdr,
                "mouse_nominal_units": mouse_nom,
                "best_mouse_unit": best_mouse,
                "GSE223245_severity_r_all": severity_r,
                "GSE223245_severity_p_all": severity_p,
                "integrated_evidence_score": score,
            }
        )
    out = pd.DataFrame(rows).sort_values("integrated_evidence_score", ascending=False)
    out.to_csv(TABLEDIR / "v3_integrated_8gene_priority_20260604.csv", index=False, encoding="utf-8-sig")
    return out


def make_fig1(evidence: pd.DataFrame) -> None:
    fig = plt.figure(figsize=(10.8, 7.2))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.0, 1.1], width_ratios=[1.15, 1.0], hspace=0.42, wspace=0.35)
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])

    ax_a.axis("off")
    boxes = [
        ("Human acute severe TBI\nGSE209552", 0.05, 0.62, OKABE["red"]),
        ("Human chronic CTE stage\nGSE193407/GSE319253", 0.52, 0.62, OKABE["blue"]),
        ("Mouse CCI time-region\nGSE163415", 0.05, 0.18, OKABE["green"]),
        ("Peripheral severity\nGSE223245", 0.52, 0.18, OKABE["orange"]),
    ]
    for text, x, y, color in boxes:
        rect = mpl.patches.FancyBboxPatch((x, y), 0.4, 0.24, boxstyle="round,pad=0.02,rounding_size=0.015", fc=color, ec="black", alpha=0.13, lw=0.8)
        ax_a.add_patch(rect)
        ax_a.text(x + 0.2, y + 0.12, text, ha="center", va="center", fontsize=8.3)
    ax_a.annotate("", xy=(0.52, 0.74), xytext=(0.45, 0.74), arrowprops=dict(arrowstyle="->", lw=0.9))
    ax_a.annotate("", xy=(0.52, 0.30), xytext=(0.45, 0.30), arrowprops=dict(arrowstyle="->", lw=0.9))
    ax_a.text(0.5, 0.5, "8 fixed genes\nSLC3A2/SLC7A11 + actin endpoints", ha="center", va="center", fontsize=9.2, fontweight="bold")
    ax_a.set_xlim(0, 1)
    ax_a.set_ylim(0, 1)
    ax_a.set_title("Evidence-layer design")
    add_panel_label(ax_a, "A")

    layer_counts = evidence["evidence_axis"].value_counts().reindex(
        ["acute_time_window", "remote_region_window", "CTE_course_stage", "CTE_external_validation", "animal_time_region", "peripheral_severity", "wet_validation"]
    )
    ax_b.barh(np.arange(len(layer_counts)), layer_counts.fillna(0), color=[OKABE["red"], OKABE["gray"], OKABE["blue"], OKABE["sky"], OKABE["green"], OKABE["orange"], OKABE["purple"]])
    ax_b.set_yticks(np.arange(len(layer_counts)))
    ax_b.set_yticklabels(["acute", "remote region", "CTE course", "CTE external", "mouse CCI", "blood severity", "planned validation"])
    ax_b.set_xlabel("Evidence entries")
    ax_b.set_title("Time, region and severity axes")
    sns.despine(ax=ax_b)
    add_panel_label(ax_b, "B")

    gene_map = pd.DataFrame(
        {
            "gene": PRIORITY_GENES,
            "axis": ["Transporter", "Transporter", "WRC/remodeling", "Adhesion/tension", "Actin scaffold", "Myosin/tension", "Myosin light chain", "Filamin scaffold"],
            "mechanistic_position": [1, 1, 2, 2, 3, 3, 3, 3],
        }
    )
    colors = {1: OKABE["red"], 2: OKABE["blue"], 3: OKABE["green"]}
    ax_c.scatter(gene_map["mechanistic_position"], np.arange(len(gene_map)), s=160, c=[colors[x] for x in gene_map["mechanistic_position"]], edgecolor="black")
    for i, row in gene_map.iterrows():
        ax_c.text(row["mechanistic_position"] + 0.05, i, f"{row['gene']}  {row['axis']}", va="center", fontsize=8)
    ax_c.set_yticks([])
    ax_c.set_xticks([1, 2, 3])
    ax_c.set_xticklabels(["cystine import", "actin control", "cytoskeletal endpoint"])
    ax_c.set_xlim(0.75, 3.85)
    ax_c.set_title("Pre-fixed transporter-actin panel")
    sns.despine(ax=ax_c, left=True)
    add_panel_label(ax_c, "C")

    guardrail = pd.DataFrame(
        {
            "Claim": ["Human brain FDR", "Nominal signal", "Mouse CCI", "Peripheral blood", "Bulk marker proxy"],
            "Allowed wording": ["association", "candidate clue", "supportive evidence", "severity context", "prioritization"],
            "Avoid": ["causality", "biomarker claim", "human proof", "brain mechanism", "cell-type proof"],
        }
    )
    ax_d.axis("off")
    tbl = ax_d.table(cellText=guardrail.values, colLabels=guardrail.columns, loc="center", cellLoc="left", colLoc="left")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(7.2)
    tbl.scale(1.0, 1.45)
    for (r, c), cell in tbl.get_celld().items():
        cell.set_edgecolor("#D1D5DB")
        if r == 0:
            cell.set_facecolor("#F3F4F6")
            cell.set_text_props(weight="bold")
    ax_d.set_title("Interpretation guardrails")
    add_panel_label(ax_d, "D")

    savefig(fig, "Fig1_v3_study_design_evidence_layers_20260604")


def make_fig2_human(focus8: pd.DataFrame, human_summary: pd.DataFrame) -> None:
    fig = plt.figure(figsize=(11.5, 8.2))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.2, 1.0], width_ratios=[1.35, 1.0], hspace=0.48, wspace=0.35)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0])
    ax_c = fig.add_subplot(gs[1, 1])

    selected_order = [
        "GSE209552_human_acute_brain_severe_TBI_vs_Control",
        "GSE193407_human_prefrontal_BA9_CTE_stage_trend",
        "GSE193407_human_prefrontal_BA9_late_CTE_stage3_4_vs_stage0",
        "GSE319253_human_superior_frontal_cortex_CTE_vs_Control",
        "GSE104687_human_brain_FWM_TBI_vs_Control",
        "GSE104687_human_brain_HIP_TBI_vs_Control",
        "GSE104687_human_brain_PCx_TBI_vs_Control",
        "GSE104687_human_brain_TCx_TBI_vs_Control",
    ]
    heat = focus8[focus8["comparison"].isin(selected_order)].pivot_table(index="comparison", columns="gene_symbol", values="effect", aggfunc="mean").reindex(selected_order)[PRIORITY_GENES]
    labels = [
        "Acute severe TBI",
        "CTE stage trend",
        "Late CTE vs stage0",
        "External CTE",
        "Remote FWM",
        "Remote HIP",
        "Remote PCx",
        "Remote TCx",
    ]
    sns.heatmap(heat, cmap="RdBu_r", center=0, linewidths=0.3, linecolor="white", ax=ax_a, cbar_kws={"label": "logFC or stage r"})
    ax_a.set_yticklabels(labels, rotation=0)
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    ax_a.set_title("Human brain 8-gene effect map across acute, chronic and regional comparisons")
    for y, comp in enumerate(selected_order):
        sub = focus8[focus8["comparison"].eq(comp)].set_index("gene_symbol")
        for x, gene in enumerate(PRIORITY_GENES):
            if gene in sub.index:
                row = sub.loc[gene]
                marker = "●" if bool(row.get("significant_FDR_0_05", False)) else ("○" if bool(row.get("significant_nominal_0_05", False)) else "")
                if marker:
                    ax_a.text(x + 0.5, y + 0.5, marker, ha="center", va="center", fontsize=8, color="black")
    add_panel_label(ax_a, "A")

    cte_trend = focus8[focus8["comparison"].eq("GSE193407_human_prefrontal_BA9_CTE_stage_trend")].set_index("gene_symbol").reindex(PRIORITY_GENES)
    ax_b.bar(np.arange(len(PRIORITY_GENES)), cte_trend["stage_correlation"], color=[OKABE["blue"] if f else "#BFC7D5" for f in cte_trend["significant_FDR_0_05"].fillna(False)])
    ax_b.axhline(0, color="black", lw=0.8)
    for i, gene in enumerate(PRIORITY_GENES):
        p = cte_trend.loc[gene, "FDR"] if gene in cte_trend.index else np.nan
        ax_b.text(i, cte_trend.loc[gene, "stage_correlation"] + 0.025, p_to_star(p), ha="center", fontsize=9)
    ax_b.set_xticks(np.arange(len(PRIORITY_GENES)))
    ax_b.set_xticklabels(PRIORITY_GENES, rotation=45, ha="right")
    ax_b.set_ylabel("Stage correlation (r)")
    ax_b.set_title("GSE193407 continuous CTE stage")
    sns.despine(ax=ax_b)
    add_panel_label(ax_b, "B")

    summ = human_summary.set_index("gene_symbol").reindex(PRIORITY_GENES).reset_index()
    ax_c.scatter(summ["human_nominal_supported"], summ["gene_symbol"], s=(summ["human_fdr_supported"] + 1) * 95, c=summ["max_abs_effect"], cmap="viridis", edgecolor="black")
    ax_c.set_xlabel("Nominally supported human comparisons")
    ax_c.set_ylabel("")
    ax_c.set_title("Cross-human evidence recurrence")
    cbar = fig.colorbar(ax_c.collections[0], ax=ax_c, shrink=0.82)
    cbar.set_label("Max |effect|")
    sns.despine(ax=ax_c)
    add_panel_label(ax_c, "C")

    savefig(fig, "Fig2_v3_human_brain_multidataset_8gene_20260604")


def make_fig3_gse209552(gse209: dict[str, pd.DataFrame]) -> None:
    fig = plt.figure(figsize=(11.5, 8.4))
    gs = fig.add_gridspec(2, 2, hspace=0.45, wspace=0.34)
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])

    module_long = gse209["module_long"].copy()
    if not module_long.empty:
        plot = module_long[module_long["analysis_set"].eq("1_双硫死亡相关")].copy()
        order = ["双硫死亡_核心转运与应激", "双硫死亡_肌动蛋白骨架靶蛋白", "双硫死亡_线粒体协同命中"]
        plot = plot[plot["submodule"].isin(order)]
        sns.boxplot(data=plot, x="submodule", y="module_score_zmean", hue="condition", order=order, ax=ax_a, palette={"Control": "#BFC7D5", "TBI": OKABE["red"]}, fliersize=0)
        sns.stripplot(data=plot, x="submodule", y="module_score_zmean", hue="condition", order=order, ax=ax_a, dodge=True, color="black", size=3.2, alpha=0.65)
        ax_a.legend_.remove()
        ax_a.set_xticklabels(["Transporter/stress", "Actin targets", "Mito synergy"], rotation=20, ha="right")
        ax_a.set_ylabel("Z-mean module score")
        ax_a.set_xlabel("")
        ax_a.set_title("Acute severe TBI module scores")
    add_panel_label(ax_a, "A")

    summary = gse209["summary"].set_index("gene_symbol").reindex(PRIORITY_GENES).reset_index()
    ax_b.bar(np.arange(len(summary)), summary["GSE209552_logFC_TBI_vs_Control"], color=[OKABE["red"] if p < 0.05 else "#BFC7D5" for p in summary["GSE209552_p_value_ttest"].fillna(1)])
    ax_b.axhline(0, color="black", lw=0.8)
    for i, row in summary.iterrows():
        ax_b.text(i, row["GSE209552_logFC_TBI_vs_Control"] + 0.08 * np.sign(row["GSE209552_logFC_TBI_vs_Control"] or 1), p_to_star(row["GSE209552_p_value_ttest"]), ha="center", fontsize=9)
    ax_b.set_xticks(np.arange(len(summary)))
    ax_b.set_xticklabels(summary["gene_symbol"], rotation=45, ha="right")
    ax_b.set_ylabel("log2CPM logFC")
    ax_b.set_title("8-gene acute severe TBI direction")
    sns.despine(ax=ax_b)
    add_panel_label(ax_b, "B")

    running = gse209["running"].copy()
    if {"rank_index", "running_ES"}.issubset(running.columns):
        candidates = running["gene_set"].dropna().unique().tolist() if "gene_set" in running.columns else []
        if candidates:
            target = candidates[0]
            if any("双硫死亡" in str(c) for c in candidates):
                target = [c for c in candidates if "双硫死亡" in str(c)][0]
            r = running[running["gene_set"].eq(target)] if "gene_set" in running.columns else running
        else:
            r = running
        ax_c.plot(r["rank_index"], r["running_ES"], color=OKABE["blue"], lw=1.8)
        ax_c.axhline(0, color="black", lw=0.7)
        ax_c.set_xlabel("Ranked genes")
        ax_c.set_ylabel("Running enrichment score")
        ax_c.set_title("Disulfidptosis-ranked enrichment")
    elif {"analysis_set", "enrichment_score", "FDR"}.issubset(running.columns):
        r = running.sort_values("enrichment_score", ascending=True).copy()
        colors = [OKABE["blue"] if f < 0.05 else "#BFC7D5" for f in r["FDR"]]
        ax_c.barh(np.arange(len(r)), r["enrichment_score"], color=colors, edgecolor="black", lw=0.4)
        ax_c.axvline(0, color="black", lw=0.8)
        ax_c.set_yticks(np.arange(len(r)))
        ax_c.set_yticklabels(
            r["analysis_set"]
            .astype(str)
            .str.replace("1_双硫死亡相关", "Disulfidptosis")
            .str.replace("2_铜代谢及铜死亡相关", "Copper/cuproptosis")
            .str.replace("3_甲硫氨酸代谢相关", "Methionine metabolism")
            .str.replace("4_硫化氢代谢酶", "H2S enzymes")
        )
        for i, row in r.reset_index(drop=True).iterrows():
            ax_c.text(row["enrichment_score"] + 0.03 * np.sign(row["enrichment_score"] or 1), i, p_to_star(row["FDR"]), va="center", fontsize=9)
        ax_c.set_xlabel("Enrichment score")
        ax_c.set_title("Ranked gene-set enrichment summary")
    else:
        ax_c.axis("off")
    sns.despine(ax=ax_c)
    add_panel_label(ax_c, "C")

    corr = gse209["marker_corr"].pivot_table(index="module", columns="celltype_proxy", values="spearman_r")
    corr = corr.reindex(["score_disulfidptosis_8gene", "score_transporter_2gene", "score_actin_6gene"])
    sns.heatmap(corr, cmap="RdBu_r", center=0, vmin=-1, vmax=1, ax=ax_d, linewidths=0.3, cbar_kws={"label": "Spearman r"})
    ax_d.set_yticklabels(["8-gene", "Transporter", "Actin"], rotation=0)
    ax_d.set_xlabel("")
    ax_d.set_ylabel("")
    ax_d.set_title("Bulk marker-proxy cell-type prioritization")
    add_panel_label(ax_d, "D")

    savefig(fig, "Fig3_v3_GSE209552_acute_severe_and_marker_proxy_20260604")


def make_fig4_mouse(gse163: dict[str, pd.DataFrame]) -> None:
    fig = plt.figure(figsize=(11.5, 8.4))
    gs = fig.add_gridspec(2, 2, hspace=0.62, wspace=0.34)
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])

    module = gse163["module"].copy()
    module["unit"] = module["time"].astype(str) + " " + module["region"].astype(str) + " " + module["treatment"].astype(str)
    heat = module.pivot_table(index="unit", columns="module", values="delta_module_score_TBI_minus_NoTBI", aggfunc="mean")
    preferred_cols = [c for c in ["Core transporter/stress", "Actin cytoskeleton targets", "Mitochondrial synergy"] if c in heat.columns]
    sns.heatmap(heat[preferred_cols], cmap="RdBu_r", center=0, ax=ax_a, linewidths=0.3, cbar_kws={"label": "TBI-NoTBI delta"})
    short_cols = {"Core transporter/stress": "Transporter", "Actin cytoskeleton targets": "Actin", "Mitochondrial synergy": "Mito"}
    ax_a.set_xticklabels([short_cols.get(t.get_text(), t.get_text()) for t in ax_a.get_xticklabels()], rotation=20, ha="right")
    ax_a.set_title("Mouse CCI module score shifts")
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    add_panel_label(ax_a, "A")

    focused = gse163["focused"].copy()
    focused["unit"] = focused["time"].astype(str) + " " + focused["region"].astype(str) + " " + focused["treatment"].astype(str)
    effect_col = "logFC_TBI_vs_NoTBI" if "logFC_TBI_vs_NoTBI" in focused.columns else "logFC"
    gene_heat = focused.pivot_table(index="unit", columns="gene_symbol", values=effect_col, aggfunc="mean")
    gene_heat = gene_heat.reindex(columns=[g for g in PRIORITY_GENES if g in gene_heat.columns])
    sns.heatmap(gene_heat, cmap="RdBu_r", center=0, ax=ax_b, linewidths=0.3, cbar_kws={"label": "logFC"})
    ax_b.set_title("8-gene spatiotemporal logFC")
    ax_b.set_xlabel("")
    ax_b.set_ylabel("")
    add_panel_label(ax_b, "B")

    pathway = gse163["pathway"].copy()
    fdr_col = "FDR_pathway_overlap" if "FDR_pathway_overlap" in pathway.columns else "FDR"
    pathway = (
        pathway.sort_values(fdr_col)
        .groupby("pathway", as_index=False)
        .agg(best_FDR=(fdr_col, "min"), best_unit=("analysis_unit", "first"), max_FDR_hits=("FDR_hits", "max"))
        .sort_values("best_FDR")
        .head(8)
    )
    ax_c.barh(np.arange(len(pathway)), -np.log10(pathway["best_FDR"].astype(float)), color=OKABE["green"], edgecolor="black", lw=0.4)
    ax_c.set_yticks(np.arange(len(pathway)))
    ax_c.set_yticklabels(pathway["pathway"].astype(str).str.replace("_", " ").str.wrap(34).str.replace("\n", "\n"))
    ax_c.invert_yaxis()
    ax_c.set_xlabel("-log10(FDR)")
    ax_c.set_title("Pathway overlap in mouse CCI")
    sns.despine(ax=ax_c)
    add_panel_label(ax_c, "C")

    nominal_count = (
        focused.groupby("gene_symbol")["p_value"]
        .apply(lambda s: int((pd.to_numeric(s, errors="coerce") < 0.05).sum()))
        .reindex(PRIORITY_GENES)
        .fillna(0)
    )
    ax_d.barh(np.arange(len(PRIORITY_GENES)), nominal_count.values, color=OKABE["purple"], edgecolor="black", lw=0.4)
    ax_d.set_yticks(np.arange(len(PRIORITY_GENES)))
    ax_d.set_yticklabels(PRIORITY_GENES)
    ax_d.invert_yaxis()
    ax_d.set_xlabel("Nominally supported CCI units")
    ax_d.set_title("Mouse-derived 8-gene priority")
    sns.despine(ax=ax_d)
    add_panel_label(ax_d, "D")

    savefig(fig, "Fig4_v3_GSE163415_mouse_CCI_spatiotemporal_suite_20260604")


def make_fig5_severity(gse223: dict[str, pd.DataFrame]) -> None:
    fig = plt.figure(figsize=(11.5, 8.4))
    gs = fig.add_gridspec(2, 2, hspace=0.46, wspace=0.34)
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])

    pca = gse223["pca"].copy()
    order = ["Control", "Mild", "Moderate", "Severe"]
    colors = {"Control": "#8D99AE", "Mild": OKABE["sky"], "Moderate": OKABE["orange"], "Severe": OKABE["red"]}
    for group in order:
        sub = pca[pca["group"].eq(group)]
        ax_a.scatter(sub["PC1"], sub["PC2"], label=group, s=55, color=colors[group], edgecolor="black", lw=0.4)
    ax_a.set_xlabel(f"PC1 ({pca['PC1_var'].iloc[0]*100:.1f}%)")
    ax_a.set_ylabel(f"PC2 ({pca['PC2_var'].iloc[0]*100:.1f}%)")
    ax_a.set_title("GSE223245 8-gene PCA (whole blood)")
    ax_a.legend(frameon=False, ncol=2)
    sns.despine(ax=ax_a)
    add_panel_label(ax_a, "A")

    modules = gse223["modules"].melt(
        id_vars=["sample_id", "group", "severity_score"],
        value_vars=["disulfidptosis_8gene", "transporter_2gene", "actin_6gene"],
        var_name="module",
        value_name="module_score",
    )
    sns.boxplot(data=modules, x="group", y="module_score", hue="module", order=order, ax=ax_b, fliersize=0, palette=[OKABE["blue"], OKABE["red"], OKABE["green"]])
    sns.stripplot(data=modules, x="group", y="module_score", hue="module", order=order, dodge=True, color="black", size=2.7, alpha=0.55, ax=ax_b)
    ax_b.legend_.remove()
    ax_b.set_xlabel("")
    ax_b.set_ylabel("Z-mean score")
    ax_b.set_title("Peripheral severity module score")
    sns.despine(ax=ax_b)
    add_panel_label(ax_b, "B")

    gene_stats = gse223["gene_stats"].set_index("gene_symbol").reindex(PRIORITY_GENES).reset_index()
    ax_c.bar(np.arange(len(gene_stats)), gene_stats["spearman_r_severity_all_groups"], color=[OKABE["orange"] if p < 0.05 else "#BFC7D5" for p in gene_stats["p_value_severity_all_groups"].fillna(1)])
    ax_c.axhline(0, color="black", lw=0.8)
    for i, row in gene_stats.iterrows():
        ax_c.text(i, row["spearman_r_severity_all_groups"] + 0.04 * np.sign(row["spearman_r_severity_all_groups"] or 1), p_to_star(row["p_value_severity_all_groups"]), ha="center")
    ax_c.set_xticks(np.arange(len(gene_stats)))
    ax_c.set_xticklabels(gene_stats["gene_symbol"], rotation=45, ha="right")
    ax_c.set_ylabel("Spearman r, severity 0-3")
    ax_c.set_title("8-gene peripheral severity trend")
    sns.despine(ax=ax_c)
    add_panel_label(ax_c, "C")

    module_stats = gse223["module_stats"].set_index("module").reindex(["disulfidptosis_8gene", "transporter_2gene", "actin_6gene"]).reset_index()
    ax_d.bar(np.arange(len(module_stats)), module_stats["spearman_r_severity_all_groups"], color=[OKABE["blue"], OKABE["red"], OKABE["green"]], edgecolor="black", lw=0.5)
    ax_d.axhline(0, color="black", lw=0.8)
    ax_d.set_xticks(np.arange(len(module_stats)))
    ax_d.set_xticklabels(["8-gene", "Transporter", "Actin"], rotation=20, ha="right")
    ax_d.set_ylabel("Spearman r, severity 0-3")
    ax_d.set_title("Module-level severity association")
    sns.despine(ax=ax_d)
    add_panel_label(ax_d, "D")

    savefig(fig, "Fig5_v3_GSE223245_peripheral_severity_suite_20260604")


def make_fig6_integrated(priority: pd.DataFrame, gse209: dict[str, pd.DataFrame]) -> None:
    fig = plt.figure(figsize=(12.4, 8.8))
    gs = fig.add_gridspec(2, 2, hspace=0.62, wspace=0.58, width_ratios=[1.0, 1.12])
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])

    metrics = priority.set_index("gene_symbol")[
        ["human_FDR_comparisons", "human_nominal_comparisons", "mouse_nominal_units", "GSE209552_acute_logFC", "GSE223245_severity_r_all"]
    ].reindex(PRIORITY_GENES)
    metrics = metrics.rename(
        columns={
            "human_FDR_comparisons": "Human FDR",
            "human_nominal_comparisons": "Human nominal",
            "mouse_nominal_units": "Mouse nominal",
            "GSE209552_acute_logFC": "Acute logFC",
            "GSE223245_severity_r_all": "Blood severity r",
        }
    )
    sns.heatmap(metrics, cmap="viridis", ax=ax_a, linewidths=0.3, cbar_kws={"label": "Observed value"})
    ax_a.set_title("Integrated dry-evidence matrix")
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    ax_a.set_xticklabels(ax_a.get_xticklabels(), rotation=35, ha="right")
    add_panel_label(ax_a, "A")

    ranked = priority.sort_values("integrated_evidence_score")
    ax_b.barh(ranked["gene_symbol"], ranked["integrated_evidence_score"], color=OKABE["purple"], edgecolor="black", lw=0.4)
    ax_b.set_xlabel("Integrated evidence score")
    ax_b.set_title("Prioritized validation genes")
    sns.despine(ax=ax_b)
    add_panel_label(ax_b, "B")

    ax_c.axis("off")
    steps = [
        ("1", "3DPI cortex/hippocampus", "qPCR/WB: SLC3A2, SLC7A11, WASF2, TLN1"),
        ("2", "Cytoskeleton and redox", "F-actin, MYH9/MYL6/FLNA, NADPH/GSH"),
        ("3", "Cell localization", "Co-stain with NeuN, GFAP, IBA1, OLIG2, CD31"),
        ("4", "Mechanistic criterion", "Same time, region and cell context move coherently"),
    ]
    for i, (num, title, detail) in enumerate(steps):
        y = 0.82 - i * 0.22
        circle = mpl.patches.Circle((0.08, y), 0.045, fc=OKABE["purple"], ec="black", lw=0.6, alpha=0.85)
        ax_c.add_patch(circle)
        ax_c.text(0.08, y, num, ha="center", va="center", fontsize=8, color="white", fontweight="bold")
        ax_c.text(0.17, y + 0.035, title, ha="left", va="center", fontsize=8.4, fontweight="bold")
        ax_c.text(0.17, y - 0.035, detail, ha="left", va="center", fontsize=7.4)
        if i < len(steps) - 1:
            ax_c.plot([0.08, 0.08], [y - 0.055, y - 0.165], color="#9CA3AF", lw=1.0)
    ax_c.set_xlim(0, 1)
    ax_c.set_ylim(0, 1)
    ax_c.set_title("Mechanistic upgrade workflow")
    add_panel_label(ax_c, "C")

    corr = gse209["marker_corr"].copy()
    corr["abs_r"] = corr["spearman_r"].abs()
    top = corr.sort_values("abs_r", ascending=False).head(8).copy()
    module_short = {
        "score_disulfidptosis_8gene": "8-gene",
        "score_transporter_2gene": "Transporter",
        "score_actin_6gene": "Actin",
    }
    cell_short = {
        "Endothelial": "Endo",
        "Neuron": "Neuron",
        "Microglia": "Micro",
        "OPC": "OPC",
        "Astrocyte": "Astro",
        "Oligodendrocyte": "Oligo",
    }
    top["label"] = top["module"].map(module_short).fillna(top["module"]) + " | " + top["celltype_proxy"].map(cell_short).fillna(top["celltype_proxy"])
    ax_d.barh(np.arange(len(top)), top["spearman_r"], color=[OKABE["red"] if v > 0 else OKABE["blue"] for v in top["spearman_r"]], edgecolor="black", lw=0.4)
    ax_d.set_yticks(np.arange(len(top)))
    ax_d.set_yticklabels(top["label"], fontsize=7.2)
    ax_d.invert_yaxis()
    ax_d.axvline(0, color="black", lw=0.8)
    ax_d.set_xlabel("Spearman r")
    ax_d.set_title("Top marker-proxy correlations")
    sns.despine(ax=ax_d)
    add_panel_label(ax_d, "D")

    savefig(fig, "Fig6_v3_integrated_priority_validation_suite_20260604")


def fmt(x: object, digits: int = 3) -> str:
    if x is None or pd.isna(x):
        return "NA"
    if isinstance(x, str):
        return x
    try:
        return f"{float(x):.{digits}g}"
    except Exception:
        return str(x)


def markdown_table(df: pd.DataFrame, cols: list[str], max_rows: int | None = None) -> str:
    work = df.copy()
    if max_rows:
        work = work.head(max_rows)
    lines = ["| " + " | ".join(cols) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
    for _, row in work[cols].iterrows():
        vals = [fmt(row[c]).replace("\n", " ") for c in cols]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def figure_block(n: int, filename: str, caption: str) -> str:
    return f'<a id="fig-{n}"></a>\n\n**Fig. {n}. {caption}**\n\n![Fig. {n}](../figures/{filename}.png)'


def build_tables_for_manuscript(evidence: pd.DataFrame, priority: pd.DataFrame, gse223: dict[str, pd.DataFrame], marker_corr: pd.DataFrame) -> dict[str, str]:
    table1_cols = ["dataset_or_source", "species", "tissue_or_model", "time_or_course", "severity_or_stage", "current_role_in_manuscript", "guardrail"]
    table2_cols = ["gene_symbol", "human_FDR_comparisons", "human_nominal_comparisons", "GSE209552_acute_logFC", "mouse_nominal_units", "GSE223245_severity_r_all", "integrated_evidence_score"]
    gene_cols = ["gene_symbol", "spearman_r_severity_all_groups", "p_value_severity_all_groups", "spearman_r_TBI_only", "p_value_TBI_only"]
    corr_cols = ["module", "celltype_proxy", "spearman_r", "p_value", "FDR"]
    corr_top = marker_corr.assign(abs_r=marker_corr["spearman_r"].abs()).sort_values("abs_r", ascending=False)
    return {
        "dataset": markdown_table(evidence, table1_cols),
        "priority": markdown_table(priority, table2_cols),
        "severity_gene": markdown_table(gse223["gene_stats"], gene_cols),
        "marker_corr": markdown_table(corr_top, corr_cols, max_rows=12),
    }


def key_numbers(human: pd.DataFrame, gse209: dict[str, pd.DataFrame], gse163: dict[str, pd.DataFrame], gse223: dict[str, pd.DataFrame], priority: pd.DataFrame) -> dict[str, str]:
    stage = human[human["comparison"].eq("GSE193407_human_prefrontal_BA9_CTE_stage_trend")]
    stage_fdr = stage[stage["significant_FDR_0_05"] == True].sort_values("FDR")
    late = human[human["comparison"].eq("GSE193407_human_prefrontal_BA9_late_CTE_stage3_4_vs_stage0")]
    late_fdr = late[late["significant_FDR_0_05"] == True].sort_values("FDR")
    acute = gse209["summary"].set_index("gene_symbol")
    module_stats = gse209["module_stats"]
    actin_mod = module_stats[(module_stats["analysis_set"] == "1_双硫死亡相关") & (module_stats["submodule"] == "双硫死亡_肌动蛋白骨架靶蛋白")]
    core_mod = module_stats[(module_stats["analysis_set"] == "1_双硫死亡相关") & (module_stats["submodule"] == "双硫死亡_核心转运与应激")]
    mouse_mod = gse163["module"]
    mouse_h = mouse_mod[(mouse_mod["time"].eq("3DPI")) & (mouse_mod["region"].eq("Hipp"))]
    sev_mod = gse223["module_stats"].sort_values("p_value_severity_all_groups").iloc[0]
    top_gene = priority.iloc[0]
    marker_top = gse209["marker_corr"].assign(abs_r=lambda x: x["spearman_r"].abs()).sort_values("abs_r", ascending=False).iloc[0]
    return {
        "stage_fdr_genes": ", ".join(stage_fdr["gene_symbol"].tolist()) or "none",
        "stage_details": "; ".join(f"{r.gene_symbol}: r={r.stage_correlation:.3f}, FDR={r.FDR:.3g}" for r in stage_fdr.itertuples()),
        "late_fdr_genes": ", ".join(late_fdr["gene_symbol"].tolist()) or "none",
        "late_details": "; ".join(f"{r.gene_symbol}: logFC={r.logFC:.3f}, FDR={r.FDR:.3g}" for r in late_fdr.itertuples()),
        "acute_actb": f"ACTB logFC={acute.loc['ACTB','GSE209552_logFC_TBI_vs_Control']:.3f}, P={acute.loc['ACTB','GSE209552_p_value_ttest']:.3g}" if "ACTB" in acute.index else "",
        "acute_wasf2": f"WASF2 logFC={acute.loc['WASF2','GSE209552_logFC_TBI_vs_Control']:.3f}, P={acute.loc['WASF2','GSE209552_p_value_ttest']:.3g}" if "WASF2" in acute.index else "",
        "actin_module": f"delta={actin_mod['delta_TBI_minus_Control'].iloc[0]:.3f}, P={actin_mod['p_value'].iloc[0]:.4f}, FDR={actin_mod['FDR'].iloc[0]:.3f}" if not actin_mod.empty else "NA",
        "core_module": f"delta={core_mod['delta_TBI_minus_Control'].iloc[0]:.3f}, P={core_mod['p_value'].iloc[0]:.4f}, FDR={core_mod['FDR'].iloc[0]:.3f}" if not core_mod.empty else "NA",
        "mouse_hipp": "; ".join(f"{r.module}/{r.treatment}: delta={r.delta_module_score_TBI_minus_NoTBI:.3f}, P={r.p_value:.3g}" for r in mouse_h.itertuples()),
        "sev_best_module": f"{sev_mod['module']}: r={sev_mod['spearman_r_severity_all_groups']:.3f}, P={sev_mod['p_value_severity_all_groups']:.3g}",
        "top_gene": f"{top_gene['gene_symbol']} (score={top_gene['integrated_evidence_score']:.1f})",
        "top_marker": f"{marker_top['module']} vs {marker_top['celltype_proxy']} marker proxy: r={marker_top['spearman_r']:.3f}, P={marker_top['p_value']:.3g}",
    }


def build_zh(vals: dict[str, str], tables: dict[str, str]) -> str:
    return f"""# TBI 后双硫死亡样转运-肌动蛋白骨架应激的时空异质性、损伤程度关联和细胞类型优先级：基于人类 TBI/CTE 与小鼠 CCI 公共转录组的整合分析

日期：{DATE}  
版本：v3 增强论文稿（中文）  
写作定位：参考 TBI 生物信息学加验证论文、急性 TBI 严重程度组学论文和双硫死亡机制论文形成的正式论文草稿；本版补充多面板结果图、GSE223245 focused 严重程度分析、GSE209552 bulk marker-proxy 细胞类型推断，以及综合优先级模型。

## 摘要

背景：创伤性脑损伤（traumatic brain injury, TBI）后的继发性损伤并非单一分子事件，而是随损伤时间、脑区、病程阶段和损伤程度改变的多层级反应。双硫死亡是一种由 SLC7A11/SLC3A2 介导的胱氨酸输入、NADPH 还原力消耗和肌动蛋白骨架二硫键压力共同驱动的新型细胞死亡形式。TBI 后是否存在双硫死亡样转运-骨架应激，以及该信号在何时、何处、何种损伤程度和哪类细胞背景中更突出，仍缺少系统整合。

方法：本研究预先固定 8 个关键基因作为主面板，其中 SLC3A2 和 SLC7A11 代表胱氨酸转运入口，WASF2、TLN1、ACTB、MYH9、MYL6 和 FLNA 代表 WRC/肌动蛋白骨架张力与结构终点。我们整合人类急性 severe TBI 脑组织 GSE209552、人类 CTE stage 数据 GSE193407、人类 CTE 外部验证 GSE319253、人类远期 TBI 多脑区 GSE104687、小鼠 CCI 时空数据 GSE163415，以及人类外周血 mild/moderate/severe TBI 数据 GSE223245。统计解释严格区分 FDR 显著与 nominal 探索性信号；GSE104687 遵守 donor/sample 独立性边界；动物数据仅作为跨物种旁证；外周血数据仅用于损伤程度相关外周线索。新增分析包括 GSE223245 原始矩阵的 8 基因和模块严重程度趋势、GSE209552 的样本级模块评分和 bulk marker-proxy 细胞类型优先级，以及跨数据集综合证据评分。

结果：人类脑组织结果显示 8 基因面板存在重复但异质的支持。GSE193407 提供最稳健的慢性病程证据，CTE stage trend 中达到 FDR 支持的基因为 {vals['stage_fdr_genes']}（{vals['stage_details']}），late CTE stage 3-4 vs stage 0 中达到 FDR 支持的基因为 {vals['late_fdr_genes']}（{vals['late_details']}）。GSE209552 急性 severe TBI 脑组织中，肌动蛋白骨架模块呈方向性增强（{vals['actin_module']}），核心转运/应激模块也呈正向但统计边界较弱（{vals['core_module']}）；单基因层面以 {vals['acute_actb']} 和 {vals['acute_wasf2']} 最具急性方向性。小鼠 GSE163415 显示 3DPI hippocampus 是最集中的动物验证窗口（{vals['mouse_hipp']}），提示早期亚急性阶段更适合验证转运-骨架模块。GSE223245 focused 分析显示外周血严重程度维度可补充 severity-aware 叙事，模块层面最强趋势为 {vals['sev_best_module']}，但不能替代脑内机制证据。GSE209552 bulk marker-proxy 分析提示 {vals['top_marker']}，这只能作为后续细胞定位实验的优先级线索，而非真正的细胞类型归属。

结论：公共转录组证据支持 TBI/CTE 后存在候选的双硫死亡样转运-肌动蛋白骨架应激。当前最稳健的主结论是：慢性 CTE stage 中 SLC7A11/SLC3A2/WASF2/TLN1 轴随病程增强；急性 severe TBI 中骨架模块已有方向性改变；小鼠 CCI 的 3DPI hippocampus 是最适合优先湿实验验证的时空窗口。机制升级需要在同一时间窗和脑区内同时证明 SLC3A2/SLC7A11 蛋白改变、F-actin/骨架终点异常、NADPH/GSH 读数改变和细胞类型共定位。

关键词：创伤性脑损伤；慢性创伤性脑病；双硫死亡；SLC7A11；SLC3A2；肌动蛋白骨架；损伤时间；损伤严重程度；公共转录组

## 引言

TBI 是导致死亡、残疾和长期神经功能障碍的重要原因。机械性原发损伤之后，脑组织会经历能量代谢危机、离子稳态紊乱、兴奋性毒性、线粒体功能障碍、氧化还原压力、血脑屏障破坏、胶质细胞激活和慢性神经炎症等继发反应。这些反应并不在所有时间点、所有脑区和所有细胞类型中均匀发生。急性数小时至数天内，能量危机和氧化还原压力往往更突出；亚急性阶段，胶质反应、细胞骨架重塑和组织修复可能逐渐参与；慢性或反复损伤后，CTE 相关病程又可呈现神经炎症、细胞类型组成改变和突触/骨架相关通路的长期偏移。因此，一个聚焦 TBI 后细胞死亡或骨架应激的研究，必须回答“何时、何处、何种损伤程度和何种细胞背景”这几个问题，而不能停留在简单的 TBI vs control 比较。

双硫死亡为这一问题提供了新的机制入口。Liu 等提出，在 SLC7A11 高表达且葡萄糖不足的背景下，细胞持续摄入胱氨酸会加重 NADPH 消耗，导致二硫键异常积累，最终诱发肌动蛋白骨架蛋白二硫键化和骨架塌陷。这一机制把氨基酸转运、还原力不足和细胞骨架机械稳定性连接起来。TBI 后脑组织恰恰存在能量危机、线粒体损伤和氧化还原失衡，因此双硫死亡样转运-骨架应激有合理的病理生理背景。不过，在疾病转录组中使用“发生双硫死亡”这一表述必须谨慎，因为 mRNA 层面的转运或骨架基因变化不能直接证明蛋白二硫键化、F-actin 塌陷或细胞死亡形式已经发生。

基于这一边界，本研究不把所有含硫代谢、氧化应激或细胞死亡基因全部混入主线，而是把问题收束到 8 个预先固定的关键基因。SLC3A2 与 SLC7A11 代表胱氨酸输入入口；WASF2 与 TLN1 代表肌动蛋白调控、黏附和张力传递；ACTB、MYH9、MYL6 与 FLNA 代表更接近双硫死亡终点的肌动蛋白骨架结构和张力节点。这样的设计牺牲了一部分广义通路覆盖，但提高了主线可检验性：如果 TBI/CTE 后确实存在双硫死亡样应激，至少应能在特定时间窗、脑区或病程阶段观察到转运入口与骨架终点的协调性改变。

本研究的目标是将已有公共数据和参考论文格式拓展成一份真正的论文式整合分析。具体问题包括：第一，8 基因面板是否在人类 TBI/CTE 脑组织中重复出现；第二，该信号是否与急性 severe TBI、慢性 CTE stage 或外周血严重程度分层相关；第三，小鼠 CCI 中是否存在明确的时间和脑区窗口；第四，当前 bulk 数据能否为神经元、星形胶质细胞、小胶质细胞、少突胶质细胞和内皮细胞的后续验证提供优先级线索。整篇论文的所有结论均保持在“候选机制”和“探索性支持”层面，避免将 nominal 候选、外周血结果或动物旁证直接写成因果或诊断标志物。

## 材料与方法

### 研究设计与参考格式

本研究的文章结构参考 TBI 生物信息学加验证型论文：先说明数据筛选和候选面板定义，再报告差异表达、模块评分、疾病程度关联、外部验证、动物旁证和后续实验设计。Zhao 等关于 TBI 中 Nrf2/pyroptosis 的文章用于参考公共数据挖掘、候选分子排序和临床严重程度关联的写法；Thomas 等急性 TBI 代谢组论文用于参考损伤严重程度、采样窗口和结果解释边界的组织方式；Liu 等双硫死亡机制论文用于定义 SLC7A11/SLC3A2、NADPH 和肌动蛋白骨架逻辑；小鼠 TBI/CCI 文献和本课题既往工作用于指导 3DPI、7D/29DPI、皮层/海马等验证窗口。

### 数据集与证据分层

本研究纳入六个公共转录组证据层。GSE209552 为人类急性 severe TBI 脑组织，损伤后时间约 4 h 至 8 d，代表急性重型损伤方向性证据。GSE193407 为人类 BA9 CTE stage 0-4 数据，适合评估慢性病程梯度。GSE319253 为 superior frontal cortex CTE vs control 数据，用于慢性外部验证。GSE104687 为远期 TBI 多脑区数据，因同一 donor 多脑区样本存在非独立性，区域结果仅作为空间探索层。GSE163415 为小鼠 CCI 数据，覆盖 3DPI/29DPI、hippocampus/thalamus/hypothalamus 和 vehicle/drug 分层，用于跨物种时空旁证。GSE223245 为人类 whole blood/PBMC mild/moderate/severe TBI 微阵列数据，只用于外周 severity-aware 线索。

{tables['dataset']}

### 候选基因与模块定义

候选面板在分析前固定为 SLC3A2、SLC7A11、WASF2、TLN1、ACTB、MYH9、MYL6 和 FLNA。模块层面定义三个读数：8 基因总模块、SLC3A2/SLC7A11 转运入口模块，以及 WASF2/TLN1/ACTB/MYH9/MYL6/FLNA 肌动蛋白骨架模块。模块评分采用每个基因在样本内表达矩阵中的 z-score 后求均值。对于 GSE209552 和 GSE223245 这种样本量较小的数据，模块统计解释优先报告效应方向、P 值和 FDR，不将 nominal 结果作为正式阳性结论。

### GSE223245 focused 严重程度再分析

本版新增 GSE223245 原始 series matrix 与 family SOFT 平台注释下载，按平台注释将探针折叠到基因符号；同一基因多探针时保留平均表达最高的探针。样本按 title 和 characteristics 分为 Control、Mild、Moderate 和 Severe，严重程度记为 0、1、2、3。对 8 个基因分别计算所有组（0-3）和 TBI-only（1-3）的 Spearman 严重程度相关，并对 Mild/Moderate/Severe vs Control 进行 Welch t 检验。模块评分同样计算所有组和 TBI-only 的 Spearman 趋势。由于每组 n=4 且组织来源为外周血，本分析只用于“损伤程度相关外周线索”，不用于脑内机制推断。

### GSE209552 bulk marker-proxy 细胞类型推断

由于当前本地可用的 GSE209552 基因矩阵为 4 个 TBI 与 3 个 control 的 bulk-like 脑组织计数矩阵，而非可直接细胞注释的 snRNA-seq 对象，本研究新增 bulk marker-proxy 分析作为细胞类型优先级线索。我们使用 RBFOX3/SNAP25/SYT1/MAP2、GFAP/AQP4/ALDH1L1/SLC1A3、P2RY12/CX3CR1/AIF1/CSF1R、MBP/MOG/PLP1/MAG、PDGFRA/CSPG4/VCAN/SOX10 和 PECAM1/VWF/CLDN5/FLT1 分别代表神经元、星形胶质、小胶质、少突胶质、OPC 和内皮 proxy 分数，并计算它们与 8 基因、转运入口和骨架模块的 Spearman 相关。该分析不能作为细胞类型归属，只能用于决定后续 IF/IHC 共定位优先级。

### 可视化与论文输出

参考正式论文套图习惯，本版将结果组织为 6 张多面板图。每张图包含 A-D 或 A-C 子图，分别承担研究设计、人脑结果、急性 severe TBI、动物 CCI、外周 severity 和综合优先级。所有图同时保存 PNG 和 PDF；中文和英文报告均以 Markdown 为中间格式，并导出 Word。正文中所有图引用均使用 [Fig. 1](#fig-1) 至 [Fig. 6](#fig-6) 的内部跳转形式。

## 结果

### 多证据层设计将 TBI 后双硫死亡问题聚焦为可检验的转运-骨架轴

[Fig. 1](#fig-1) 总结了本研究的证据分层和解释边界。与泛化的细胞死亡通路筛选不同，本研究在写作前固定 8 个关键基因，并将证据按急性 severe TBI、慢性 CTE stage、远期脑区探索、小鼠 CCI 时空旁证和外周血严重程度分层。这样的设计有两个优点。第一，它避免了在大量候选通路中事后挑选阳性基因；第二，它迫使每个结果都回答特定问题：急性重型损伤中是否已出现模块方向性改变，慢性病程中是否出现随 stage 增强的转运/骨架轴，动物 CCI 中哪个脑区和时间点最适合验证，以及外周血严重程度能否为临床分层提供补充线索。

{figure_block(1, "Fig1_v3_study_design_evidence_layers_20260604", "Study design, evidence layers, pre-fixed 8-gene panel and interpretation guardrails.")}

### 人类脑组织支持 SLC7A11/SLC3A2/WASF2/TLN1 在慢性 CTE 病程中增强

人类脑组织分析首先回答该面板是否在 TBI/CTE 背景中重复出现。[Fig. 2](#fig-2)A 显示，不同数据集和比较的效应并不均一。GSE104687 远期 TBI 多脑区存在若干 nominal 方向性变化，但 donor-level 主分析没有 FDR DEG，因此只能作为空间探索。相比之下，GSE193407 CTE stage 结果是当前最强的人脑证据：stage trend 中 FDR 支持基因为 {vals['stage_fdr_genes']}，具体为 {vals['stage_details']}；late CTE stage 3-4 vs stage 0 中 FDR 支持基因为 {vals['late_fdr_genes']}，具体为 {vals['late_details']}。[Fig. 2](#fig-2)B 将这一病程梯度单独呈现，可以看到转运入口和 WRC/张力调控节点更稳定，而 ACTB、MYH9、MYL6 和 FLNA 在 mRNA 层面未形成同等强度的 FDR 信号。[Fig. 2](#fig-2)C 进一步显示，跨人脑比较的复现度并非均匀分布，SLC7A11、SLC3A2、WASF2 和 TLN1 更适合作为主读数，而骨架终点更适合与蛋白和形态学实验结合解释。

这一结果的生物学含义是，TBI/CTE 相关慢性病程中更可能首先表现为“胱氨酸转运入口 + actin remodeling/张力调控节点”的转录增强，而不是所有骨架结构基因同步上调。该模式与双硫死亡机制相吻合，但仍不能证明双硫死亡已经发生。真正的机制证明需要检测 SLC3A2/SLC7A11 蛋白、NADPH/GSH 状态、F-actin 形态和骨架蛋白二硫键化。

{figure_block(2, "Fig2_v3_human_brain_multidataset_8gene_20260604", "Human brain multidataset evidence for the 8-gene transporter-actin panel. Filled dots in panel A denote FDR support and open dots denote nominal support.")}

### 急性 severe TBI 中骨架模块已有方向性变化，且可用 marker proxy 提出细胞定位优先级

GSE209552 代表急性 severe TBI 人脑证据。模块层面，肌动蛋白骨架模块在 TBI 相对 control 中升高，统计结果为 {vals['actin_module']}；核心转运/应激模块也呈正向趋势，统计结果为 {vals['core_module']}。[Fig. 3](#fig-3)A 展示样本级模块评分，提示急性重型损伤中骨架终点模块比转运入口模块更明显。单基因层面，[Fig. 3](#fig-3)B 显示 ACTB 和 WASF2 在该数据中方向性较强，分别为 {vals['acute_actb']} 和 {vals['acute_wasf2']}。由于样本量只有 TBI n=4、control n=3，这些结果不能写成急性 TBI 的正式分子标志物，但适合作为“急性重型损伤窗口已出现转运-骨架方向性改变”的支持。

本版还新增 bulk marker-proxy 分析，用来弥补当前细胞类型归属不足。[Fig. 3](#fig-3)D 显示 8 基因模块、转运入口模块和骨架模块与神经元、星形胶质、小胶质、少突胶质、OPC 和内皮 marker proxy 的相关矩阵。其中最强的相关线索为 {vals['top_marker']}。这一结果不能替代 snRNA-seq 细胞注释，也不能证明某一细胞类型发生双硫死亡；它的价值在于为后续共定位实验排序。例如，如果转运模块与内皮或胶质 marker proxy 更相关，后续 SLC3A2/SLC7A11 与 CD31、GFAP、IBA1 或 OLIG2 的共定位就应优先于单纯 NeuN 定位。

{figure_block(3, "Fig3_v3_GSE209552_acute_severe_and_marker_proxy_20260604", "Acute severe TBI evidence in GSE209552, including module scores, 8-gene logFC, enrichment curve and bulk marker-proxy cell-type prioritization.")}

### 小鼠 CCI 显示 3DPI hippocampus 是最集中的转运-骨架验证窗口

动物数据用于回答时间和脑区问题。[Fig. 4](#fig-4)A 显示 GSE163415 中不同 time-region-treatment 单元的模块评分变化。3DPI hippocampus 是最集中的动物窗口，多个模块呈 TBI-NoTBI 正向差值；其中 {vals['mouse_hipp']}。[Fig. 4](#fig-4)B 的 8 基因 logFC 热图进一步显示，早期 hippocampus 中转运入口和骨架终点均更容易出现方向性变化。[Fig. 4](#fig-4)C 的 pathway overlap 也支持 3DPI hippocampus/all treatments 中 actin cytoskeleton and tension、disulfidptosis core/PDF genes 和 amino-acid transport/cystine axis 的共同命中。

29DPI 结果与 3DPI 不完全相同。后期广泛骨架模块信号减弱，但 amino-acid transport/cystine axis 仍有残留支持。这提示 TBI 后双硫死亡样应激可能具有阶段性：早期亚急性阶段更偏“转运入口 + 骨架张力”共同增强，后期可能保留部分转运/代谢轴变化。基于这一结果，湿实验第一轮不宜平均铺开过多时间点，而应优先选择 3DPI cortex/hippocampus，并将 7D 或 29DPI 作为持续性观察窗口。

{figure_block(4, "Fig4_v3_GSE163415_mouse_CCI_spatiotemporal_suite_20260604", "Mouse CCI spatiotemporal evidence in GSE163415, including module-score shifts, 8-gene logFC, pathway overlap and validation priority.")}

### GSE223245 外周血 focused 再分析补充损伤严重程度维度，但不代表脑内机制

为补充损伤程度，本版重新下载并分析 GSE223245 原始矩阵。[Fig. 5](#fig-5)A 的 8 基因 PCA 显示 Control、Mild、Moderate 和 Severe 之间存在一定分布差异，但样本数只有每组 n=4，且来源为外周血/PBMC。[Fig. 5](#fig-5)B-D 展示模块和单基因的严重程度相关。模块层面最强趋势为 {vals['sev_best_module']}。单基因层面的详细 Spearman 结果列于 Table 3。

GSE223245 的意义不是证明脑组织内发生双硫死亡，而是提醒论文设计需要纳入“severity-aware”视角。急性 severe TBI 脑组织、CTE stage 和外周血 mild/moderate/severe 不是同一个尺度，不能合并为统一 severity score。更专业的写法是把它们放在互补证据层：GSE209552 代表急性重型脑组织窗口，GSE193407 代表慢性病程梯度，GSE223245 代表外周血损伤程度分层线索。

{figure_block(5, "Fig5_v3_GSE223245_peripheral_severity_suite_20260604", "Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data.")}

Table 3. GSE223245 8 基因严重程度趋势结果。

{tables['severity_gene']}

### 综合优先级显示 SLC3A2/SLC7A11/WASF2/TLN1 是主读数，ACTB/MYH9/MYL6/FLNA 更适合作为蛋白和形态学终点

最后，我们将人脑 FDR/nominal 支持、GSE209552 急性方向性、小鼠 CCI 支持和 GSE223245 外周严重程度相关整合成 8 基因优先级矩阵。[Fig. 6](#fig-6)A-B 显示当前综合优先基因为 {vals['top_gene']}，但这并不意味着其可作为诊断标志物，而是表示它最适合进入下一轮实验验证。整体上，SLC7A11、SLC3A2、WASF2 和 TLN1 更适合作为主表达读数；ACTB、MYH9、MYL6 和 FLNA 更接近骨架终点，推荐结合 WB、免疫荧光、非还原胶、F-actin 染色和形态学分析。

[Fig. 6](#fig-6)C 给出机制升级路径：第一步在 3DPI cortex/hippocampus 做 qPCR/WB 验证 SLC3A2/SLC7A11/WASF2/TLN1；第二步检测 ACTB/MYH9/MYL6/FLNA 和 phalloidin F-actin 形态；第三步同步检测 NADPH/NADP+ 和 GSH/GSSG；第四步做 SLC3A2/SLC7A11 与 NeuN、GFAP、IBA1、OLIG2、CD31 的共定位。只有当这些读数在同一时间窗、同一脑区和相近细胞背景中同向改变时，才能把“转录组候选”推进为“机制支持”。

{figure_block(6, "Fig6_v3_integrated_priority_validation_suite_20260604", "Integrated 8-gene priority matrix, validation ranking, mechanistic upgrade workflow and marker-proxy correlation summary.")}

Table 2. 8 基因综合证据优先级。

{tables['priority']}

Table 4. GSE209552 bulk marker-proxy 相关性最高的细胞类型优先级线索。

{tables['marker_corr']}

## 讨论

本研究将 TBI 后双硫死亡问题从广义含硫代谢和细胞死亡框架收束为一个更可检验的转运-肌动蛋白骨架主线。相比只报告差异基因列表，本版分析的改进在于同时引入了时间、脑区、病程阶段、损伤程度和细胞类型优先级。结果并不支持“所有双硫死亡相关基因在 TBI 后同步上调”这种简单叙事，而更支持一个分层模型：慢性 CTE stage 中，SLC7A11/SLC3A2/WASF2/TLN1 轴较稳定增强；急性 severe TBI 中，骨架模块已出现方向性改变；小鼠 CCI 中，3DPI hippocampus 是最集中的验证窗口；外周血严重程度数据提供补充线索，但不能代表脑内机制。

这一模型与双硫死亡的原始机制存在合理衔接。SLC7A11/SLC3A2 负责胱氨酸输入，可能在能量危机和 NADPH 不足背景下放大二硫键压力；WASF2 和 TLN1 位于 actin remodeling、黏附和张力传递层面，可能反映骨架脆弱性或修复性重塑；ACTB、MYH9、MYL6 和 FLNA 是更接近终点的结构读数，但 mRNA 变化未必等同于蛋白二硫键化或骨架塌陷。因此，本研究在术语上使用“双硫死亡样转运-骨架应激”，而不是直接宣布“双硫死亡已经发生”。

损伤程度的处理需要特别谨慎。GSE209552 的 severe TBI、GSE193407 的 CTE stage 和 GSE223245 的 mild/moderate/severe 外周血分层代表不同维度：急性临床严重程度、慢性神经病理病程和外周分子分层。把三者合并成单一 severity score 会制造统计幻觉。更可靠的写法是把 severe TBI 作为急性重型脑组织窗口，把 CTE stage 作为慢性病程梯度，把 GSE223245 作为外周严重程度补充。这样的证据分层与 Thomas 等急性 TBI 严重程度组学论文的叙事方式更接近，也符合本课题 AGENTS 规范。

细胞类型问题是当前最重要的证据缺口。本研究新增的 bulk marker-proxy 分析可以提出候选方向，但不能替代单细胞或空间定位。TBI 后 bulk 表达变化可能来自细胞组成改变，也可能来自同一细胞类型内的状态改变。尤其是 SLC3A2/SLC7A11 在内皮、胶质和神经元中都可能有不同功能，ACTB/MYH9/MYL6/FLNA 又是广泛表达的骨架基因，因此单靠 bulk 相关性无法判断敏感细胞类型。正式机制论文需要提取 GSE209552 或其他 snRNA-seq 数据中带细胞注释的表达矩阵，并在动物组织中进行 NeuN、GFAP、IBA1、OLIG2 和 CD31 共定位。

本研究还存在若干局限。第一，各公共数据集平台、组织来源、疾病定义和协变量不同，不能直接合并成统一效应量。第二，GSE104687 的多脑区样本存在 donor 内非独立性，区域 nominal 信号只能探索性使用。第三，GSE209552 样本量小且代表 severe TBI 外科切除脑组织，不能外推至所有 TBI。第四，GSE223245 为外周血/PBMC，不应被写成脑组织机制证据。第五，转录组不能替代蛋白、还原力、F-actin 形态和细胞死亡验证。第六，本版综合评分用于实验排序，不是统计学模型或临床预测模型。

## 结论

本研究基于人类 TBI/CTE、小鼠 CCI 和外周血严重程度公共数据，形成了一个更完整的 TBI 后双硫死亡样转运-肌动蛋白骨架应激论文框架。当前证据支持：慢性 CTE stage 中 SLC7A11/SLC3A2/WASF2/TLN1 轴较稳健；急性 severe TBI 中骨架模块已有方向性增强；小鼠 CCI 的 3DPI hippocampus 是优先验证窗口；外周血 severe-aware 分析可作为补充但不能替代脑组织证据。后续实验应优先围绕 3DPI cortex/hippocampus 的 SLC3A2/SLC7A11、WASF2/TLN1、ACTB/MYH9/MYL6/FLNA、F-actin、NADPH/GSH 和细胞类型共定位展开。

## 图注

**Fig. 1.** Study design, evidence layers, pre-fixed 8-gene panel and interpretation guardrails.  
**Fig. 2.** Human brain multidataset evidence for the 8-gene transporter-actin panel.  
**Fig. 3.** Acute severe TBI evidence in GSE209552, including module scores and marker-proxy prioritization.  
**Fig. 4.** Mouse CCI spatiotemporal evidence in GSE163415.  
**Fig. 5.** Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data.  
**Fig. 6.** Integrated 8-gene priority matrix and validation workflow.

## 数据可用性

本研究使用 GEO 公共数据 GSE104687、GSE209552、GSE193407、GSE319253、GSE163415 和 GSE223245。所有重分析表格、PNG/PDF 图件、中英文 Markdown 与 Word 论文稿保存在 `Phase3_深化优化与最终报告_20260506-0513/11_双硫死亡聚焦论文设计_20260604/`。

## 伦理声明

本阶段仅使用公开去标识化数据和既有文献资料，不涉及新增人体样本。后续动物实验需在开展前获得所在单位动物伦理委员会批准。

## 作者贡献

待根据实际团队分工补充。

## 资金

待补充。

## 利益冲突

作者声明待补充。

## 参考文献

1. Liu X, Nie L, Zhang Y, et al. Actin cytoskeleton vulnerability to disulfide stress mediates disulfidptosis. Nature Cell Biology. 2023;25:404-414. https://doi.org/10.1038/s41556-023-01091-2  
2. Machesky LM. Deadly actin collapse by disulfidptosis. Nature Cell Biology. 2023.  
3. Zhao G, Zhao J, Lang J, Sun G. Nrf2 functions as a pyroptosis-related mediator in traumatic brain injury and is correlated with cytokines and disease severity: a bioinformatics analysis and retrospective clinical study. Frontiers in Neurology. 2024;15:1341342. https://doi.org/10.3389/fneur.2024.1341342  
4. Thomas I, Dickens AM, Posti JP, et al. Serum metabolome associated with severity of acute traumatic brain injury. Nature Communications. 2022;13:2545. https://doi.org/10.1038/s41467-022-30227-5  
5. Zhang M, Shan H, Chang P, et al. Hydrogen Sulfide Offers Neuroprotection on Traumatic Brain Injury in Parallel with Reduced Apoptosis and Autophagy in Mice. PLoS ONE. 2014;9:e87241. https://doi.org/10.1371/journal.pone.0087241  
6. Labadorf A, Agus F, Aytan N, et al. Inflammation and neuronal gene expression changes differ in early versus late chronic traumatic encephalopathy brain. BMC Medical Genomics. 2023;16:49. https://doi.org/10.1186/s12920-023-01471-5  
7. Garza R, Sharma Y, Atacho DAM, et al. Single-cell transcriptomics of human traumatic brain injury reveals activation of endogenous retroviruses in oligodendroglia. Cell Reports. 2023;42:113395. https://doi.org/10.1016/j.celrep.2023.113395  
8. Attilio PJ, Snapper DM, Rusnak M, et al. Transcriptomic Analysis of Mouse Brain After Traumatic Brain Injury Reveals That the Angiotensin Receptor Blocker Candesartan Acts Through Novel Pathways. Frontiers in Neuroscience. 2021;15:636259. https://doi.org/10.3389/fnins.2021.636259
"""


def build_en(vals: dict[str, str], tables: dict[str, str]) -> str:
    return f"""# Spatiotemporal, Severity-Related and Cell-Type-Prioritized Disulfidptosis-Like Transporter-Actin Cytoskeletal Stress After Traumatic Brain Injury: An Integrative Transcriptomic Analysis of Human TBI/CTE and Mouse CCI Datasets

Date: {DATE}  
Version: v3 enhanced manuscript draft (English)  
Positioning: a full manuscript-style draft modeled on TBI bioinformatics-validation papers, acute TBI severity omics papers and the original disulfidptosis mechanism literature.

## Abstract

Background: Secondary injury after traumatic brain injury (TBI) varies by time after injury, brain region, disease course, injury severity and cellular context. Disulfidptosis is a recently described cell-death mechanism driven by SLC7A11/SLC3A2-dependent cystine uptake, NADPH depletion and disulfide stress on the actin cytoskeleton. Whether TBI induces a disulfidptosis-like transporter-actin cytoskeletal stress program, and when, where and in which injury contexts this program is most evident, remains insufficiently defined.

Methods: We pre-specified an 8-gene panel. SLC3A2 and SLC7A11 represented the cystine-import entry point, whereas WASF2, TLN1, ACTB, MYH9, MYL6 and FLNA represented actin remodeling, adhesion/tension and cytoskeletal endpoint biology. We integrated human acute severe TBI brain tissue (GSE209552), human CTE stage data (GSE193407), an external CTE dataset (GSE319253), remote multi-region human TBI data (GSE104687), mouse controlled cortical impact (CCI) time-region data (GSE163415), and human peripheral blood mild/moderate/severe TBI data (GSE223245). All analyses distinguished FDR-supported findings from nominal exploratory signals, respected donor/sample independence, treated animal data as supportive evidence only, and used peripheral blood data only as a severity-aware external context. New dry-lab analyses included focused GSE223245 severity trends, GSE209552 sample-level module scores and bulk marker-proxy cell-type prioritization, and an integrated 8-gene evidence score.

Results: Human brain datasets showed recurrent but heterogeneous support for the 8-gene panel. GSE193407 provided the strongest chronic disease-course evidence: FDR-supported stage-trend genes were {vals['stage_fdr_genes']} ({vals['stage_details']}), and FDR-supported late CTE stage 3-4 versus stage 0 genes were {vals['late_fdr_genes']} ({vals['late_details']}). In acute severe TBI brain tissue, the actin cytoskeletal module increased directionally ({vals['actin_module']}), while the core transporter/stress module was positive but statistically weaker ({vals['core_module']}). Mouse GSE163415 identified 3DPI hippocampus as the most concentrated validation window ({vals['mouse_hipp']}). GSE223245 added a peripheral severity-aware layer, with the strongest module trend being {vals['sev_best_module']}, but this result does not constitute brain-mechanism evidence. GSE209552 bulk marker-proxy analysis highlighted {vals['top_marker']}, which should be interpreted only as a prioritization clue for cell-type localization.

Conclusions: Public transcriptomic evidence supports a candidate disulfidptosis-like transporter-actin cytoskeletal stress program after TBI/CTE. The most robust current conclusion is that the SLC7A11/SLC3A2/WASF2/TLN1 axis strengthens along chronic CTE stage, acute severe TBI already shows directional actin-module change, and mouse CCI 3DPI hippocampus is the most rational first validation window. Mechanistic upgrading requires protein, F-actin, redox and cell-type co-localization evidence in the same time window and brain region.

Keywords: traumatic brain injury; chronic traumatic encephalopathy; disulfidptosis; SLC7A11; SLC3A2; actin cytoskeleton; injury time; injury severity; public transcriptomics

## Introduction

TBI is a major cause of death, disability and long-term neurological dysfunction. After the primary mechanical insult, brain tissue undergoes energy failure, ionic imbalance, excitotoxicity, mitochondrial injury, redox stress, blood-brain barrier disruption, glial activation and chronic neuroinflammation. These responses are not uniform across time, brain region or cell type. Early hours to days are dominated by metabolic crisis and redox stress, whereas subacute and chronic phases may involve cytoskeletal remodeling, inflammatory persistence, cell-composition shifts and neurodegenerative cascades. A focused mechanism study therefore needs to address not only whether a pathway is perturbed after TBI, but also when, where, at what injury severity and in what cellular context the perturbation is most plausible.

Disulfidptosis provides a mechanistic entry point connecting amino-acid transport, redox failure and cytoskeletal fragility. In the original model, high SLC7A11 expression under glucose limitation promotes cystine uptake, increases NADPH consumption, causes abnormal disulfide accumulation and ultimately collapses the actin cytoskeleton. TBI brain tissue is characterized by energy crisis and oxidative-reductive imbalance, making a disulfidptosis-like process biologically plausible. However, disease transcriptomics cannot by itself prove disulfidptosis. Increased transporter mRNA, altered cytoskeletal transcripts or broad oxidative stress signals do not demonstrate cytoskeletal protein disulfidation, F-actin collapse or a specific cell-death modality.

For this reason, the present study narrows the question to a pre-fixed transporter-actin panel. SLC3A2 and SLC7A11 represent the cystine-import entry point; WASF2 and TLN1 represent actin remodeling, adhesion and tension transfer; ACTB, MYH9, MYL6 and FLNA represent cytoskeletal structural and contractile endpoints. This focused design makes the hypothesis testable: if TBI/CTE induces a disulfidptosis-like state, then specific time windows, regions or disease stages should show coordinated evidence across transporter entry and actin endpoint genes.

## Materials and Methods

### Study Design and Template Sources

The manuscript structure follows a TBI bioinformatics-validation format: dataset selection, candidate panel definition, differential expression, module scoring, disease-course or severity association, external validation, animal supportive evidence and experimental prioritization. The writing style was modeled on a TBI Nrf2/pyroptosis bioinformatics-validation paper, an acute TBI severity metabolomics paper, the original disulfidptosis mechanism paper, and mouse TBI/CCI validation literature.

### Datasets and Evidence Tiers

Six public transcriptomic evidence tiers were used. GSE209552 represented human acute severe TBI brain tissue sampled approximately 4 h to 8 d after injury. GSE193407 represented human BA9 CTE stage 0-4 and was used for chronic disease-course analysis. GSE319253 represented external superior frontal cortex CTE versus control validation. GSE104687 represented remote multi-region human TBI and was treated as a regional exploratory layer because multiple regions from the same donor are not independent observations. GSE163415 represented mouse CCI across 3DPI/29DPI, hippocampus/thalamus/hypothalamus and treatment strata. GSE223245 represented peripheral blood/PBMC mild/moderate/severe TBI and was used only as a peripheral severity-aware context.

{tables['dataset']}

### Candidate Genes, Modules and Statistics

The primary panel was fixed before analysis as SLC3A2, SLC7A11, WASF2, TLN1, ACTB, MYH9, MYL6 and FLNA. We calculated an 8-gene module, a 2-gene transporter module and a 6-gene actin module from z-scored expression values. Focused GSE223245 analysis used Spearman correlation across severity scores 0-3 and across TBI-only scores 1-3, plus Welch tests for each TBI severity group versus control. GSE209552 bulk marker-proxy analysis used canonical markers for neurons, astrocytes, microglia, oligodendrocytes, OPCs and endothelial cells to prioritize, but not prove, cell-type involvement.

## Results

### A Multi-Tier Design Focused the Question Into a Testable Transporter-Actin Axis

[Fig. 1](#fig-1) summarizes the evidence design and interpretation guardrails. The key choice was to keep the 8-gene panel fixed and then ask distinct questions in distinct evidence tiers: acute severe human brain, chronic CTE stage, remote regional TBI, mouse CCI time-region validation, peripheral severity context and experimental upgrade planning. This structure reduces post hoc gene selection and prevents peripheral or animal data from being overinterpreted as human brain mechanism.

{figure_block(1, "Fig1_v3_study_design_evidence_layers_20260604", "Study design, evidence layers, pre-fixed 8-gene panel and interpretation guardrails.")}

### Human Brain Data Supported SLC7A11/SLC3A2/WASF2/TLN1 Along Chronic CTE Course

[Fig. 2](#fig-2) shows that the 8-gene panel was recurrent but heterogeneous across human brain comparisons. GSE104687 provided only exploratory regional nominal clues, whereas GSE193407 provided the strongest FDR-supported disease-course evidence. Stage-trend FDR-supported genes were {vals['stage_fdr_genes']} ({vals['stage_details']}), and late CTE stage 3-4 versus stage 0 FDR-supported genes were {vals['late_fdr_genes']} ({vals['late_details']}). This pattern suggests a chronic transporter/WRC-actin tension axis rather than synchronous up-regulation of every cytoskeletal endpoint transcript.

{figure_block(2, "Fig2_v3_human_brain_multidataset_8gene_20260604", "Human brain multidataset evidence for the 8-gene transporter-actin panel.")}

### Acute Severe TBI Showed Directional Actin-Module Change and Marker-Proxy Prioritization

In GSE209552, [Fig. 3](#fig-3) shows that the actin cytoskeletal module increased in acute severe TBI ({vals['actin_module']}), whereas the core transporter/stress module was positive but less statistically secure ({vals['core_module']}). ACTB and WASF2 showed notable acute directions ({vals['acute_actb']}; {vals['acute_wasf2']}). Bulk marker-proxy analysis identified {vals['top_marker']} as the strongest prioritization clue, but this analysis cannot replace annotated snRNA-seq or tissue co-localization.

{figure_block(3, "Fig3_v3_GSE209552_acute_severe_and_marker_proxy_20260604", "Acute severe TBI evidence in GSE209552, including module scores, 8-gene logFC, enrichment curve and marker-proxy prioritization.")}

### Mouse CCI Identified 3DPI Hippocampus as the Most Concentrated Validation Window

GSE163415 provided the clearest time-region evidence. As shown in [Fig. 4](#fig-4), the strongest signals concentrated in 3DPI hippocampus, where {vals['mouse_hipp']}. Pathway overlap supported actin cytoskeleton and tension, disulfidptosis core/PDF genes and amino-acid transport/cystine-axis terms. Later 29DPI data showed weaker broad cytoskeletal effects but retained some transport/metabolic-axis support, suggesting a stage-dependent response.

{figure_block(4, "Fig4_v3_GSE163415_mouse_CCI_spatiotemporal_suite_20260604", "Mouse CCI spatiotemporal evidence in GSE163415.")}

### Peripheral GSE223245 Added a Severity-Aware Context Without Substituting for Brain Evidence

Focused reanalysis of GSE223245 whole blood/PBMC data added an injury-severity layer. [Fig. 5](#fig-5) shows the 8-gene PCA, group-level module scores and severity trend results. The strongest module trend was {vals['sev_best_module']}. These findings are useful for severity-aware framing but cannot be interpreted as brain-cell disulfidptosis evidence because the tissue is peripheral blood and each group contains only four samples.

{figure_block(5, "Fig5_v3_GSE223245_peripheral_severity_suite_20260604", "Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data.")}

Table 3. GSE223245 8-gene severity trend results.

{tables['severity_gene']}

### Integrated Prioritization Favored Transporter/WRC-Tension Readouts for Expression and Cytoskeletal Endpoints for Protein/Morphology

[Fig. 6](#fig-6) shows that the integrated evidence matrix ranked {vals['top_gene']} as the current top validation candidate. Across the panel, SLC7A11, SLC3A2, WASF2 and TLN1 are best suited as expression-level primary readouts, while ACTB, MYH9, MYL6 and FLNA should be interpreted together with protein abundance, non-reducing assays, F-actin morphology and co-localization.

{figure_block(6, "Fig6_v3_integrated_priority_validation_suite_20260604", "Integrated 8-gene priority matrix, validation ranking, mechanistic upgrade workflow and marker-proxy correlation summary.")}

Table 2. Integrated 8-gene evidence priority.

{tables['priority']}

Table 4. Top GSE209552 bulk marker-proxy correlation clues.

{tables['marker_corr']}

## Discussion

This study reframes post-TBI disulfidptosis as a focused, testable transporter-actin cytoskeletal stress hypothesis. The data do not support a simple claim that all disulfidptosis-related genes are uniformly up-regulated after TBI. Instead, they support a layered model: chronic CTE stage strengthens the SLC7A11/SLC3A2/WASF2/TLN1 axis; acute severe TBI shows directional actin-module change; mouse CCI localizes the most concentrated validation window to 3DPI hippocampus; and peripheral severity data provide only a complementary blood-based context.

The findings are compatible with the original disulfidptosis mechanism but remain short of proving it. Transcriptomic evidence cannot demonstrate protein disulfidation, F-actin collapse or a specific cell-death modality. The most accurate term is therefore “disulfidptosis-like transporter-actin cytoskeletal stress.” Mechanistic upgrading should require concordant transporter protein changes, cytoskeletal endpoint changes, F-actin abnormalities, NADPH/GSH alterations and cell-type co-localization in the same time window and region.

Several limitations require emphasis. Public datasets differ in platform, species, tissue, injury definition and covariate availability. GSE104687 regional samples are not independent donor-level observations. GSE209552 is small and represents severe surgically sampled brain tissue. GSE223245 is peripheral blood and cannot substitute for brain tissue. The marker-proxy analysis is only a prioritization layer, not cell-type attribution. The integrated evidence score is intended for experimental prioritization, not statistical inference or clinical prediction.

## Conclusion

Public transcriptomic data support a candidate disulfidptosis-like transporter-actin cytoskeletal stress program after TBI/CTE. The strongest dry-lab conclusion is a chronic CTE-stage association of SLC7A11/SLC3A2/WASF2/TLN1, directional acute severe TBI actin-module change, and a mouse CCI 3DPI hippocampal validation window. The next decisive step is tissue-level validation of SLC3A2/SLC7A11, WASF2/TLN1, ACTB/MYH9/MYL6/FLNA, F-actin, NADPH/GSH and cell-type co-localization.

## Figure Legends

**Fig. 1.** Study design, evidence layers, pre-fixed 8-gene panel and interpretation guardrails.  
**Fig. 2.** Human brain multidataset evidence for the 8-gene transporter-actin panel.  
**Fig. 3.** Acute severe TBI evidence in GSE209552.  
**Fig. 4.** Mouse CCI spatiotemporal evidence in GSE163415.  
**Fig. 5.** Peripheral severity-focused analysis of GSE223245.  
**Fig. 6.** Integrated 8-gene priority matrix and validation workflow.

## Data Availability

All public datasets were obtained from GEO: GSE104687, GSE209552, GSE193407, GSE319253, GSE163415 and GSE223245. Reanalysis tables, PNG/PDF figures, and Chinese/English Markdown and Word manuscripts are stored in `Phase3_深化优化与最终报告_20260506-0513/11_双硫死亡聚焦论文设计_20260604/`.

## Ethics Statement

This stage used only public de-identified datasets and literature materials. Any subsequent animal experiment requires institutional animal ethics approval before initiation.

## Author Contributions

To be completed according to the actual team roles.

## Funding

To be completed.

## Conflict of Interest

To be completed.

## References

1. Liu X, Nie L, Zhang Y, et al. Actin cytoskeleton vulnerability to disulfide stress mediates disulfidptosis. Nature Cell Biology. 2023;25:404-414. https://doi.org/10.1038/s41556-023-01091-2  
2. Machesky LM. Deadly actin collapse by disulfidptosis. Nature Cell Biology. 2023.  
3. Zhao G, Zhao J, Lang J, Sun G. Nrf2 functions as a pyroptosis-related mediator in traumatic brain injury and is correlated with cytokines and disease severity. Frontiers in Neurology. 2024;15:1341342. https://doi.org/10.3389/fneur.2024.1341342  
4. Thomas I, Dickens AM, Posti JP, et al. Serum metabolome associated with severity of acute traumatic brain injury. Nature Communications. 2022;13:2545. https://doi.org/10.1038/s41467-022-30227-5  
5. Zhang M, Shan H, Chang P, et al. Hydrogen Sulfide Offers Neuroprotection on Traumatic Brain Injury in Parallel with Reduced Apoptosis and Autophagy in Mice. PLoS ONE. 2014;9:e87241. https://doi.org/10.1371/journal.pone.0087241  
6. Labadorf A, Agus F, Aytan N, et al. Inflammation and neuronal gene expression changes differ in early versus late chronic traumatic encephalopathy brain. BMC Medical Genomics. 2023;16:49. https://doi.org/10.1186/s12920-023-01471-5  
7. Garza R, Sharma Y, Atacho DAM, et al. Single-cell transcriptomics of human traumatic brain injury reveals activation of endogenous retroviruses in oligodendroglia. Cell Reports. 2023;42:113395. https://doi.org/10.1016/j.celrep.2023.113395  
8. Attilio PJ, Snapper DM, Rusnak M, et al. Transcriptomic Analysis of Mouse Brain After Traumatic Brain Injury Reveals That the Angiotensin Receptor Blocker Candesartan Acts Through Novel Pathways. Frontiers in Neuroscience. 2021;15:636259. https://doi.org/10.3389/fnins.2021.636259
"""


def add_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_markdown_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    pattern = re.compile(r"\[Fig\. (\d+)\]\(#fig-(\d+)\)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        add_hyperlink(p, f"Fig. {m.group(1)}", f"fig_{m.group(2)}")
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def markdown_to_docx(md: str, out_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bookmark_next: str | None = None
    bookmark_id = 1
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        anchor_match = re.match(r'<a id="fig-(\d+)"></a>', line)
        if anchor_match:
            bookmark_next = f"fig_{anchor_match.group(1)}"
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("!["):
            img_match = re.search(r"\]\(([^)]+)\)", line)
            if img_match:
                img_rel = img_match.group(1).replace("../figures/", "")
                img_path = FIGDIR / img_rel
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6.6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
                rows.append(parts)
                i += 1
            header = rows[0]
            data = rows[2:]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, h in enumerate(header):
                table.rows[0].cells[j].text = h
            for row in data:
                cells = table.add_row().cells
                for j, value in enumerate(row[: len(header)]):
                    cells[j].text = value
            i -= 1
        elif line.startswith("**Fig. "):
            caption = re.sub(r"\*\*", "", line)
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if bookmark_next:
                add_bookmark(p, bookmark_next, bookmark_id)
                bookmark_id += 1
                bookmark_next = None
        else:
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            add_markdown_paragraph(doc, clean)
        i += 1
    doc.save(out_path)


def write_reference_format_update() -> None:
    refs = [
        ("Zhao 2024 TBI bioinformatics-validation template", "09_Zhao2024_TBI_bioinfo_clinical_validation_template.pdf", "borrowed IMRAD structure, candidate ranking, disease-severity discussion"),
        ("Thomas 2022 acute TBI severity omics template", "10_Thomas2022_acute_TBI_metabolomics_template.pdf", "borrowed severity/time-window framing and cautious clinical interpretation"),
        ("Liu 2023 original disulfidptosis mechanism", "01_Liu2023_original_disulfidptosis.pdf", "borrowed SLC7A11/SLC3A2-NADPH-actin logic"),
        ("Machesky 2023 disulfidptosis commentary", "05_Machesky2023_NCB_deadly_actin_collapse_disulfidptosis.pdf", "borrowed actin-collapse conceptual language"),
        ("Zhang 2014 mouse TBI model", "04_Zhang2013_H2S_neuroprotection_TBI.pdf", "borrowed mouse TBI validation window and tissue readout style"),
    ]
    lines = ["# v3 参考格式论文与套图借鉴记录", ""]
    for title, fname, note in refs:
        pdf = WORKDIR / "references" / fname
        lines.append(f"- [{title}]({pdf.resolve()}): {note}.")
    lines.append("")
    lines.append("v3 图件改为多面板套图：Fig. 1 为研究设计和证据边界，Fig. 2 为人脑多数据集，Fig. 3 为急性 severe TBI 与 marker-proxy，Fig. 4 为小鼠 CCI 时空，Fig. 5 为外周严重程度，Fig. 6 为综合优先级和验证路线。")
    (REPORTDIR / "reference_format_papers_for_review_v3_20260604.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    configure_style()
    print("Loading existing human brain focused results...", flush=True)
    evidence = pd.read_csv(TABLEDIR / "time_severity_extension_matrix_20260604.csv")
    human, human_summary, _cte = load_human_focus()
    print("Running focused GSE223245 severity analysis...", flush=True)
    gse223 = analyze_gse223245()
    print("Loading GSE209552 acute severe TBI analysis...", flush=True)
    gse209 = load_gse209552()
    print("Loading GSE163415 mouse CCI analysis...", flush=True)
    gse163 = load_gse163415()
    print("Building integrated priority table...", flush=True)
    priority = integrated_priority(human_summary, gse209, gse163, gse223)

    print("Drawing multi-panel figures...", flush=True)
    make_fig1(evidence)
    make_fig2_human(human, human_summary)
    make_fig3_gse209552(gse209)
    make_fig4_mouse(gse163)
    make_fig5_severity(gse223)
    make_fig6_integrated(priority, gse209)

    print("Writing bilingual manuscripts...", flush=True)
    tables = build_tables_for_manuscript(evidence, priority, gse223, gse209["marker_corr"])
    vals = key_numbers(human, gse209, gse163, gse223, priority)
    zh = build_zh(vals, tables)
    en = build_en(vals, tables)
    zh_md = REPORTDIR / "TBI_disulfidptosis_enhanced_manuscript_v3_ZH_20260604.md"
    en_md = REPORTDIR / "TBI_disulfidptosis_enhanced_manuscript_v3_EN_20260604.md"
    zh_docx = REPORTDIR / "TBI_disulfidptosis_enhanced_manuscript_v3_ZH_20260604.docx"
    en_docx = REPORTDIR / "TBI_disulfidptosis_enhanced_manuscript_v3_EN_20260604.docx"
    zh_md.write_text(zh, encoding="utf-8")
    en_md.write_text(en, encoding="utf-8")
    markdown_to_docx(zh, zh_docx)
    markdown_to_docx(en, en_docx)
    write_reference_format_update()

    manifest = pd.DataFrame(
        [
            {"type": "report", "path": str(zh_md), "description": "Chinese enhanced v3 Markdown manuscript"},
            {"type": "report", "path": str(zh_docx), "description": "Chinese enhanced v3 Word manuscript"},
            {"type": "report", "path": str(en_md), "description": "English enhanced v3 Markdown manuscript"},
            {"type": "report", "path": str(en_docx), "description": "English enhanced v3 Word manuscript"},
            {"type": "report", "path": str(REPORTDIR / "reference_format_papers_for_review_v3_20260604.md"), "description": "Reference-format papers and borrowed layout elements"},
        ]
        + [
            {"type": "figure", "path": str(FIGDIR / f"Fig{i}_v3_{suffix}_20260604.png"), "description": "v3 multi-panel figure"}
            for i, suffix in []
        ]
    )
    figure_rows = []
    for p in sorted(FIGDIR.glob("Fig*_v3_*_20260604.png")):
        figure_rows.append({"type": "figure", "path": str(p), "description": "v3 multi-panel PNG"})
        pdf = p.with_suffix(".pdf")
        if pdf.exists():
            figure_rows.append({"type": "figure_pdf", "path": str(pdf), "description": "v3 multi-panel vector PDF"})
    table_rows = [{"type": "table", "path": str(p), "description": "v3 supporting analysis table"} for p in sorted(TABLEDIR.glob("v3_*_20260604.csv"))]
    manifest = pd.concat([manifest, pd.DataFrame(figure_rows + table_rows)], ignore_index=True)
    manifest.to_csv(REPORTDIR / "enhanced_v3_artifact_manifest_20260604.csv", index=False, encoding="utf-8-sig")
    print("Generated v3 manuscript package")
    print(zh_docx)
    print(en_docx)
    print(REPORTDIR / "enhanced_v3_artifact_manifest_20260604.csv")


if __name__ == "__main__":
    main()
