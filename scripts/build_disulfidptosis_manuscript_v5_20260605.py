from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_disulfidptosis_manuscript_v4_20260605 as v4


WORKDIR = v4.WORKDIR
TABLEDIR = v4.TABLEDIR
FIGDIR = v4.FIGDIR
REPORTDIR = v4.REPORTDIR


def figure_block(n: int, filename: str, caption: str) -> str:
    return f'<a id="fig-{n}"></a>\n\n![Fig. {n}](../figures/{filename}.png)\n\n**Fig. {n}. {caption}**'


def table_caption(number: int, title: str, detail: str) -> str:
    return f"**Table {number}. {title}.** {detail}"


def table_note(text: str) -> str:
    return f"Note. {text}"


def display_claim_map() -> pd.DataFrame:
    rows = [
        {
            "display_item": "Fig. 1",
            "question_answered": "How is the hypothesis constrained?",
            "primary_content": "evidence tiers, time/region/severity axes, fixed 8-gene panel and claim guardrails",
            "claim_boundary": "Defines the analysis framework; does not provide statistical evidence by itself",
        },
        {
            "display_item": "Fig. 2",
            "question_answered": "Which human brain contexts support the transporter-actin axis?",
            "primary_content": "human acute, chronic CTE, external CTE and remote regional comparisons",
            "claim_boundary": "FDR-supported only for the CTE-stage subset; other human signals are exploratory",
        },
        {
            "display_item": "Fig. 3",
            "question_answered": "What is visible in acute severe TBI brain tissue?",
            "primary_content": "module scores, 8-gene logFC, pathway enrichment summary and bulk marker proxies",
            "claim_boundary": "Small-sample directional support, not a diagnostic or causal claim",
        },
        {
            "display_item": "Fig. 4",
            "question_answered": "When and where is the mouse CCI support strongest?",
            "primary_content": "3DPI/29DPI, hippocampus/thalamus/hypothalamus and treatment-stratified logFC patterns",
            "claim_boundary": "Animal supportive evidence; prioritizes validation windows rather than proving human mechanism",
        },
        {
            "display_item": "Fig. 5",
            "question_answered": "Does injury severity add an independent axis?",
            "primary_content": "GSE223245 PCA, module scores, gene-level severity trends and module-level correlations",
            "claim_boundary": "Peripheral blood/PBMC severity context only; not brain mechanism evidence",
        },
        {
            "display_item": "Fig. 6",
            "question_answered": "Which genes and readouts should be validated first?",
            "primary_content": "integrated evidence matrix, ranked genes, mechanistic workflow and top marker-proxy correlations",
            "claim_boundary": "Experimental prioritization, not a predictive model",
        },
        {
            "display_item": "Fig. 7",
            "question_answered": "Do comparison-level and validation-window summaries support the same hierarchy?",
            "primary_content": "human comparison directionality, mouse window ranking, treatment-stratified module modulation and upgrade priorities",
            "claim_boundary": "Evidence audit linking existing results to next-step validation design",
        },
        {
            "display_item": "Fig. 8",
            "question_answered": "What evidence is still missing for a true disulfidptosis claim?",
            "primary_content": "claim ladder, cell-proxy heatmap, wet-lab time-course design and peripheral severity genes",
            "claim_boundary": "Makes the inferential ceiling explicit: transcriptomics supports a candidate state, not mechanism proof",
        },
    ]
    out = pd.DataFrame(rows)
    out.to_csv(TABLEDIR / "v5_display_claim_map_20260605.csv", index=False, encoding="utf-8-sig")
    return out


def build_table_strings(data: dict[str, pd.DataFrame], ext: dict[str, pd.DataFrame]) -> dict[str, str]:
    display = display_claim_map()
    dataset_cols = ["dataset_or_source", "species", "tissue_or_model", "time_or_course", "severity_or_stage", "current_role_in_manuscript", "guardrail"]
    human_cols = ["dataset", "comparison_short", "mean_effect", "positive_genes", "nominal_genes", "fdr_genes", "n_genes"]
    mouse_cols = ["time", "region", "treatment", "mean_logFC", "positive_genes", "nominal_genes", "panel_FDR_genes", "min_p"]
    severity_cols = ["gene_symbol", "spearman_r_severity_all_groups", "p_value_severity_all_groups", "FDR_severity_all_groups", "spearman_r_TBI_only", "p_value_TBI_only"]
    marker_cols = ["module", "celltype_proxy", "spearman_r", "p_value", "FDR", "interpretation"]
    priority_cols = ["gene_symbol", "human_FDR_comparisons", "human_nominal_comparisons", "GSE209552_acute_logFC", "mouse_nominal_units", "GSE223245_severity_r_all", "integrated_evidence_score"]
    validation_cols = ["evidence_gap", "readout", "best_current_support", "required_upgrade", "priority"]
    display_cols = ["display_item", "question_answered", "primary_content", "claim_boundary"]
    return {
        "datasets": v4.markdown_table(data["evidence"], dataset_cols),
        "display": v4.markdown_table(display, display_cols),
        "human": v4.markdown_table(ext["human_comp"], human_cols),
        "mouse": v4.markdown_table(ext["mouse_units"], mouse_cols, max_rows=12),
        "severity": v4.markdown_table(data["sev_gene"].sort_values("p_value_severity_all_groups"), severity_cols),
        "marker": v4.markdown_table(ext["marker_priority"], marker_cols, max_rows=12),
        "priority": v4.markdown_table(data["priority"], priority_cols),
        "validation": v4.markdown_table(ext["validation_matrix"], validation_cols),
    }


CAPTIONS_ZH = {
    0: ("Graphical abstract. 该图用一条从公共转录组到机制验证的证据链概括本文主线：先固定 SLC3A2/SLC7A11 与 6 个肌动蛋白骨架终点基因，再按急性 severe TBI、慢性 CTE、鼠 CCI、外周严重程度和 bulk marker-proxy 分层解释。图中的最后一步不是“已证明双硫死亡”，而是指出 3DPI hippocampus/cortex 是最值得进入蛋白、F-actin、氧化还原和细胞共定位验证的窗口。"),
    1: ("Study design and inferential guardrails for the fixed transporter-actin panel. Panel A shows the evidence-layer design in which human acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets are assigned different roles rather than merged into one analysis. Panel B lists the biological axes that the manuscript attempts to answer: acute injury, remote region, CTE course, external CTE validation, mouse CCI, blood severity and planned validation. Panel C places the eight genes along the mechanistic route from cystine import to actin regulation and cytoskeletal endpoint. Panel D defines the wording rules used throughout the paper, separating FDR-supported associations from nominal clues, animal support, peripheral context and bulk marker-proxy prioritization."),
    2: ("Human brain transcriptomic evidence across acute, chronic and regional comparisons. Panel A is the 8-gene effect map; rows represent comparison contexts and columns represent the pre-fixed genes, with filled or open marks indicating different levels of statistical support. Panel B focuses on the GSE193407 continuous CTE-stage trend and shows that SLC3A2, SLC7A11, WASF2 and TLN1 are the FDR-supported stage-associated genes. Panel C summarizes recurrence across human comparisons by combining nominal support count with maximum absolute effect size, showing why the transporter-regulatory genes and some cytoskeletal endpoint genes should be interpreted differently."),
    3: ("Acute severe TBI evidence from GSE209552. Panel A compares sample-level transporter/stress, actin-target and mitochondrial-synergy module scores between TBI and control samples. Panel B shows gene-level logFC for all eight genes; ACTB, MYH9 and MYL6 carry the strongest nominal cytoskeletal directionality. Panel C places the disulfidptosis gene set in the ranked pathway context and shows that the disulfidptosis-like module is more positive than several comparator pathways. Panel D shows bulk marker-proxy correlations, which are used only to prioritize later cell-type co-localization, not to assign the signal to a cell type."),
    4: ("Mouse CCI spatiotemporal support in GSE163415. Panel A summarizes module score shifts across 3DPI/29DPI, brain regions and treatment strata, showing the most concentrated transporter and actin changes in 3DPI hippocampus. Panel B resolves the pattern at the single-gene level and shows that WASF2 and TLN1 are particularly strong within the 3DPI hippocampal units. Panel C shows pathway-level overlap in mouse CCI, where actin cytoskeleton/tension and amino-acid transport/cystine axes rank above broader oxidative and mitochondrial categories. Panel D ranks the eight genes by mouse-derived support, highlighting SLC3A2, TLN1, SLC7A11 and WASF2 as practical validation targets."),
    5: ("Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data. Panel A shows PCA of the eight-gene expression matrix, illustrating partial but not complete separation of control and TBI severity groups. Panel B compares module scores across control, mild, moderate and severe groups. Panel C shows gene-level Spearman correlations across severity 0-3; FLNA and SLC3A2 are the strongest negative trends, while MYL6 is a positive TBI-only clue. Panel D summarizes module-level severity associations, which support a severity-aware narrative but remain peripheral rather than brain-mechanistic evidence."),
    6: ("Integrated prioritization of genes and validation readouts. Panel A combines human FDR support, human nominal support, mouse nominal support, acute logFC and blood severity association into a dry-evidence matrix. Panel B converts the same evidence into an integrated validation score, ranking SLC3A2, SLC7A11, WASF2 and TLN1 above the remaining cytoskeletal endpoint genes. Panel C translates the ranking into an experimental sequence: 3DPI cortex/hippocampus qPCR/WB, cytoskeleton-redox assays, cell co-localization and finally a same-window mechanistic criterion. Panel D displays top marker-proxy correlations used to prioritize CD31 and NeuN co-staining, while preserving the limitation that bulk correlations do not prove cell-type origin."),
    7: ("Evidence audit and validation priorities. Panel A audits human comparison-level directionality and shows that acute severe TBI has the largest mean effect but lacks FDR support, whereas chronic CTE-stage comparisons provide the clearest FDR-supported human evidence. Panel B ranks mouse CCI validation windows and confirms that 3DPI hippocampus is the most coherent animal support layer. Panel C compares Drug-Veh module differences within GSE163415; these data are interpreted as stage-stratified modulation rather than treatment efficacy. Panel D lists the mechanistic gaps that must be upgraded before the manuscript can claim true disulfidptosis."),
    8: ("Claim ladder, cellular-context clues and wet-lab design. Panel A defines four evidence levels from peripheral clue to mechanistic proof; the current study is intentionally kept below the mechanistic-proof tier. Panel B expands the GSE209552 marker-proxy map and highlights positive endothelial and negative neuronal correlations as co-localization priorities. Panel C converts the transcriptomic results into a wet-lab time-course design, with 2D-7D and especially 3D prioritized for transporter, actin/F-actin, redox and co-localization readouts. Panel D shows peripheral severity gene trends, emphasizing that FLNA and SLC3A2 severity associations are informative but cannot replace brain-tissue validation."),
}


CAPTIONS_EN = {
    0: ("Graphical abstract. The figure summarizes the manuscript logic from public transcriptomes to mechanistic validation. The analysis fixes SLC3A2/SLC7A11 and six actin-cytoskeletal endpoint genes, stratifies evidence by acute severe TBI, chronic CTE, mouse CCI, peripheral severity and bulk marker proxies, and finally nominates 3DPI hippocampus/cortex for protein, F-actin, redox and cell-localization validation. The final node is a validation target, not proof that disulfidptosis has already occurred."),
    1: ("Study design and inferential guardrails for the fixed transporter-actin panel. Panel A shows the evidence-layer design, assigning acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets to distinct roles. Panel B lists the time, region and severity axes addressed by the study. Panel C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints. Panel D states the wording rules used in the manuscript, separating FDR-supported associations, nominal clues, animal support, peripheral context and bulk marker-proxy prioritization."),
    2: ("Human brain transcriptomic evidence across acute, chronic and regional comparisons. Panel A shows the 8-gene effect map across comparison contexts. Panel B focuses on GSE193407 continuous CTE stage and identifies SLC3A2, SLC7A11, WASF2 and TLN1 as FDR-supported stage-associated genes. Panel C summarizes recurrence across human comparisons by combining nominal support count and maximum absolute effect, distinguishing transporter-regulatory genes from downstream cytoskeletal endpoint genes."),
    3: ("Acute severe TBI evidence from GSE209552. Panel A compares sample-level transporter/stress, actin-target and mitochondrial-synergy module scores between TBI and control samples. Panel B shows gene-level logFC for all eight genes, with ACTB, MYH9 and MYL6 carrying the strongest nominal cytoskeletal directionality. Panel C places the disulfidptosis gene set in a ranked pathway context. Panel D shows bulk marker-proxy correlations used for later co-localization prioritization, not for direct cell-type assignment."),
    4: ("Mouse CCI spatiotemporal support in GSE163415. Panel A summarizes module score shifts across 3DPI/29DPI, brain regions and treatment strata. Panel B resolves the pattern at the single-gene level, with WASF2 and TLN1 prominent in 3DPI hippocampal units. Panel C summarizes pathway overlap in mouse CCI. Panel D ranks the eight genes by mouse-derived support, highlighting SLC3A2, TLN1, SLC7A11 and WASF2 as practical validation targets."),
    5: ("Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data. Panel A shows PCA of the eight-gene expression matrix. Panel B compares module scores across control, mild, moderate and severe groups. Panel C shows gene-level Spearman correlations across severity 0-3. Panel D summarizes module-level severity associations. These results support a severity-aware peripheral narrative but remain separate from brain-mechanism evidence."),
    6: ("Integrated prioritization of genes and validation readouts. Panel A combines human FDR support, human nominal support, mouse nominal support, acute logFC and blood severity association into a dry-evidence matrix. Panel B converts the same information into an integrated validation score. Panel C translates the ranking into a wet-lab sequence. Panel D displays top marker-proxy correlations used to prioritize co-staining while preserving the limitation that bulk correlations do not prove cell-type origin."),
    7: ("Evidence audit and validation priorities. Panel A audits human comparison-level directionality. Panel B ranks mouse CCI validation windows. Panel C compares Drug-Veh module differences within GSE163415 and is interpreted as stage-stratified modulation rather than treatment efficacy. Panel D lists the mechanistic gaps that must be closed before a true disulfidptosis claim is defensible."),
    8: ("Claim ladder, cellular-context clues and wet-lab design. Panel A defines four evidence levels from peripheral clue to mechanistic proof. Panel B expands the GSE209552 marker-proxy map. Panel C converts transcriptomic findings into a wet-lab time-course design. Panel D shows peripheral severity gene trends, emphasizing that blood-based FLNA and SLC3A2 trends cannot replace matched brain-tissue validation."),
}


FIG_FILES = {
    0: "GraphicalAbstract_v4_TBI_disulfidptosis_20260605",
    1: "Fig1_v3_study_design_evidence_layers_20260604",
    2: "Fig2_v3_human_brain_multidataset_8gene_20260604",
    3: "Fig3_v3_GSE209552_acute_severe_and_marker_proxy_20260604",
    4: "Fig4_v3_GSE163415_mouse_CCI_spatiotemporal_suite_20260604",
    5: "Fig5_v3_GSE223245_peripheral_severity_suite_20260604",
    6: "Fig6_v3_integrated_priority_validation_suite_20260604",
    7: "Fig7_v4_evidence_audit_and_validation_priorities_20260605",
    8: "Fig8_v4_claim_ladder_cell_proxy_and_wetlab_design_20260605",
}


def fig(n: int, lang: str) -> str:
    captions = CAPTIONS_ZH if lang == "zh" else CAPTIONS_EN
    return figure_block(n, FIG_FILES[n], captions[n])


def zh_manuscript(vals: dict[str, str], tables: dict[str, str]) -> str:
    return f"""# TBI 后双硫死亡样转运-肌动蛋白骨架应激的时空定位、损伤程度关联与细胞类型优先级：基于人类 TBI/CTE 与小鼠 CCI 公共转录组的整合生物信息学研究

## 摘要

背景：创伤性脑损伤（traumatic brain injury, TBI）后的继发性损伤具有强烈的时间、脑区、损伤程度和细胞背景异质性。双硫死亡由 SLC7A11/SLC3A2 介导的胱氨酸摄入、NADPH 消耗、二硫键压力和肌动蛋白骨架崩塌共同驱动，但公共转录组不能直接证明蛋白二硫键化或 F-actin 崩塌。因此，本研究将问题表述为“TBI 后是否存在双硫死亡样转运-肌动蛋白骨架应激”，并预先固定 SLC3A2、SLC7A11、WASF2、TLN1、ACTB、MYH9、MYL6 和 FLNA 8 个关键基因。方法：整合 GSE209552、GSE193407、GSE319253、GSE104687、GSE163415 和 GSE223245，按急性 severe TBI、慢性 CTE 病程、远期脑区、鼠 CCI 时间-脑区和外周严重程度分层分析。统计解释严格区分 FDR 支持与 nominal 线索，同 donor 多脑区样本不作为独立主分析，动物和外周血数据仅作为旁证或严重程度背景。结果：慢性 CTE 病程中最稳定的 FDR 支持来自 GSE193407 stage trend，具体为 {vals['stage_details']}；late CTE stage 3-4 vs stage 0 中为 {vals['late_details']}。急性 severe TBI 中所有 8 个基因方向均为正，较强骨架终点信号包括 {vals['acute_genes']}。鼠 CCI 中最集中窗口为 {vals['mouse_top']}，提示 3DPI hippocampus/cortex 是后续验证优先窗口。外周血严重程度只提供 peripheral clue，模块层面为 {vals['severity_module']}，单基因最强为 {vals['severity_gene']}。bulk marker-proxy 最强线索为 {vals['marker_top']}，用于排序共定位实验而非细胞归属。结论：现有公共数据支持一个候选的 TBI/CTE 双硫死亡样转运-骨架应激模型，但尚未达到机制证明层级。下一步应在 3DPI cortex/hippocampus 同时检测 SLC3A2/SLC7A11 蛋白、WASF2/TLN1 与 ACTB/MYH9/MYL6/FLNA 终点、F-actin、NADPH/GSH 和 NeuN/GFAP/IBA1/OLIG2/CD31 共定位。

关键词：创伤性脑损伤；慢性创伤性脑病；双硫死亡；SLC7A11；SLC3A2；肌动蛋白骨架；损伤时间；损伤严重程度；公共转录组

## 图形摘要

{fig(0, 'zh')}

## 引言

TBI 并不是一个单一时间点的分子事件。机械性原发损伤之后，脑组织会在数小时内进入能量代谢危机、离子稳态破坏、兴奋性毒性、线粒体损伤和氧化还原压力状态；数天至数周内，血脑屏障破坏、血管反应、胶质细胞激活、轴突和突触结构改变、细胞骨架重塑以及炎症持续化继续塑造继发性损伤。慢性或反复轻度损伤背景下，CTE 又表现为长期神经炎症、神经退行性改变和细胞组成重排。因此，若只把全部 TBI 样本与 control 相比，容易把不同阶段、不同脑区和不同细胞背景下的信号混成一个平均值。更适合机制论文的设计应当把问题拆成“何时、何处、何种损伤程度、哪类细胞背景更值得验证”。

双硫死亡为这个问题提供了新的机制入口。原始机制研究表明，SLC7A11 高表达细胞在葡萄糖不足时持续摄入胱氨酸，胱氨酸还原消耗 NADPH，进而造成二硫键压力累积，最终使肌动蛋白骨架蛋白发生异常二硫键化并导致 F-actin 崩塌。TBI 后脑组织本身存在能量供应不足、氧化还原失衡、线粒体障碍和骨架损伤，因此从病理生理背景看，双硫死亡样应激具有合理性。然而，公共 bulk RNA-seq 或 microarray 只能看到 mRNA 层面的入口和部分终点变化，不能替代蛋白二硫键化、F-actin 形态或细胞死亡表型检测。本文因此避免直接宣称“TBI 后发生双硫死亡”，而是使用“候选双硫死亡样转运-肌动蛋白骨架应激”这一更可验证、也更符合证据边界的表述。

本研究没有扩大到所有氧化应激、硫代谢或细胞死亡基因，而是把假说收束到 8 个预先固定的关键基因。SLC3A2 和 SLC7A11 定义胱氨酸转运入口，WASF2 和 TLN1 定义肌动蛋白调控、黏附和张力传递节点，ACTB、MYH9、MYL6 和 FLNA 定义更接近结构终点和机械应力的骨架蛋白。这样的设计牺牲了一部分通路覆盖面，但换来了清晰的可检验性：如果 TBI/CTE 背景下确有与双硫死亡相邻的应激状态，应当能在特定时间窗、脑区或病程阶段看到转运入口与骨架终点的协同或阶段性改变。

本文的目标是把现有公共数据转化为一篇可以继续导向湿实验验证的论文，而不是把公共转录组过度包装为机制证明。为此，文章采用参考 TBI 生信验证论文和急性 TBI severity omics 论文的结构：先说明数据层级和 claim boundary，再按 human brain、acute severe TBI、mouse CCI、peripheral severity、cell-type proxy 和 integrated validation priority 逐层推进。所有图表都被写入结果论证链中，每个多面板图的子图均在图注中解释，并在正文中承担一个明确问题。

## 材料与方法

### 数据集纳入与证据分层

本研究纳入六个 GEO 公共数据层。GSE209552 代表人类急性 severe TBI 外科切除脑组织，损伤后时间约 4 小时至 8 天，适合提供急性重型脑组织方向性证据。GSE193407 代表人类 BA9 脑区 CTE stage 0-4，适合分析慢性病程梯度。GSE319253 代表 superior frontal cortex CTE vs control，可作为慢性外部验证层。GSE104687 是远期 TBI 多脑区死后脑组织数据，由于同一 donor 多脑区样本不能作为独立观测，因此只作为空间探索层。GSE163415 是小鼠 CCI 数据，覆盖 3DPI/29DPI、hippocampus/thalamus/hypothalamus 和 vehicle/drug 分层，用于跨物种时间-脑区旁证。GSE223245 是 whole blood/PBMC mild、moderate 和 severe TBI 微阵列数据，只用于外周严重程度线索。Table 1 总结了这些数据层级及其不可越界解释。

{table_caption(1, 'Dataset layers and evidence boundaries', 'The table specifies the biological role of each dataset before analysis. This pre-specified role is essential because acute severe TBI, chronic CTE stage, remote postmortem regions, mouse CCI and peripheral blood severity are not interchangeable measures of the same process.')}

{tables['datasets']}

{table_note('GSE104687 multi-region samples were treated as exploratory because multiple regions from the same donor are not independent observations. GSE223245 was retained only as a peripheral severity-aware layer.')}

### 候选基因、模块评分和统计原则

候选基因在分析前固定为 SLC3A2、SLC7A11、WASF2、TLN1、ACTB、MYH9、MYL6 和 FLNA。为了避免把一个长基因列表误写为机制，本文只构建三个模块读数：8 基因总模块、SLC3A2/SLC7A11 转运入口模块，以及 WASF2/TLN1/ACTB/MYH9/MYL6/FLNA 肌动蛋白骨架模块。模块评分采用样本内 z-score 标准化后求均值；单基因比较使用现有差异分析或 Welch t 检验；严重程度趋势使用 Spearman 相关；多重检验使用 Benjamini-Hochberg FDR。正文中只有 FDR<0.05 被称为 FDR-supported，P<0.05 但 FDR 未达标只称为 nominal exploratory clue。

### 图表与正文整合原则

参考论文中，图并不是在结果之后附加展示，而是直接决定 Results 的段落顺序：workflow 图定义分析边界，主结果图回答核心生物学问题，后续图依次处理严重程度、验证窗口和机制升级。因此 v5 版本将 Fig. 1 through Fig. 8 重新写入结果链条。Table 2 记录每个 display item 对应的问题、展示内容和 claim boundary，用于保证图、表和正文分析不再割裂。

{table_caption(2, 'Display-item-to-claim map used for the v5 rewrite', 'Each figure is assigned one primary biological question and one explicit claim boundary. This table was used as a writing audit so that captions, results paragraphs and discussion claims point to the same evidence layer.')}

{tables['display']}

{table_note('This table is not a statistical result; it is a manuscript-level evidence audit following the PaperOrchestra-style idea-log-figure-refinement workflow.')}

## 结果

### 证据层设计把“双硫死亡”转换为可检验的转运-骨架轴

Fig. 1 是整篇文章的逻辑起点，而不是装饰性 workflow。Fig. 1A 将人类急性 severe TBI、慢性 CTE、鼠 CCI 和外周严重程度放在不同证据层，而不是把它们合并为一个“大 TBI 队列”。这种分层很关键，因为 GSE209552 代表急性重型脑组织，GSE193407 代表慢性病程，GSE163415 代表动物模型时间-脑区验证，GSE223245 代表外周血严重程度背景；它们回答的问题相邻但不相同。Fig. 1B 进一步把文章问题拆成 acute、remote region、CTE course、external CTE、mouse CCI、blood severity 和 planned validation 七个轴，说明本文不是简单筛选差异基因，而是围绕“何时、何处、何种损伤程度、何种验证路径”组织证据。

Fig. 1C 是本文避免泛化的关键：8 个基因被放置在从 cystine import 到 actin control 再到 cytoskeletal endpoint 的机制坐标上。SLC3A2 和 SLC7A11 代表入口是否打开，WASF2 和 TLN1 代表骨架调控和张力传递是否参与，ACTB、MYH9、MYL6 和 FLNA 代表结构终点是否出现表达或机械应激线索。Fig. 1D 则定义了全文用语规则：human brain FDR 只能写 association，nominal signal 只能写 candidate clue，mouse CCI 是 supportive evidence，peripheral blood 是 severity context，bulk marker proxy 只用于 prioritization。换句话说，Fig. 1 先规定“哪些结论可以写，哪些不能写”，后续所有结果都在这个边界内展开。

{fig(1, 'zh')}

### 人类脑组织结果显示慢性 CTE 病程中的转运-骨架调控轴最稳定

人类脑组织是本文证据链中最重要、但也最需要分层解释的部分。Fig. 2A 显示，急性 severe TBI、CTE stage trend、late CTE vs stage0、external CTE 和远期 TBI 脑区比较的 8 基因效应方向并不完全相同。急性 severe TBI 行的整体颜色最深，说明平均效应大，但这来自小样本急性脑组织；CTE stage trend 和 late CTE 行的颜色更温和，却有更多 FDR 标记，说明慢性病程中的统计稳定性更强。远期 TBI 多脑区行接近零，提示 GSE104687 在 donor-level 主分析下不应被当作强证据来源。

Fig. 2B 解释了为什么本文把慢性 CTE 病程视为最稳健的人脑层。GSE193407 continuous stage trend 中，达到 FDR 支持的基因为 {vals['stage_details']}。这四个基因构成一个清晰轴：SLC3A2/SLC7A11 对应胱氨酸转运入口，WASF2/TLN1 对应肌动蛋白调控和张力结构。late CTE stage 3-4 vs stage 0 中，FDR 支持基因为 {vals['late_details']}，其中 SLC7A11、WASF2 和 SLC3A2 延续了 stage trend 的方向。Fig. 2C 则把“出现多少次 nominal 支持”和“最大效应量”结合起来，显示 MYH9、ACTB、MYL6 等骨架终点虽然在人类比较中有较高方向性或复现线索，但 FDR 稳定性不如 SLC3A2/SLC7A11/WASF2/TLN1。因此，本文不把 8 个基因写成一个同质模块，而是把它拆成入口、调控和终点三个层次解释。

{fig(2, 'zh')}

{table_caption(3, 'Human comparison-level evidence audit', 'The table summarizes each human-brain comparison at the 8-gene panel level. Mean effect indicates the average direction across the fixed panel; positive genes, nominal genes and FDR genes separate directionality, exploratory support and multiple-testing-supported evidence.')}

{tables['human']}

{table_note('The strongest FDR-supported human evidence comes from GSE193407 CTE-stage analyses. GSE209552 has a larger mean effect but remains small-sample acute directional evidence.')}

### 急性 severe TBI 更像早期骨架终点应激，而不是完整机制证明

Fig. 3 将 GSE209552 从样本模块、单基因、路径和 cell proxy 四个角度放在同一张图中，因此它不能只被概括为“8 个基因上调”。Fig. 3A 的模块评分显示，TBI 样本在 transporter/stress 和 actin targets 两个模块上整体高于 control，但点的离散度也很大，提示急性 severe TBI 脑组织中存在样本异质性。Fig. 3B 的单基因 logFC 更直观：{vals['acute_genes']}，其中 ACTB、MYH9 和 MYL6 等骨架终点在 nominal 层面较突出。这个结果与急性损伤后机械应力、血管破坏、胶质反应和细胞骨架重塑的病理背景吻合，但由于 TBI n=4、control n=3，它不能承担生物标志物发现或因果证明的角色。

Fig. 3C 把 disulfidptosis-like gene set 放到几个相邻机制模块的排序中。与 copper/cuproptosis、methionine metabolism 和 H2S enzymes 相比，disulfidptosis 模块的 enrichment score 更高，说明急性 severe TBI 的现有转录组读数确实更靠近本文预设的转运-骨架轴，而不是一个泛化的硫代谢或细胞死亡信号。Fig. 3D 则给出 cell-type prioritization 的第一层线索：8-gene、transporter 和 actin 模块均与 endothelial proxy 呈正相关，与 neuron proxy 呈负相关。这里的正确解释不是“内皮细胞发生双硫死亡”，而是“CD31 与 NeuN 应进入后续共定位实验”，因为 bulk marker-proxy 无法区分细胞比例变化和细胞内表达状态变化。

{fig(3, 'zh')}

### 鼠 CCI 把可验证窗口收敛到 3DPI hippocampus/cortex

Fig. 4 回答本文的“何时、何处”问题。Fig. 4A 显示，不同 time-region-treatment 单元的 transporter 和 actin 模块并不均匀，3DPI Hipp Drug 和 3DPI Hipp Veh 的模块色阶最深，提示 3DPI hippocampus 同时具有入口和骨架终点信号。Fig. 4B 的单基因热图显示，3DPI Hipp all 和 3DPI Hipp Drug 中 WASF2 与 TLN1 尤为突出，SLC3A2/SLC7A11 和 MYH9/FLNA 也呈一致正向。这个模式与人类 CTE 中的 SLC3A2/SLC7A11/WASF2/TLN1 轴形成跨物种旁证，但仍只能写作 animal supportive evidence。

Fig. 4C 显示，mouse CCI 中与本文主线最接近的是 actin cytoskeleton and tension、amino-acid transport and cystine axis 以及 disulfidptosis core/PDF genes，强于更泛化的 oxidative stress/glutathione handling 和 mitochondrial respiratory support。这说明 3DPI mouse CCI 的信号不是单纯“炎症或氧化应激”，而是更靠近转运-骨架结构。Fig. 4D 进一步把小鼠支持转换为基因层验证优先级，SLC3A2、TLN1、SLC7A11 和 WASF2 位居前列。结合 v4 validation-window 排序，最强单元为 {vals['mouse_top']}；按脑区时间汇总，{vals['mouse_hipp']}。因此，后续湿实验的第一轮不应平均铺开所有时间点，而应优先检测 3DPI cortex/hippocampus，并用 7D 或 29DPI 作为持续性和转归观察。

{fig(4, 'zh')}

{table_caption(4, 'Top mouse CCI time-region-treatment validation windows', 'The table ranks GSE163415 units by the coherence of the fixed 8-gene panel. Mean logFC summarizes effect direction, positive genes count direction consistency, nominal genes and panel-FDR genes separate exploratory and panel-level statistical support.')}

{tables['mouse']}

{table_note('Mouse CCI is used to select validation windows. It supports the biological plausibility of the 3DPI hippocampal window but does not prove the same mechanism in human brain.')}

### Comparison-level audit explains why acute, chronic and mouse evidence should not be collapsed

Fig. 7 was rewritten in v5 as a bridge between the primary results and the validation plan. Fig. 7A shows that acute severe TBI has the largest mean 8-gene effect across human comparisons, but it has no FDR-supported genes in the current small dataset. In contrast, GSE193407 CTE stage trend and late CTE stage3-4 vs stage0 have more FDR-supported genes despite smaller average effect size. This is the central reason for the paper’s cautious wording: acute severe TBI provides directional biological plausibility, whereas chronic CTE provides the strongest human statistical support.

Fig. 7B independently ranks mouse CCI validation windows and again places 3DPI hippocampus at the top, showing that the same validation window emerges when the mouse dataset is summarized at the unit level rather than only by module heatmaps. Fig. 7C adds treatment-stratified information, but it is deliberately not written as drug efficacy. The strongest treatment-stratified readout is {vals['treatment_top']}，which suggests that later 29DPI transporter/stress modulation may persist or shift under treatment strata. Fig. 7D then translates these observations into evidence gaps: transporter entry, actin cytoskeletal endpoint, redox prerequisite and cell-type localization all remain high-priority upgrades before a true disulfidptosis claim can be made.

{fig(7, 'zh')}

### 外周血严重程度提供 severity-aware clue，但不能替代脑组织机制证据

Fig. 5 让损伤程度问题进入正文，但它的解释必须独立于脑组织机制。Fig. 5A 的 PCA 显示，whole blood/PBMC 的 8 基因表达不能把 control、mild、moderate 和 severe 完全分开，说明外周血信号存在个体差异和组内重叠。Fig. 5B 的模块评分显示 control、mild、moderate、severe 之间有一定梯度和离散，但并非单调增强。Fig. 5C 显示单基因层面最强全组 severity 相关为 {vals['severity_gene']}，SLC3A2 也呈负相关，而 MYL6 在 TBI-only 分析中有正向线索。Fig. 5D 的模块层面最强为 {vals['severity_module']}，但 FDR 未达到 0.05。

这些结果对论文的价值不在于证明脑内双硫死亡，而在于提醒后续设计必须纳入 injury severity。GSE209552 的 severe TBI、GSE193407 的 CTE stage 和 GSE223245 的 mild/moderate/severe blood groups 不是同一个 severity 尺度，不能强行合并为一个连续变量。合理写法是：severe acute brain tissue 支持早期方向性，CTE stage 支持慢性病程关联，peripheral blood 支持外周严重程度背景。这样既回答了“损伤程度”问题，又避免跨组织、跨病程的统计混淆。

{fig(5, 'zh')}

{table_caption(5, 'Peripheral blood 8-gene severity trends in GSE223245', 'The table reports gene-level Spearman correlations across severity 0-3 and TBI-only severity 1-3. FDR values are shown to separate robust peripheral severity associations from exploratory trends.')}

{tables['severity']}

{table_note('Because GSE223245 is whole blood/PBMC with small group sizes, these results are peripheral severity clues only. They do not establish brain-region or cell-type mechanism.')}

### 细胞类型线索应被写成共定位优先级，而不是细胞归属

bulk marker-proxy 是当前最容易被过度解释的部分。Table 6 显示，GSE209552 中最强线索为 {vals['marker_top']}，此外 actin module 与 endothelial proxy 也为正相关，8-gene module 与 neuron proxy 为负相关。Fig. 3D 和 Fig. 8B 都展示了这个模式。真正值得讨论的是它的实验含义：如果 endothelial proxy 与 8-gene module 同向，后续 IHC/IF 中必须把 SLC3A2/SLC7A11、WASF2/TLN1 或 F-actin 与 CD31 共同检测；如果 neuron proxy 呈负相关，则需要同时加入 NeuN 判断这是神经元丢失、神经元状态改变，还是非神经元细胞比例升高造成的 bulk 相关。仅凭 bulk marker-proxy，不能说“内皮细胞发生双硫死亡”或“神经元不参与”。

{table_caption(6, 'Bulk marker-proxy prioritization in acute severe TBI brain tissue', 'The table ranks Spearman correlations between module scores and marker-proxy cell-type scores in GSE209552. The interpretation column intentionally labels these findings as proxy clues because bulk tissue cannot separate cell abundance from within-cell transcriptional state.')}

{tables['marker']}

{table_note('The strongest nominal proxy clue points to endothelial markers, but FDR does not reach 0.05. Co-localization or annotated single-nucleus data are required for cell-type assignment.')}

### 综合优先级把表达入口、骨架调控和结构终点分开验证

Fig. 6 是将公共数据转化为实验计划的核心图。Fig. 6A 的 integrated dry-evidence matrix 显示，SLC3A2、SLC7A11、WASF2 和 TLN1 在 human FDR、human nominal、mouse nominal、acute logFC 和 blood severity r 之间形成较高综合支持，而 ACTB、MYH9、MYL6 和 FLNA 更偏向急性方向或外周严重程度线索。Fig. 6B 将这一矩阵压缩成 integrated evidence score，当前第一优先级为 {vals['priority_top']}。不过这个排名不是“生物标志物模型”，而是实验投入顺序。

Fig. 6C 将基因排名翻译成四步验证路径。第一步在 3DPI cortex/hippocampus 做 qPCR/WB，优先检测 SLC3A2、SLC7A11、WASF2 和 TLN1；第二步检测 F-actin、MYH9/MYL6/FLNA、NADPH/NADP+ 和 GSH/GSSG；第三步与 NeuN、GFAP、IBA1、OLIG2 和 CD31 共定位；第四步要求同一时间、同一脑区、同一细胞背景中入口、骨架、氧化还原和定位信号一致。Fig. 6D 把 marker-proxy 线索嵌入这个计划中，说明 CD31 与 NeuN 不是可选染色，而是决定机制归属的必要读数。

{fig(6, 'zh')}

{table_caption(7, 'Integrated 8-gene evidence priority', 'The table combines human FDR support, human nominal support, acute severe TBI directionality, mouse CCI support and peripheral severity association into an experimental priority score. The score is used to rank validation targets, not to classify patients or infer causality.')}

{tables['priority']}

{table_note('SLC3A2 ranks first because it combines human CTE-stage support, mouse CCI recurrence and a peripheral severity trend. Endpoint genes still require protein and F-actin validation even when mRNA evidence is weaker.')}

### 机制升级需要同一窗口中的转运、骨架、氧化还原和细胞定位证据

Fig. 8 把本文的 inferential ceiling 画出来。Fig. 8A 的 claim ladder 将证据分为 peripheral clue、public brain transcriptome、tissue validation 和 mechanistic proof 四级。当前论文最强只能达到 public brain transcriptome plus animal support：慢性 CTE stage 有 FDR 支持，急性 severe TBI 有方向性，mouse CCI 有时间-脑区旁证，但没有蛋白二硫键化和 F-actin 崩塌证据。Fig. 8B 的 marker-proxy heatmap 和 Fig. 8D 的 peripheral severity trends 都是优先级信息，而不是机制证明。

Fig. 8C 是后续实验最直接的时间课程建议。1h-12h 可以观察早期氧化还原和应激启动，1D-2D 可以观察入口和骨架变化形成，3D-7D 是最值得集中检测 transporter、actin/F-actin、redox 和 co-localization 的窗口，29DPI 更适合判断残留或转归。Table 8 将这些要求写成机制升级矩阵。只有当同一时间窗和脑区中同时观察到 SLC3A2/SLC7A11 蛋白变化、WASF2/TLN1 与骨架终点异常、F-actin 形态改变、NADPH/GSH 改变以及细胞类型共定位时，论文才能从“候选双硫死亡样应激”升级为“机制支持双硫死亡相关过程”。

{fig(8, 'zh')}

{table_caption(8, 'Mechanistic upgrade matrix for the next experimental stage', 'The table lists the missing evidence classes required to move from public transcriptomic association to tissue validation and then to mechanistic proof. Each row links a current evidence gap to a concrete experimental readout.')}

{tables['validation']}

{table_note('The priority column refers to the next experimental stage, not to clinical urgency. Protein, F-actin, redox and cell co-localization must be measured in matched tissue windows.')}

## 讨论

本研究的核心贡献不是发现某一个 TBI 基因，而是把一个新型细胞死亡机制拆解为可以被公共转录组初步定位、再被湿实验验证的证据链。双硫死亡的定义需要胱氨酸摄入、还原力消耗、二硫键压力和肌动蛋白骨架崩塌连续出现；公共转录组最多只能捕捉入口和部分终点。因此，本文使用“disulfidptosis-like transporter-actin cytoskeletal stress”这一表述。这个措辞看似保守，但它让论文可以在不越过证据边界的情况下提出清晰机制假说。

人类脑组织结果支持一个阶段性模型。慢性 CTE 中，SLC7A11/SLC3A2/WASF2/TLN1 轴具有较稳定 FDR 支持，提示反复或慢性损伤背景下，胱氨酸转运入口和骨架调控节点可能随病程推进而增强。急性 severe TBI 中，ACTB、MYH9、MYL6 等骨架终点方向性更突出，这可能反映机械应力、血管破坏、细胞比例改变和急性骨架重塑。两者并不矛盾：急性重型损伤更容易暴露骨架终点应激，慢性 CTE 更容易积累转运入口和调控轴改变。

小鼠 CCI 结果给出了最实际的实验设计信息。3DPI hippocampus 是最集中的验证窗口，且 WASF2/TLN1、SLC3A2/SLC7A11 与多个终点基因在同一窗口中呈正向。对后续实验来说，这比“所有时间点都检测一遍”更有价值。第一轮应集中在 3DPI cortex/hippocampus，以 qPCR/WB 确认入口和调控节点，以 phalloidin 或 F-actin 染色确认形态，以 NADPH/NADP+ 和 GSH/GSSG 确认还原前提，再以 CD31、NeuN、GFAP、IBA1 和 OLIG2 共定位判断细胞背景。如果这些读数不能在同一窗口中相互支持，就应将机制重新表述为更广义的骨架重塑或氧化还原应激。

损伤程度是本文新增的重要维度，但它不能被粗暴合并。GSE209552 的 severe TBI、GSE193407 的 CTE stage 和 GSE223245 的 mild/moderate/severe blood groups 分别代表急性重型脑组织、慢性病程和外周严重程度背景。把它们统一成单一 severity 分数会制造错误的统计连续性。本文更合适的做法是分层讨论：急性 severe TBI 说明在重型脑组织中有方向性骨架应激，CTE stage 说明慢性病程中入口-调控轴有 FDR 支持，外周血说明部分基因与严重程度相关但不能替代脑内证据。

细胞类型解释仍然是当前最大的证据缺口。bulk marker-proxy 指向 endothelial 和 neuronal 背景，但这既可能来自细胞比例变化，也可能来自血管成分、坏死区域、采样部位或细胞状态变化。真正能推进论文的是共定位或单细胞/单核层面的验证。理想情况下，后续应从可获得的 snRNA-seq 对象中提取 SLC3A2/SLC7A11、WASF2/TLN1 和骨架终点在神经元、星形胶质、小胶质、少突胶质、OPC 和内皮细胞中的表达；若暂时没有合适对象，则在 3DPI cortex/hippocampus 中做多重 IF/IHC 是更可行的替代路径。

本文还有若干限制。第一，各数据集平台、物种、组织来源、病程定义和协变量不同，不能进行简单合并。第二，GSE104687 同 donor 多脑区样本不独立，不能把区域样本当作主分析重复。第三，GSE209552 样本量小，且代表外科切除 severe TBI 脑组织，外推性有限。第四，GSE223245 来自外周血/PBMC，只能提供 severity-aware context。第五，integrated evidence score 是实验优先级工具，不是统计预测模型。第六，mRNA 证据不能替代蛋白二硫键化、非还原胶、F-actin 形态和细胞死亡读数。

总体而言，本研究把 TBI 后双硫死亡问题从一个容易过度表述的通路概念，改写成一个可测试的转运-骨架应激模型。当前公共数据最支持三个层面：慢性 CTE 中的 SLC7A11/SLC3A2/WASF2/TLN1 病程轴，急性 severe TBI 中的骨架终点方向性增强，以及 mouse CCI 中 3DPI hippocampus 的验证窗口。下一步的关键不是继续增加公共数据图，而是在同一时间、同一脑区、同一细胞背景中补齐蛋白、F-actin、氧化还原和共定位证据。

## 结论

现有公共转录组支持 TBI/CTE 后存在候选双硫死亡样转运-肌动蛋白骨架应激，而不是直接证明双硫死亡已经发生。最稳健的人脑证据来自慢性 CTE stage 中的 SLC7A11/SLC3A2/WASF2/TLN1 轴；急性 severe TBI 提供骨架终点方向性线索；mouse CCI 将验证窗口收敛到 3DPI hippocampus/cortex；外周血严重程度数据只能作为 severity-aware context。后续应围绕 3DPI cortex/hippocampus 开展 SLC3A2/SLC7A11、WASF2/TLN1、ACTB/MYH9/MYL6/FLNA、F-actin、NADPH/GSH 和细胞共定位实验，以判断该候选转录组信号能否升级为真正的机制证据。

## 数据可用性

本研究使用 GEO 公共数据 GSE104687、GSE209552、GSE193407、GSE319253、GSE163415 和 GSE223245。重分析表格、PNG/PDF 图件、Markdown 和 Word 稿件均保存于本项目 Phase3 双硫死亡聚焦论文设计文件夹。

## 伦理声明

本阶段仅使用公开去标识化数据和既有文献资料，不涉及新增人体样本。后续动物实验需在开展前获得所在单位动物伦理委员会批准。

## References

1. Liu X, Nie L, Zhang Y, et al. Actin cytoskeleton vulnerability to disulfide stress mediates disulfidptosis. Nature Cell Biology. 2023;25:404-414. https://doi.org/10.1038/s41556-023-01091-2  
2. Machesky LM. Deadly actin collapse by disulfidptosis. Nature Cell Biology. 2023.  
3. Zhao G, Zhao J, Lang J, Sun G. Nrf2 functions as a pyroptosis-related mediator in traumatic brain injury and is correlated with cytokines and disease severity. Frontiers in Neurology. 2024;15:1341342. https://doi.org/10.3389/fneur.2024.1341342  
4. Thomas I, Dickens AM, Posti JP, et al. Serum metabolome associated with severity of acute traumatic brain injury. Nature Communications. 2022;13:2545. https://doi.org/10.1038/s41467-022-30227-5  
5. Zhang M, Shan H, Chang P, et al. Hydrogen Sulfide Offers Neuroprotection on Traumatic Brain Injury in Parallel with Reduced Apoptosis and Autophagy in Mice. PLoS ONE. 2014;9:e87241. https://doi.org/10.1371/journal.pone.0087241  
6. Labadorf A, Agus F, Aytan N, et al. Inflammation and neuronal gene expression changes differ in early versus late chronic traumatic encephalopathy brain. BMC Medical Genomics. 2023;16:49. https://doi.org/10.1186/s12920-023-01471-5  
7. Garza R, Sharma Y, Atacho DAM, et al. Single-cell transcriptomics of human traumatic brain injury reveals activation of endogenous retroviruses in oligodendroglia. Cell Reports. 2023;42:113395. https://doi.org/10.1016/j.celrep.2023.113395  
8. Attilio PJ, Snapper DM, Rusnak M, et al. Transcriptomic Analysis of Mouse Brain After Traumatic Brain Injury Reveals That the Angiotensin Receptor Blocker Candesartan Acts Through Novel Pathways. Frontiers in Neuroscience. 2021;15:636259. https://doi.org/10.3389/fnins.2021.636259
"""


def en_manuscript(vals: dict[str, str], tables: dict[str, str]) -> str:
    return f"""# Spatiotemporal, Severity-Linked and Cell-Context Prioritization of a Disulfidptosis-Like Transporter-Actin Stress Axis After TBI: An Integrative Bioinformatic Study of Human TBI/CTE and Mouse CCI Transcriptomes

## Abstract

Background: Secondary injury after traumatic brain injury (TBI) varies by time, region, injury severity and cellular context. Disulfidptosis is driven by SLC7A11/SLC3A2-dependent cystine uptake, NADPH depletion, disulfide stress and actin-cytoskeletal collapse, but public transcriptomes cannot directly prove protein disulfidation or F-actin collapse. We therefore asked whether TBI/CTE transcriptomes support a candidate disulfidptosis-like transporter-actin cytoskeletal stress state. Methods: We fixed an eight-gene panel before analysis: SLC3A2, SLC7A11, WASF2, TLN1, ACTB, MYH9, MYL6 and FLNA. Public datasets GSE209552, GSE193407, GSE319253, GSE104687, GSE163415 and GSE223245 were interpreted as distinct evidence tiers: acute severe human brain tissue, chronic CTE course, external chronic CTE support, remote multi-region TBI, mouse CCI time-region support and peripheral severity context. FDR-supported and nominal findings were explicitly separated. Results: The strongest human FDR-supported evidence came from GSE193407 CTE stage trend: {vals['stage_details']}. Late CTE stage 3-4 versus stage 0 supported {vals['late_details']}. Acute severe TBI showed positive directionality across all eight genes, including {vals['acute_genes']}. Mouse CCI localized the most coherent validation window to {vals['mouse_top']}. Peripheral blood severity added context only: {vals['severity_module']}; the strongest single-gene trend was {vals['severity_gene']}. Bulk marker-proxy analysis nominated co-localization priorities rather than cell-type assignment, with the top clue being {vals['marker_top']}. Conclusions: Public transcriptomes support a candidate post-TBI/CTE disulfidptosis-like transporter-actin stress model, but not mechanistic proof. A matched 3DPI cortex/hippocampus validation design should test transporter proteins, cytoskeletal endpoints, F-actin morphology, NADPH/GSH redox state and cell-type co-localization.

Keywords: traumatic brain injury; chronic traumatic encephalopathy; disulfidptosis; SLC7A11; SLC3A2; actin cytoskeleton; injury time; injury severity; public transcriptomics

## Graphical Abstract

{fig(0, 'en')}

## Introduction

TBI is not a single molecular event. After the primary mechanical insult, injured brain tissue enters a dynamic secondary phase characterized by metabolic crisis, ionic imbalance, excitotoxicity, mitochondrial dysfunction, redox stress, blood-brain barrier disruption, vascular responses, glial activation and cytoskeletal remodeling. In chronic or repetitive injury, CTE adds long-term neuroinflammatory, degenerative and cell-composition changes. A useful mechanistic transcriptomic paper should therefore not ask only whether TBI differs from control, but when, where, at what injury severity and in which cellular context a candidate process is most plausible.

Disulfidptosis provides a mechanistic entry point into this problem. In the original model, high SLC7A11/SLC3A2-dependent cystine uptake under limited glucose availability consumes NADPH, increases disulfide stress and collapses the actin cytoskeleton. TBI brain tissue contains several relevant prerequisites, including energy stress, mitochondrial dysfunction, redox imbalance and cytoskeletal injury. However, public bulk transcriptomes cannot detect protein disulfidation, non-reducing migration, F-actin collapse or cell death morphology. For this reason, the present manuscript deliberately uses the term disulfidptosis-like transporter-actin cytoskeletal stress.

The hypothesis was constrained before analysis to eight genes. SLC3A2 and SLC7A11 define transporter entry. WASF2 and TLN1 define actin regulation, adhesion and tension transfer. ACTB, MYH9, MYL6 and FLNA define structural or mechanical cytoskeletal endpoints. This sacrifices pathway breadth but improves interpretability: if a disulfidptosis-adjacent stress state is present after TBI/CTE, transporter entry and cytoskeletal endpoints should show coordinated or stage-dependent changes in specific time windows, regions or disease-course strata.

## Materials and Methods

### Datasets and Evidence Tiers

Six GEO evidence layers were included. GSE209552 represented acute severe surgically sampled human TBI brain tissue. GSE193407 represented human BA9 CTE stage 0-4 and was used for chronic disease-course analysis. GSE319253 represented external superior frontal cortex CTE versus control support. GSE104687 represented remote multi-region postmortem TBI and was treated as exploratory because multiple regions from the same donors are not independent observations. GSE163415 represented mouse CCI across 3DPI/29DPI, regions and treatment strata. GSE223245 represented whole blood/PBMC mild, moderate and severe TBI and was used only as peripheral severity context.

{table_caption(1, 'Dataset layers and evidence boundaries', 'This table defines the biological role and interpretation boundary of each dataset before analysis. Acute severe TBI, chronic CTE stage, remote regions, mouse CCI and peripheral severity are not interchangeable measures and are therefore interpreted separately.')}

{tables['datasets']}

{table_note('Animal and peripheral datasets are supportive layers. They are not used as direct proof of human brain mechanism.')}

### Candidate Genes, Modules and Statistical Principles

The candidate panel was fixed as SLC3A2, SLC7A11, WASF2, TLN1, ACTB, MYH9, MYL6 and FLNA. Three module readouts were calculated: the full eight-gene module, the two-gene transporter-entry module and the six-gene actin-cytoskeletal module. Module scores were computed as the mean of sample-wise z-scored expression values. Gene-level comparisons used existing differential-expression outputs or Welch tests. Severity trends used Spearman correlations. Benjamini-Hochberg FDR was used for multiple testing. Only FDR<0.05 is described as FDR-supported; P<0.05 without FDR support is described as a nominal exploratory clue.

### Display-Item Integration

The v5 manuscript follows a figure-driven Results structure modeled on TBI bioinformatics-validation papers, acute TBI severity omics papers and the original disulfidptosis mechanism paper. Workflow figures define inference boundaries; primary multi-panel figures answer specific biological questions; tables provide numerical audit trails. Table 2 links each figure to a question, content layer and claim boundary.

{table_caption(2, 'Display-item-to-claim map used for the v5 rewrite', 'Each figure was assigned a primary biological question and an explicit claim boundary so that figure captions, Results paragraphs and Discussion claims refer to the same evidence layer.')}

{tables['display']}

{table_note('This is a manuscript-level audit table rather than a statistical test.')}

## Results

### A Tiered Design Converted Disulfidptosis Into a Testable Transporter-Actin Axis

Fig. 1 is the logical starting point of the manuscript. Fig. 1A assigns acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets to different evidence layers rather than merging them into a single TBI dataset. Fig. 1B makes the biological axes explicit: acute injury, remote region, CTE course, external CTE validation, mouse CCI, blood severity and planned validation. Fig. 1C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints, clarifying why SLC3A2/SLC7A11, WASF2/TLN1 and ACTB/MYH9/MYL6/FLNA should not be interpreted as a homogeneous block. Fig. 1D states the claim boundary used throughout the manuscript: FDR-supported human brain results can support association, nominal results are candidate clues, mouse CCI is supportive evidence, peripheral blood is severity context and bulk marker proxies guide prioritization only.

{fig(1, 'en')}

### Human Brain Data Supported a Chronic CTE-Stage Transporter-Regulatory Axis

Human brain datasets provided the highest-value but most heterogeneous evidence. Fig. 2A shows that the largest average effect occurred in acute severe TBI, whereas FDR-supported human evidence was concentrated in the GSE193407 CTE-stage rows. Fig. 2B explains the core chronic finding: the continuous CTE-stage analysis identified {vals['stage_details']}. Late CTE stage 3-4 versus stage 0 supported {vals['late_details']}. These genes form a transporter-regulatory axis rather than a broad cytoskeletal endpoint signature. Fig. 2C combines nominal recurrence and maximum absolute effect, showing why endpoint genes such as MYH9, ACTB and MYL6 remain useful follow-up targets but should be interpreted with less statistical confidence than SLC3A2, SLC7A11, WASF2 and TLN1.

{fig(2, 'en')}

{table_caption(3, 'Human comparison-level evidence audit', 'This table summarizes each human-brain comparison at the panel level and separates mean direction, positive gene count, nominal support and FDR support.')}

{tables['human']}

{table_note('GSE193407 provides the strongest FDR-supported human evidence; GSE209552 provides acute directional support with small-sample constraints.')}

### Acute Severe TBI Revealed Directional Cytoskeletal Endpoint Stress

GSE209552 added the acute severe TBI brain window. Fig. 3A shows higher transporter/stress and actin-target module scores in TBI samples, but also substantial sample heterogeneity. Fig. 3B shows positive logFC for all eight genes, with notable effects including {vals['acute_genes']}. Fig. 3C places the disulfidptosis-like panel in a ranked pathway context, where it exceeds several adjacent sulfur, copper and metabolic comparator pathways. Fig. 3D shows that module scores are positively correlated with endothelial proxies and negatively correlated with neuronal proxies. The correct inference is not cell-type assignment, but co-localization prioritization.

{fig(3, 'en')}

### Mouse CCI Localized the First Validation Window to 3DPI Hippocampus

GSE163415 answered the time-region question. Fig. 4A shows that 3DPI hippocampal units have the strongest combined transporter and actin module shifts. Fig. 4B resolves the pattern at the gene level and highlights WASF2/TLN1 together with transporter and endpoint genes in 3DPI hippocampal units. Fig. 4C shows that actin cytoskeleton/tension, amino-acid transport/cystine axis and disulfidptosis core/PDF genes rank above broader oxidative and mitochondrial categories. Fig. 4D converts the mouse evidence into gene-level validation priority. The highest-ranked window was {vals['mouse_top']}, and the hippocampal summary was {vals['mouse_hipp']}. These findings nominate 3DPI hippocampus/cortex for first-round wet-lab validation.

{fig(4, 'en')}

{table_caption(4, 'Top mouse CCI time-region-treatment validation windows', 'The table ranks GSE163415 analysis units by the coherence of the fixed eight-gene panel.')}

{tables['mouse']}

{table_note('Mouse CCI is supportive evidence and is used to choose validation windows, not to prove human brain mechanism.')}

### Evidence Audit Prevented Collapsing Acute, Chronic and Mouse Layers Into One Claim

Fig. 7 connects the primary results to the validation plan. Fig. 7A shows why acute severe TBI and chronic CTE must be interpreted differently: acute TBI has the largest mean effect but no FDR-supported genes, whereas CTE-stage comparisons contain the clearest FDR-supported human evidence. Fig. 7B independently confirms that 3DPI hippocampus ranks highest in mouse CCI. Fig. 7C shows treatment-stratified Drug-Veh module differences, with the strongest effect being {vals['treatment_top']}; this is interpreted as stage-stratified modulation, not treatment efficacy. Fig. 7D lists the evidence gaps that must be upgraded before a true disulfidptosis claim can be made.

{fig(7, 'en')}

### Peripheral Blood Added Severity Context but Not Brain-Mechanism Evidence

GSE223245 was used to add injury-severity context. Fig. 5A shows partial but incomplete separation of control and severity groups in the eight-gene PCA space. Fig. 5B shows module-score variation across control, mild, moderate and severe groups without a perfectly monotonic pattern. Fig. 5C identifies gene-level severity trends, with the strongest result being {vals['severity_gene']}. Fig. 5D summarizes module-level associations, with {vals['severity_module']}. These results are useful because they remind future studies to include injury severity, but they remain peripheral blood/PBMC signals and cannot replace matched brain tissue evidence.

{fig(5, 'en')}

{table_caption(5, 'Peripheral blood 8-gene severity trends in GSE223245', 'Gene-level Spearman correlations are reported across severity 0-3 and within TBI-only severity 1-3.')}

{tables['severity']}

{table_note('The strongest blood trends are peripheral severity clues and should not be interpreted as brain-region mechanism.')}

### Marker Proxies Prioritized Co-Localization Rather Than Assigning Cell Type

The strongest GSE209552 marker-proxy clue was {vals['marker_top']}. Fig. 3D and Fig. 8B show positive endothelial and negative neuronal associations for the eight-gene module. These results are hypothesis-generating because bulk correlations can arise from cell proportion, vascular content, lesion sampling or within-cell expression state. The experimental implication is concrete: SLC3A2/SLC7A11, WASF2/TLN1 and F-actin readouts should be co-stained with CD31 and NeuN, and also with GFAP, IBA1 and OLIG2, before any cell-type claim is made.

{table_caption(6, 'Bulk marker-proxy prioritization in acute severe TBI brain tissue', 'Module-cell proxy correlations are ranked by absolute Spearman correlation. They are labeled as proxy clues because bulk tissue cannot distinguish cell abundance from cell state.')}

{tables['marker']}

{table_note('FDR support is limited; the table is intended to prioritize co-localization experiments.')}

### Integrated Ranking Separated Transporter Entry, Actin Regulation and Endpoint Readouts

Fig. 6A combines human FDR support, human nominal support, mouse nominal support, acute logFC and blood severity association into an integrated dry-evidence matrix. Fig. 6B converts this into an experimental priority score, with {vals['priority_top']} ranking first. Fig. 6C translates the rank into a validation sequence: first test SLC3A2/SLC7A11/WASF2/TLN1 in 3DPI cortex/hippocampus, then add F-actin, MYH9/MYL6/FLNA, NADPH/GSH assays and cell co-localization. Fig. 6D inserts the marker-proxy results into that plan, making CD31 and NeuN co-staining necessary rather than optional.

{fig(6, 'en')}

{table_caption(7, 'Integrated 8-gene evidence priority', 'The score ranks experimental validation targets by combining human, mouse, acute and peripheral evidence layers.')}

{tables['priority']}

{table_note('The score is a prioritization tool, not a clinical classifier or causal model.')}

### Mechanistic Upgrading Requires Matched Transporter, Cytoskeletal, Redox and Localization Evidence

Fig. 8 defines the inferential ceiling of the manuscript. Fig. 8A places current evidence below the mechanistic-proof tier. Fig. 8B and Fig. 8D provide cellular-context and peripheral-severity clues, but neither proves disulfidptosis. Fig. 8C converts the transcriptomic pattern into a wet-lab time-course design, prioritizing 3DPI and 7D for transporter, actin/F-actin, redox and co-localization readouts. Table 8 specifies the missing evidence classes required to move from transcriptomic association to tissue validation and finally to mechanistic support.

{fig(8, 'en')}

{table_caption(8, 'Mechanistic upgrade matrix for the next experimental stage', 'Each row links a missing evidence class to a concrete experimental readout required for mechanistic upgrading.')}

{tables['validation']}

{table_note('Mechanistic proof requires matched time, region and cell context.')}

## Discussion

The main contribution of this study is not the nomination of a single TBI gene, but the conversion of a new cell-death mechanism into a testable transcriptomic and experimental framework. Disulfidptosis requires a chain linking cystine import, reducing-power depletion, disulfide stress and actin cytoskeletal collapse. Public transcriptomes can only capture the entry point and part of the endpoint. The manuscript therefore uses conservative wording while still preserving mechanistic direction.

The human brain results support a stage-dependent model. Chronic CTE contains the strongest FDR-supported transporter-regulatory axis, centered on SLC7A11/SLC3A2/WASF2/TLN1. Acute severe TBI shows larger directional effects, especially among cytoskeletal endpoint genes, but remains limited by sample size. These observations are complementary: acute severe injury may expose cytoskeletal stress, whereas chronic disease may accumulate transporter and regulatory-axis changes.

Mouse CCI provides the most actionable validation design. The strongest window is 3DPI hippocampus, which argues against spreading the first wet-lab round evenly across many regions and time points. A focused design should test 3DPI cortex/hippocampus for transporter proteins, regulatory nodes, endpoint proteins, F-actin morphology, NADPH/GSH redox state and cell co-localization. Later 7D or 29DPI samples can then determine persistence or transition.

Severity must be handled carefully. Severe acute TBI, CTE stage and mild/moderate/severe blood groups are related but not identical biological scales. Combining them into one severity score would be misleading. The layered structure used here keeps acute severity, chronic disease course and peripheral severity context separate while still allowing all three to inform validation design.

Cell context remains the largest evidence gap. Endothelial marker-proxy associations suggest vascular or barrier compartments should be prioritized, while neuronal negative associations require attention to neuronal loss or state shifts. However, bulk marker proxies cannot separate abundance from within-cell expression. Co-localization or annotated single-nucleus data are required before any cell-type conclusion can be made.

This study has limitations. Datasets differ in platform, species, tissue source, injury definition and covariates. GSE104687 contains non-independent multi-region samples from the same donors. GSE209552 is small and represents surgically sampled severe TBI tissue. GSE223245 is peripheral blood/PBMC and cannot substitute for brain data. The integrated score is an experimental prioritization device, not a predictive model. Finally, transcriptomics cannot replace protein disulfidation assays, F-actin morphology or cell-death phenotyping.

## Conclusion

Public transcriptomes support a candidate disulfidptosis-like transporter-actin cytoskeletal stress model after TBI/CTE, but they do not prove disulfidptosis. The strongest human evidence is the chronic CTE-stage SLC7A11/SLC3A2/WASF2/TLN1 axis; acute severe TBI adds directional cytoskeletal endpoint support; mouse CCI prioritizes 3DPI hippocampus/cortex; peripheral blood adds severity context only. The next experimental step should test transporter proteins, cytoskeletal endpoints, F-actin, redox state and cell-type co-localization in matched 3DPI cortex/hippocampus tissue.

## Data Availability

This study used public GEO datasets GSE104687, GSE209552, GSE193407, GSE319253, GSE163415 and GSE223245. Reanalysis tables, figures, Markdown manuscripts and Word files are stored in the Phase3 disulfidptosis-focused manuscript folder.

## Ethics Statement

This stage used only public de-identified datasets and literature materials. Future animal experiments require institutional animal ethics approval before initiation.

## References

1. Liu X, Nie L, Zhang Y, et al. Actin cytoskeleton vulnerability to disulfide stress mediates disulfidptosis. Nature Cell Biology. 2023;25:404-414. https://doi.org/10.1038/s41556-023-01091-2  
2. Machesky LM. Deadly actin collapse by disulfidptosis. Nature Cell Biology. 2023.  
3. Zhao G, Zhao J, Lang J, Sun G. Nrf2 functions as a pyroptosis-related mediator in traumatic brain injury and is correlated with cytokines and disease severity. Frontiers in Neurology. 2024;15:1341342. https://doi.org/10.3389/fneur.2024.1341342  
4. Thomas I, Dickens AM, Posti JP, et al. Serum metabolome associated with severity of acute traumatic brain injury. Nature Communications. 2022;13:2545. https://doi.org/10.1038/s41467-022-30227-5  
5. Zhang M, Shan H, Chang P, et al. Hydrogen Sulfide Offers Neuroprotection on Traumatic Brain Injury in Parallel with Reduced Apoptosis and Autophagy in Mice. PLoS ONE. 2014;9:e87241. https://doi.org/10.1371/journal.pone.0087241  
6. Labadorf A, Agus F, Aytan N, et al. Inflammation and neuronal gene expression changes differ in early versus late chronic traumatic encephalopathy brain. BMC Medical Genomics. 2023;16:49. https://doi.org/10.1186/s12920-023-01471-5  
7. Garza R, Sharma Y, Atacho DAM, et al. Single-cell transcriptomics of human traumatic brain injury reveals activation of endogenous retroviruses in oligodendroglia. Cell Reports. 2023;42:113395. https://doi.org/10.1016/j.celrep.2023.113395  
8. Attilio PJ, Snapper DM, Rusnak M, et al. Transcriptomic Analysis of Mouse Brain After Traumatic Brain Injury Reveals That the Angiotensin Receptor Blocker Candesartan Acts Through Novel Pathways. Frontiers in Neuroscience. 2021;15:636259. https://doi.org/10.3389/fnins.2021.636259
"""


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink_black(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    t = OxmlElement("w:t")
    t.text = text
    run.append(rpr)
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_markdown_runs(paragraph, text: str) -> None:
    token = re.compile(r"(\[Fig\. (\d+)\]\(#fig-(\d+)\)|\*\*(.*?)\*\*)")
    pos = 0
    for m in token.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(2) and m.group(3):
            add_hyperlink_black(paragraph, f"Fig. {m.group(2)}", f"fig_{m.group(3)}")
        else:
            run = paragraph.add_run(m.group(4))
            run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def link_figure_refs(md: str) -> str:
    linked_lines = []
    pattern = re.compile(r"(?<!\[)Fig\. ([0-8])(?=([A-D](?![A-Za-z])|[,\s.;:)]))")
    for line in md.splitlines():
        if line.startswith("**Fig. ") or line.startswith("![") or line.startswith('<a id="fig-') or line.startswith("| "):
            linked_lines.append(line)
        else:
            linked_lines.append(pattern.sub(lambda m: f"[Fig. {m.group(1)}](#fig-{m.group(1)})", line))
    return "\n".join(linked_lines)


def add_text_paragraph(doc: Document, text: str, *, style: str | None = None, align=None, font_size: float | None = None) -> object:
    p = doc.add_paragraph(style=style)
    add_markdown_runs(p, text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align is None else align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if font_size is not None:
        for run in p.runs:
            run.font.size = Pt(font_size)
    return p


def set_docx_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.color.rgb = RGBColor(0, 0, 0)
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 15), ("Heading 1", 13), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(7.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if row_idx == 0:
                        run.bold = True


def markdown_to_docx_v5(md: str, out_path: Path) -> None:
    doc = Document()
    set_docx_style(doc)
    bookmark_next: str | None = None
    bookmark_id = 1
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        anchor_match = re.match(r'<a id="fig-(\d+)"></a>', line)
        if anchor_match:
            bookmark_next = f"fig_{anchor_match.group(1)}"
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_markdown_runs(p, line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("!["):
            img_match = re.search(r"\]\(([^)]+)\)", line)
            if img_match:
                img_rel = img_match.group(1).replace("../figures/", "")
                img_path = FIGDIR / img_rel
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6.55))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([p.strip() for p in lines[i].strip().strip("|").split("|")])
                i += 1
            header = rows[0]
            data = rows[2:]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, h in enumerate(header):
                table.rows[0].cells[j].text = h
            for row in data:
                cells = table.add_row().cells
                for j, value in enumerate(row[: len(header)]):
                    cells[j].text = value
            format_table(table)
            i -= 1
        elif line.startswith("**Fig. "):
            p = add_text_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=9.0)
            if bookmark_next:
                add_bookmark(p, bookmark_next, bookmark_id)
                bookmark_id += 1
                bookmark_next = None
        elif line.startswith("**Table "):
            add_text_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=9.0)
        elif line.startswith("Note. "):
            add_text_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=8.5)
        else:
            add_text_paragraph(doc, line)
        i += 1
    doc.save(out_path)


def write_reference_notes(vals: dict[str, str]) -> None:
    refs = [
        ("Zhao 2024 TBI bioinformatics-validation template", "09_Zhao2024_TBI_bioinfo_clinical_validation_template.pdf", "used for the IMRAD order, workflow-first layout, DEG-to-validation progression and complete figure legends"),
        ("Thomas 2022 acute TBI severity omics template", "10_Thomas2022_acute_TBI_metabolomics_template.pdf", "used for cohort/platform-first Results, severity-stratified interpretation and dense panel-by-panel prose"),
        ("Liu 2023 original disulfidptosis mechanism", "01_Liu2023_original_disulfidptosis.pdf", "used for the transporter-NADPH-disulfide stress-actin endpoint chain and cautious mechanism criteria"),
        ("Machesky 2023 disulfidptosis commentary", "05_Machesky2023_NCB_deadly_actin_collapse_disulfidptosis.pdf", "used for conceptual emphasis on actin collapse as the endpoint rather than transporter expression alone"),
        ("Zhang 2014 mouse TBI model", "04_Zhang2013_H2S_neuroprotection_TBI.pdf", "used for mouse TBI time-course framing and cortex/hippocampus validation style"),
    ]
    lines = ["# v5 reference-format and manuscript-integration audit", ""]
    lines.append("The v5 rewrite no longer treats figures as separated report inserts. Each figure was assigned a primary biological question, each multi-panel caption was expanded with panel-by-panel descriptions, and each Results subsection now interprets the relevant panels before moving to cautious biological inference.")
    lines.append("")
    for title, fname, note in refs:
        pdf = WORKDIR / "references" / fname
        lines.append(f"- [{title}]({pdf.resolve()}): {note}.")
    lines.extend(
        [
            "",
            "Main v5 formatting changes:",
            "- Word styles were reset to black manuscript-style typography, with hierarchy expressed by font size and bold weight rather than colored headings.",
            "- Figure captions now sit below the figures in the generated Word files.",
            "- Fig. internal hyperlinks remain clickable but use black underlined text.",
            "- Table captions and notes were expanded so tables can be read independently.",
            "",
            f"Key numerical anchors retained: {vals['stage_details']}; {vals['mouse_top']}; {vals['severity_module']}; {vals['marker_top']}.",
        ]
    )
    (REPORTDIR / "reference_format_papers_for_review_v5_20260605.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    v4.configure_style()
    data = v4.load_inputs()
    ext = v4.build_extended_tables(data)
    vals = v4.build_key_values(data, ext)
    tables = build_table_strings(data, ext)
    zh = link_figure_refs(zh_manuscript(vals, tables))
    en = link_figure_refs(en_manuscript(vals, tables))

    zh_md = REPORTDIR / "TBI_disulfidptosis_integrated_manuscript_v5_ZH_20260605.md"
    en_md = REPORTDIR / "TBI_disulfidptosis_integrated_manuscript_v5_EN_20260605.md"
    zh_docx = REPORTDIR / "TBI_disulfidptosis_integrated_manuscript_v5_ZH_20260605.docx"
    en_docx = REPORTDIR / "TBI_disulfidptosis_integrated_manuscript_v5_EN_20260605.docx"
    zh_md.write_text(zh, encoding="utf-8")
    en_md.write_text(en, encoding="utf-8")
    markdown_to_docx_v5(zh, zh_docx)
    markdown_to_docx_v5(en, en_docx)
    write_reference_notes(vals)

    artifacts = [
        ("report", zh_md, "Chinese integrated v5 Markdown manuscript"),
        ("report", zh_docx, "Chinese integrated v5 Word manuscript with black manuscript-style typography"),
        ("report", en_md, "English integrated v5 Markdown manuscript"),
        ("report", en_docx, "English integrated v5 Word manuscript with black manuscript-style typography"),
        ("report", REPORTDIR / "reference_format_papers_for_review_v5_20260605.md", "v5 reference-format and integration audit"),
        ("table", TABLEDIR / "v5_display_claim_map_20260605.csv", "v5 figure-to-claim integration audit table"),
        ("script", Path(__file__), "v5 manuscript generation script"),
    ]
    manifest = pd.DataFrame(
        [{"type": kind, "path": str(path), "description": desc, "exists": path.exists()} for kind, path, desc in artifacts]
    )
    manifest_path = REPORTDIR / "integrated_v5_artifact_manifest_20260605.csv"
    manifest.to_csv(manifest_path, index=False, encoding="utf-8-sig")
    print(f"Wrote {zh_docx}")
    print(f"Wrote {en_docx}")
    print(f"Wrote {manifest_path}")


if __name__ == "__main__":
    main()
