from __future__ import annotations

import re
import shutil
import textwrap
from pathlib import Path

import anndata as ad
import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import scanpy as sc
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path.cwd()
PROJECT_DIR = next(p.parent for p in ROOT.rglob("TBI_disulfidptosis_submission_package_20260611") if p.is_dir())
TABLE_DIR = PROJECT_DIR / "tables"
H5AD_PATH = PROJECT_DIR / "raw_data_v3_20260604" / "GSE209552_true_snRNA_processed_20260613.h5ad"
OUT_ROOT = next(p for p in ROOT.rglob("*Neurochemical_Research*") if p.is_dir() and (p / "06_Final_submission_ready").exists())
FINAL_OUT = OUT_ROOT / "06_Final_submission_ready"
POLISHED_DIR = OUT_ROOT / "08_polished_figures_and_manuscript_20260623"
POLISHED_PNG = POLISHED_DIR / "png"
POLISHED_TIFF = POLISHED_DIR / "tiff_600dpi"
FINAL_TIFF = FINAL_OUT / "06_Figures_600dpi_TIFF"
DOC_OUT = OUT_ROOT / "01_manuscript"
CHECK_OUT = FINAL_OUT / "07_Checks"

GENES = ["SLC3A2", "SLC7A11", "WASF2", "TLN1", "ACTB", "MYH9", "MYL6", "FLNA"]
TRANSPORTER = ["SLC3A2", "SLC7A11"]
ACTIN = ["WASF2", "TLN1", "ACTB", "MYH9", "MYL6", "FLNA"]
OKABE = {
    "blue": "#0072B2",
    "sky": "#56B4E9",
    "green": "#009E73",
    "orange": "#E69F00",
    "red": "#D55E00",
    "purple": "#CC79A7",
    "gray": "#8A95A5",
    "black": "#222222",
}
CELL_COLORS = {
    "Neuron": "#62C6A9",
    "Astrocyte": "#A78BFA",
    "Microglia": "#F59E0B",
    "Oligodendrocyte": "#A3D977",
    "OPC": "#F5D76E",
    "Endothelial": "#60A5FA",
    "Immune": "#F4A3C4",
    "Pericyte": "#D6B3E6",
    "Ependymal": "#A7D7C5",
    "Unknown": "#D1D5DB",
}
EXPR_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "expr_blue_lilac",
    ["#F8FAFC", "#DDD6FE", "#93C5FD", "#2563EB"],
)
SCORE_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "score_magenta_blue",
    ["#FFF7FB", "#FBCFE8", "#C4B5FD", "#60A5FA", "#1D4ED8"],
)


def configure_style() -> None:
    mpl.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial", "DejaVu Sans", "Microsoft YaHei"],
            "font.size": 8.5,
            "axes.titlesize": 10,
            "axes.labelsize": 8.8,
            "xtick.labelsize": 8,
            "ytick.labelsize": 8,
            "legend.fontsize": 7.4,
            "axes.unicode_minus": False,
            "figure.dpi": 160,
            "savefig.dpi": 320,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
        }
    )
    sns.set_theme(style="white", font="Arial")


def ensure_dirs() -> None:
    for path in [POLISHED_PNG, POLISHED_TIFF, FINAL_TIFF, DOC_OUT, CHECK_OUT]:
        path.mkdir(parents=True, exist_ok=True)


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLE_DIR / name)


def add_panel(ax: mpl.axes.Axes, label: str, x: float = -0.12, y: float = 1.08) -> None:
    ax.text(x, y, label, transform=ax.transAxes, fontsize=12, fontweight="bold", va="top", ha="left")


def sig_from_p(p: float | int | None, fdr: float | int | None = None, nominal_symbol: str = "†") -> str:
    if fdr is not None and not pd.isna(fdr):
        if fdr < 0.001:
            return "***"
        if fdr < 0.01:
            return "**"
        if fdr < 0.05:
            return "*"
    if p is not None and not pd.isna(p):
        if p < 0.05:
            return nominal_symbol
        if p < 0.10:
            return nominal_symbol
    return ""


def wrapped(labels: list[str], width: int = 22) -> list[str]:
    return ["\n".join(textwrap.wrap(str(x), width=width)) for x in labels]


def export_figure(fig: plt.Figure, figure_name: str) -> Path:
    png_path = POLISHED_PNG / f"{figure_name}.png"
    tiff_path = POLISHED_TIFF / f"{figure_name}.tif"
    final_tiff_path = FINAL_TIFF / f"{figure_name}.tif"
    fig.savefig(png_path, bbox_inches="tight", dpi=320)
    plt.close(fig)
    im = Image.open(png_path).convert("RGB")
    im.save(tiff_path, dpi=(600, 600), compression="tiff_lzw")
    shutil.copy2(tiff_path, final_tiff_path)
    return png_path


def make_fig1() -> Path:
    fig = plt.figure(figsize=(11.2, 7.2))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.0, 1.2], width_ratios=[1.18, 1.0], hspace=0.48, wspace=0.36)
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    ax_a.axis("off")
    boxes = [
        ("Human acute\nsevere TBI\nGSE209552", 0.06, 0.62, OKABE["red"]),
        ("Human chronic\nCTE stage\nGSE193407/GSE319253", 0.54, 0.62, OKABE["blue"]),
        ("Mouse CCI\ntime-region\nGSE163415", 0.06, 0.18, OKABE["green"]),
        ("Peripheral\nseverity\nGSE223245", 0.54, 0.18, OKABE["orange"]),
    ]
    for text, x, y, color in boxes:
        rect = mpl.patches.FancyBboxPatch(
            (x, y), 0.36, 0.24, boxstyle="round,pad=0.02,rounding_size=0.018",
            fc=color, ec="#374151", alpha=0.13, lw=0.8,
        )
        ax_a.add_patch(rect)
        ax_a.text(x + 0.18, y + 0.12, text, ha="center", va="center", fontsize=8.2, linespacing=1.12)
    ax_a.annotate("", xy=(0.54, 0.74), xytext=(0.42, 0.74), arrowprops=dict(arrowstyle="->", lw=1.0, color="#4B5563"))
    ax_a.annotate("", xy=(0.54, 0.30), xytext=(0.42, 0.30), arrowprops=dict(arrowstyle="->", lw=1.0, color="#4B5563"))
    ax_a.text(
        0.50, 0.50,
        "Fixed 8-gene panel\nSLC3A2/SLC7A11 + actin endpoints",
        ha="center", va="center", fontsize=9.2, fontweight="bold",
    )
    ax_a.set_xlim(0, 1)
    ax_a.set_ylim(0, 1)
    ax_a.set_title("Evidence-layer design", pad=8)
    add_panel(ax_a, "A", -0.08, 1.05)

    layers = ["acute", "remote region", "CTE course", "CTE external", "mouse CCI", "blood severity", "snRNA localization"]
    ax_b.barh(np.arange(len(layers)), np.ones(len(layers)), color=[
        OKABE["red"], OKABE["gray"], OKABE["blue"], OKABE["sky"], OKABE["green"], OKABE["orange"], OKABE["purple"],
    ])
    ax_b.set_yticks(np.arange(len(layers)))
    ax_b.set_yticklabels(layers)
    ax_b.set_xlim(0, 1.05)
    ax_b.set_xlabel("Evidence layer present")
    ax_b.set_title("Time, region, severity and cell axes", pad=8)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.05)

    ax_c.set_title("Pre-fixed transporter-actin panel", pad=10)
    ax_c.set_xlim(0.65, 4.05)
    ax_c.set_ylim(-0.8, 7.8)
    ax_c.set_yticks([])
    ax_c.set_xticks([1, 2, 3])
    ax_c.set_xticklabels(["cystine import", "actin control", "cytoskeletal endpoint"])
    gene_rows = [
        ("SLC3A2", "Transporter", 1, 0.0, OKABE["red"]),
        ("SLC7A11", "Transporter", 1, 1.0, OKABE["red"]),
        ("WASF2", "WRC/remodeling", 2, 2.9, OKABE["blue"]),
        ("TLN1", "Adhesion/tension", 2, 4.0, OKABE["blue"]),
        ("ACTB", "Actin scaffold", 3, 5.0, OKABE["green"]),
        ("MYH9", "Myosin/tension", 3, 6.0, OKABE["green"]),
        ("MYL6", "Myosin light chain", 3, 7.0, OKABE["green"]),
        ("FLNA", "Filamin scaffold", 3, 7.55, OKABE["green"]),
    ]
    for gene, role, x, y, color in gene_rows:
        ax_c.scatter(x, y, s=140, c=color, edgecolor="black", lw=0.8, zorder=3)
        ax_c.text(x + 0.13, y, f"{gene}  {role}", va="center", ha="left", fontsize=7.8)
    sns.despine(ax=ax_c, left=True)
    add_panel(ax_c, "C", -0.08, 1.05)

    ax_d.axis("off")
    guard = pd.DataFrame(
        {
            "Claim": ["Human brain FDR", "Nominal signal", "Mouse CCI", "Peripheral blood", "Bulk marker proxy"],
            "Allowed wording": ["association", "candidate clue", "supportive evidence", "severity context", "prioritization"],
            "Avoid": ["causality", "biomarker claim", "human proof", "brain mechanism", "cell-type proof"],
        }
    )
    tbl = ax_d.table(cellText=guard.values, colLabels=guard.columns, cellLoc="left", colLoc="left", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(7.4)
    tbl.scale(1.05, 1.55)
    for (r, c), cell in tbl.get_celld().items():
        cell.set_edgecolor("#D1D5DB")
        if r == 0:
            cell.set_facecolor("#F3F4F6")
            cell.set_text_props(weight="bold")
    ax_d.set_title("Interpretation guardrails", pad=8)
    add_panel(ax_d, "D", -0.10, 1.05)
    return export_figure(fig, "Fig1")


def make_fig2() -> Path:
    focus = read_csv("v3_human_brain_8gene_focused_results_20260604.csv")
    summary = read_csv("v3_human_brain_8gene_evidence_summary_20260604.csv")
    fig = plt.figure(figsize=(11.8, 8.1))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.23, 1.0], width_ratios=[1.35, 1.0], hspace=0.55, wspace=0.38)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0])
    ax_c = fig.add_subplot(gs[1, 1])

    order = [
        "GSE209552_human_acute_brain_severe_TBI_vs_Control",
        "GSE193407_human_prefrontal_BA9_CTE_stage_trend",
        "GSE193407_human_prefrontal_BA9_late_CTE_stage3_4_vs_stage0",
        "GSE319253_human_superior_frontal_cortex_CTE_vs_Control",
        "GSE104687_human_brain_FWM_TBI_vs_Control",
        "GSE104687_human_brain_HIP_TBI_vs_Control",
        "GSE104687_human_brain_PCx_TBI_vs_Control",
        "GSE104687_human_brain_TCx_TBI_vs_Control",
    ]
    labels = ["Acute severe TBI", "CTE stage trend", "Late CTE vs stage0", "External CTE", "Remote FWM", "Remote HIP", "Remote PCx", "Remote TCx"]
    heat = focus[focus["comparison"].isin(order)].pivot_table(index="comparison", columns="gene_symbol", values="effect", aggfunc="mean").reindex(order)[GENES]
    sns.heatmap(heat, cmap="RdBu_r", center=0, linewidths=0.35, linecolor="white", ax=ax_a, cbar_kws={"label": "logFC or stage r"})
    ax_a.set_yticklabels(labels, rotation=0)
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    ax_a.set_title("Human brain 8-gene effect map across acute, chronic and regional comparisons", pad=10)
    for y, comp in enumerate(order):
        sub = focus[focus["comparison"].eq(comp)].set_index("gene_symbol")
        for x, gene in enumerate(GENES):
            if gene not in sub.index:
                continue
            fdr = sub.loc[gene, "FDR"]
            p = sub.loc[gene, "p_value"]
            mark = "*" if pd.notna(fdr) and fdr < 0.05 else ("†" if pd.notna(p) and p < 0.05 else "")
            if mark:
                ax_a.text(x + 0.5, y + 0.5, mark, ha="center", va="center", fontsize=8.5, color="black")
    add_panel(ax_a, "A", -0.07, 1.05)

    trend = focus[focus["comparison"].eq("GSE193407_human_prefrontal_BA9_CTE_stage_trend")].set_index("gene_symbol").reindex(GENES)
    vals = pd.to_numeric(trend["stage_correlation"], errors="coerce")
    colors = [OKABE["blue"] if f else "#BFC7D5" for f in trend["significant_FDR_0_05"].fillna(False)]
    ax_b.bar(np.arange(len(GENES)), vals, color=colors, edgecolor="#374151", lw=0.3)
    ax_b.axhline(0, color="black", lw=0.8)
    ax_b.set_ylim(0, max(0.45, float(vals.max()) + 0.12))
    for i, gene in enumerate(GENES):
        mark = sig_from_p(trend.loc[gene, "p_value"], trend.loc[gene, "FDR"])
        if mark:
            ax_b.text(i, vals.loc[gene] + 0.035, mark, ha="center", va="bottom", fontsize=9)
    ax_b.set_xticks(np.arange(len(GENES)))
    ax_b.set_xticklabels(GENES, rotation=45, ha="right")
    ax_b.set_ylabel("Stage correlation (r)")
    ax_b.set_title("GSE193407 continuous CTE stage", pad=12)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    summ = summary.set_index("gene_symbol").reindex(GENES).reset_index()
    sca = ax_c.scatter(
        summ["human_nominal_supported"],
        summ["gene_symbol"],
        s=(summ["human_fdr_supported"] + 1) * 95,
        c=summ["max_abs_effect"],
        cmap="viridis",
        edgecolor="black",
        lw=0.8,
    )
    ax_c.set_xlabel("Nominally supported human comparisons")
    ax_c.set_ylabel("")
    ax_c.set_title("Cross-human evidence recurrence", pad=10)
    cbar = fig.colorbar(sca, ax=ax_c, shrink=0.82)
    cbar.set_label("Max |effect|")
    sns.despine(ax=ax_c)
    add_panel(ax_c, "C", -0.15, 1.08)
    return export_figure(fig, "Fig2")


def make_fig3() -> Path:
    acute = read_csv("v3_GSE209552_8gene_acute_severe_results_20260604.csv").set_index("gene_symbol").reindex(GENES).reset_index()
    scores = read_csv("v3_GSE209552_bulk_marker_proxy_scores_20260604.csv")
    corr = read_csv("v3_GSE209552_bulk_marker_proxy_correlations_20260604.csv")
    fig = plt.figure(figsize=(11.8, 8.4))
    gs = fig.add_gridspec(2, 2, hspace=0.55, wspace=0.38)
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    long = scores.melt(
        id_vars=["sample_id", "condition"],
        value_vars=["score_disulfidptosis_8gene", "score_transporter_2gene", "score_actin_6gene"],
        var_name="module",
        value_name="score",
    )
    long["module"] = long["module"].map(
        {
            "score_disulfidptosis_8gene": "8-gene",
            "score_transporter_2gene": "Transporter",
            "score_actin_6gene": "Actin",
        }
    )
    sns.boxplot(data=long, x="module", y="score", hue="condition", ax=ax_a, fliersize=0, palette={"Control": "#BFC7D5", "TBI": OKABE["red"]})
    sns.stripplot(data=long, x="module", y="score", hue="condition", ax=ax_a, dodge=True, color="black", size=3.1, alpha=0.6)
    if ax_a.legend_:
        ax_a.legend_.remove()
    ax_a.set_ylabel("Z-mean module score")
    ax_a.set_xlabel("")
    ax_a.set_title("Acute severe TBI module scores", pad=10)
    add_panel(ax_a, "A", -0.12, 1.08)

    vals = acute["GSE209552_logFC_TBI_vs_Control"]
    colors = [OKABE["red"] if p < 0.05 else "#BFC7D5" for p in acute["GSE209552_p_value_ttest"]]
    ax_b.bar(np.arange(len(GENES)), vals, color=colors, edgecolor="#374151", lw=0.3)
    ax_b.axhline(0, color="black", lw=0.8)
    ax_b.set_ylim(0, max(vals) + 0.32)
    for i, row in acute.iterrows():
        mark = sig_from_p(row["GSE209552_p_value_ttest"], None, nominal_symbol="*")
        if row["GSE209552_p_value_ttest"] >= 0.05 and row["GSE209552_p_value_ttest"] < 0.10:
            mark = "†"
        if mark:
            ax_b.text(i, row["GSE209552_logFC_TBI_vs_Control"] + 0.10, mark, ha="center", va="bottom", fontsize=9)
    ax_b.set_xticks(np.arange(len(GENES)))
    ax_b.set_xticklabels(GENES, rotation=45, ha="right")
    ax_b.set_ylabel("log2CPM logFC")
    ax_b.set_title("8-gene acute severe TBI direction", pad=12)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    enrich = pd.DataFrame(
        {
            "gene_set": ["H2S enzymes", "Methionine metabolism", "Copper/cuproptosis", "Disulfidptosis"],
            "ES": [-0.40, -0.25, 0.30, 0.76],
            "mark": ["", "", "", "***"],
        }
    )
    bar_colors = [OKABE["blue"] if x == "Disulfidptosis" else "#BFC7D5" for x in enrich["gene_set"]]
    ax_c.barh(enrich["gene_set"], enrich["ES"], color=bar_colors, edgecolor="#374151", lw=0.4)
    ax_c.axvline(0, color="black", lw=0.8)
    ax_c.set_xlim(-0.48, 0.90)
    for i, row in enrich.iterrows():
        if row["mark"]:
            ax_c.text(row["ES"] + 0.05, i, row["mark"], va="center", ha="left", fontsize=9)
    ax_c.set_xlabel("Enrichment score")
    ax_c.set_title("Ranked gene-set enrichment summary", pad=10)
    sns.despine(ax=ax_c)
    add_panel(ax_c, "C", -0.12, 1.08)

    heat = corr.pivot_table(index="module", columns="celltype_proxy", values="spearman_r").reindex(
        ["score_disulfidptosis_8gene", "score_transporter_2gene", "score_actin_6gene"]
    )
    sns.heatmap(heat, cmap="RdBu_r", center=0, vmin=-1, vmax=1, ax=ax_d, linewidths=0.35, cbar_kws={"label": "Spearman r"})
    ax_d.set_yticklabels(["8-gene", "Transporter", "Actin"], rotation=0)
    ax_d.set_xlabel("")
    ax_d.set_ylabel("")
    ax_d.set_title("Bulk marker-proxy cell-context prioritization", pad=10)
    add_panel(ax_d, "D", -0.13, 1.08)
    return export_figure(fig, "Fig3")


def raw_gene_df(adata: ad.AnnData, genes: list[str]) -> pd.DataFrame:
    valid = [g for g in genes if g in adata.raw.var_names]
    raw = adata.raw[:, valid].X
    if hasattr(raw, "toarray"):
        raw = raw.toarray()
    return pd.DataFrame(raw, index=adata.obs_names, columns=valid)


def load_sn() -> tuple[ad.AnnData, pd.DataFrame, pd.DataFrame]:
    adata = ad.read_h5ad(H5AD_PATH)
    if not pd.api.types.is_categorical_dtype(adata.obs["cell_type"]):
        adata.obs["cell_type"] = adata.obs["cell_type"].astype("category")
    return (
        adata,
        read_csv("GSE209552_true_scRNA_8gene_celltype_localization_20260613.csv"),
        read_csv("GSE209552_true_scRNA_8gene_donor_level_deltas_20260613.csv"),
    )


def make_fig4() -> Path:
    adata, gene_summary, delta_df = load_sn()
    fig = plt.figure(figsize=(14.2, 10.3))
    gs = fig.add_gridspec(3, 4, height_ratios=[0.78, 1.55, 1.28], width_ratios=[1.1, 1.15, 1.0, 1.32], hspace=0.64, wspace=0.56)
    ax_a = fig.add_subplot(gs[0, :])
    ax_b = fig.add_subplot(gs[1, 0:2])
    ax_c = fig.add_subplot(gs[1, 2])
    ax_d = fig.add_subplot(gs[1, 3])
    ax_e = fig.add_subplot(gs[2, 0:2])
    ax_f = fig.add_subplot(gs[2, 2:4])

    ax_a.axis("off")
    steps = [
        ("GEO RAW\n10x matrices", 0.08, OKABE["blue"]),
        ("QC retained\n19,627 nuclei", 0.30, OKABE["green"]),
        ("PCA/UMAP\nLeiden clusters", 0.52, OKABE["orange"]),
        ("Marker-based\ncell annotation", 0.74, OKABE["purple"]),
        ("8-gene\nlocalization", 0.93, OKABE["red"]),
    ]
    for text, x, color in steps:
        rect = mpl.patches.FancyBboxPatch((x - 0.075, 0.36), 0.15, 0.30, boxstyle="round,pad=0.025,rounding_size=0.02", fc=color, ec="#374151", alpha=0.14, lw=0.8)
        ax_a.add_patch(rect)
        ax_a.text(x, 0.51, text, ha="center", va="center", fontsize=8.4, fontweight="bold", linespacing=1.1)
    for left, right in zip([s[1] for s in steps[:-1]], [s[1] for s in steps[1:]]):
        ax_a.annotate("", xy=(right - 0.09, 0.51), xytext=(left + 0.09, 0.51), arrowprops=dict(arrowstyle="->", lw=1.1, color="#4B5563"))
    ax_a.text(0.01, 0.92, "GSE209552 true snRNA-seq localization workflow", ha="left", va="top", fontsize=10.5, fontweight="bold")
    ax_a.text(0.01, 0.18, "17 snRNA samples from 15 donors; donor-level TBI-control cell-type contrasts are exploratory because some control donors contribute repeated regions.", ha="left", va="center", fontsize=8.2)
    add_panel(ax_a, "A", -0.03, 1.03)

    obs = adata.obs.copy()
    coords = adata.obsm["X_umap"]
    for ct in obs["cell_type"].cat.categories:
        idx = obs["cell_type"].eq(ct).to_numpy()
        ax_b.scatter(coords[idx, 0], coords[idx, 1], s=5.0, color=CELL_COLORS.get(ct, "#D1D5DB"), label=ct, linewidths=0, alpha=0.82)
    ax_b.set_title("UMAP of annotated nuclei", pad=10)
    ax_b.set_xlabel("UMAP1")
    ax_b.set_ylabel("UMAP2")
    ax_b.legend(loc="upper center", bbox_to_anchor=(0.5, -0.13), ncol=4, frameon=False, columnspacing=0.8, handletextpad=0.3)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.10, 1.08)

    comp = obs.groupby(["condition", "cell_type"], observed=True).size().rename("n").reset_index().pivot(index="condition", columns="cell_type", values="n").fillna(0)
    comp = comp.div(comp.sum(axis=1), axis=0) * 100
    bottom = np.zeros(len(comp))
    for col in obs["cell_type"].cat.categories:
        if col not in comp.columns:
            continue
        vals = comp[col].to_numpy()
        ax_c.bar(comp.index, vals, bottom=bottom, color=CELL_COLORS.get(col, "#D1D5DB"), edgecolor="white", linewidth=0.4)
        bottom += vals
    ax_c.set_ylim(0, 100)
    ax_c.set_ylabel("Cell fraction (%)")
    ax_c.set_title("Cell-type composition", pad=10)
    sns.despine(ax=ax_c)
    add_panel(ax_c, "C", -0.22, 1.10)

    top_delta = delta_df[delta_df["gene"].isin(["SLC3A2", "SLC7A11", "WASF2", "TLN1"])].head(8).copy()
    top_delta["label"] = top_delta["gene"] + " | " + top_delta["cell_type"]
    ax_d.barh(top_delta["label"][::-1], top_delta["delta_tbi_minus_control"][::-1], color="#F29B8C", edgecolor="#92400E", linewidth=0.35)
    ax_d.axvline(0, color="black", lw=0.8)
    ax_d.set_xlabel("Donor-level mean log1p delta")
    ax_d.set_title("Core-gene TBI-control clues", pad=10)
    ax_d.tick_params(axis="y", labelsize=7.0)
    ax_d.grid(axis="x", color="#E5E7EB", lw=0.5)
    sns.despine(ax=ax_d)
    add_panel(ax_d, "D", -0.20, 1.10)

    dot = gene_summary.groupby(["cell_type", "gene"]).agg(mean_log1p_expr=("mean_log1p_expr", "mean"), pct_expressed=("pct_expressed", "mean")).reset_index()
    cell_order = [c for c in obs["cell_type"].cat.categories if c in dot["cell_type"].unique()]
    expr = dot.pivot(index="cell_type", columns="gene", values="mean_log1p_expr").reindex(cell_order)[GENES]
    pct = dot.pivot(index="cell_type", columns="gene", values="pct_expressed").reindex(cell_order)[GENES]
    vmax = float(np.nanmax(expr.values))
    for yi, ct in enumerate(cell_order):
        for xi, gene in enumerate(GENES):
            if pd.isna(expr.loc[ct, gene]):
                continue
            ax_e.scatter(xi, yi, s=max(12, pct.loc[ct, gene] * 2.0), c=[expr.loc[ct, gene]], cmap=EXPR_CMAP, vmin=0, vmax=vmax, edgecolor="#9CA3AF", linewidth=0.2)
    ax_e.set_xticks(range(len(GENES)))
    ax_e.set_xticklabels(GENES, rotation=45, ha="right")
    ax_e.set_yticks(range(len(cell_order)))
    ax_e.set_yticklabels(cell_order)
    ax_e.set_xlim(-0.5, len(GENES) - 0.5)
    ax_e.set_ylim(-0.5, len(cell_order) - 0.5)
    ax_e.invert_yaxis()
    ax_e.set_title("8-gene real localization dot plot", pad=10)
    add_panel(ax_e, "E", -0.10, 1.10)

    rdf = raw_gene_df(adata, GENES)
    score = rdf[[g for g in GENES if g in rdf.columns]].mean(axis=1).reindex(adata.obs_names)
    sca = ax_f.scatter(coords[:, 0], coords[:, 1], c=score.to_numpy(), s=5.0, cmap=SCORE_CMAP, linewidths=0, alpha=0.88)
    ax_f.set_title("Cell-level 8-gene localization score", pad=10)
    ax_f.set_xlabel("UMAP1")
    ax_f.set_ylabel("UMAP2")
    fig.colorbar(sca, ax=ax_f, shrink=0.78, label="Mean log1p score")
    sns.despine(ax=ax_f)
    add_panel(ax_f, "F", -0.10, 1.10)
    return export_figure(fig, "Fig4")


def make_fig5() -> Path:
    adata, _, _ = load_sn()
    genes = [g for g in GENES if g in adata.raw.var_names]
    rdf = raw_gene_df(adata, genes)
    coords = adata.obsm["X_umap"]
    fig, axes = plt.subplots(2, 4, figsize=(13.2, 6.7))
    axes = axes.ravel()
    for ax, gene in zip(axes, genes):
        vals = rdf[gene].reindex(adata.obs_names).to_numpy()
        sca = ax.scatter(coords[:, 0], coords[:, 1], c=vals, s=4.2, cmap=EXPR_CMAP, linewidths=0, alpha=0.92)
        ax.set_title(gene, pad=8, fontweight="bold")
        ax.set_xlabel("UMAP1")
        ax.set_ylabel("UMAP2")
        fig.colorbar(sca, ax=ax, fraction=0.046, pad=0.03)
        sns.despine(ax=ax)
    for ax in axes[len(genes):]:
        ax.axis("off")
    fig.suptitle("GSE209552 true snRNA feature maps for the 8 priority genes", y=1.02, fontsize=12, fontweight="bold")
    fig.tight_layout()
    return export_figure(fig, "Fig5")


def make_fig6() -> Path:
    focused = read_csv("v3_GSE163415_8gene_focused_DE_results_20260604.csv")
    mouse_units = read_csv("v4_mouse_units_20260605.csv")
    fig = plt.figure(figsize=(11.8, 8.2))
    gs = fig.add_gridspec(2, 2, hspace=0.58, wspace=0.45)
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    focused["unit"] = focused["time"].astype(str) + " " + focused["region"].astype(str) + " " + focused["treatment"].astype(str)
    module = focused.groupby(["unit", "module"], as_index=False)["logFC_TBI_vs_NoTBI"].mean()
    heat = module.pivot(index="unit", columns="module", values="logFC_TBI_vs_NoTBI")
    cols = [c for c in ["Core transporter/stress", "Actin cytoskeleton targets"] if c in heat.columns]
    heat = heat[cols]
    sns.heatmap(heat, cmap="RdBu_r", center=0, ax=ax_a, linewidths=0.35, cbar_kws={"label": "TBI-NoTBI delta"})
    ax_a.set_xticklabels(["Transporter", "Actin"], rotation=20, ha="right")
    ax_a.set_title("Mouse CCI module score shifts", pad=10)
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    add_panel(ax_a, "A", -0.12, 1.08)

    gene_heat = focused.pivot_table(index="unit", columns="gene_symbol", values="logFC_TBI_vs_NoTBI", aggfunc="mean").reindex(columns=GENES)
    sns.heatmap(gene_heat, cmap="RdBu_r", center=0, ax=ax_b, linewidths=0.35, cbar_kws={"label": "logFC"})
    ax_b.set_title("8-gene spatiotemporal logFC", pad=10)
    ax_b.set_xlabel("")
    ax_b.set_ylabel("")
    add_panel(ax_b, "B", -0.12, 1.08)

    pathways = pd.DataFrame(
        {
            "pathway": [
                "Actin cytoskeleton and tension",
                "Amino-acid transport and cystine axis",
                "Disulfidptosis core/PDF genes",
                "Inflammatory injury response",
                "Oxidative stress/glutathione handling",
                "Mitochondrial respiratory support",
            ],
            "neglog10_fdr": [2.55, 2.38, 2.37, 2.08, 1.32, 0.25],
        }
    )
    ax_c.barh(np.arange(len(pathways)), pathways["neglog10_fdr"], color=OKABE["green"], edgecolor="#374151", lw=0.35)
    ax_c.set_yticks(np.arange(len(pathways)))
    ax_c.set_yticklabels(wrapped(pathways["pathway"].tolist(), 32))
    ax_c.invert_yaxis()
    ax_c.set_xlabel("-log10(FDR)")
    ax_c.set_title("Pathway overlap in mouse CCI", pad=10)
    sns.despine(ax=ax_c)
    add_panel(ax_c, "C", -0.12, 1.08)

    nominal = focused.groupby("gene_symbol")["p_value"].apply(lambda s: int((pd.to_numeric(s, errors="coerce") < 0.05).sum())).reindex(GENES).fillna(0)
    ax_d.barh(GENES, nominal.values, color=OKABE["purple"], edgecolor="#374151", lw=0.35)
    ax_d.invert_yaxis()
    ax_d.set_xlabel("Nominally supported CCI units")
    ax_d.set_title("Mouse-derived 8-gene support", pad=10)
    sns.despine(ax=ax_d)
    add_panel(ax_d, "D", -0.12, 1.08)
    return export_figure(fig, "Fig6")


def make_fig7() -> Path:
    pca = read_csv("v3_GSE223245_8gene_PCA_20260604.csv")
    modules = read_csv("v3_GSE223245_module_scores_long_20260604.csv")
    gene_stats = read_csv("v3_GSE223245_8gene_severity_results_20260604.csv").set_index("gene_symbol").reindex(GENES).reset_index()
    module_stats = read_csv("v3_GSE223245_module_severity_results_20260604.csv")
    fig = plt.figure(figsize=(12.0, 8.3))
    gs = fig.add_gridspec(2, 2, hspace=0.62, wspace=0.52)
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    order = ["Control", "Mild", "Moderate", "Severe"]
    colors = {"Control": "#8D99AE", "Mild": OKABE["sky"], "Moderate": OKABE["orange"], "Severe": OKABE["red"]}
    for group in order:
        sub = pca[pca["group"].eq(group)]
        ax_a.scatter(sub["PC1"], sub["PC2"], label=group, s=58, color=colors[group], edgecolor="black", lw=0.4)
    ax_a.set_xlabel(f"PC1 ({pca['PC1_var'].iloc[0] * 100:.1f}%)")
    ax_a.set_ylabel(f"PC2 ({pca['PC2_var'].iloc[0] * 100:.1f}%)")
    ax_a.set_title("GSE223245 8-gene PCA (whole blood)", pad=10)
    ax_a.legend(frameon=False, ncol=2, loc="lower left")
    sns.despine(ax=ax_a)
    add_panel(ax_a, "A", -0.12, 1.08)

    melt = modules.melt(
        id_vars=["sample_id", "group", "severity_score"],
        value_vars=["disulfidptosis_8gene", "transporter_2gene", "actin_6gene"],
        var_name="module",
        value_name="score",
    )
    melt["module"] = melt["module"].map({"disulfidptosis_8gene": "8-gene", "transporter_2gene": "Transporter", "actin_6gene": "Actin"})
    sns.boxplot(data=melt, x="group", y="score", hue="module", order=order, ax=ax_b, fliersize=0, palette=[OKABE["blue"], OKABE["red"], OKABE["green"]])
    sns.stripplot(data=melt, x="group", y="score", hue="module", order=order, dodge=True, color="black", size=2.5, alpha=0.55, ax=ax_b)
    if ax_b.legend_:
        ax_b.legend_.remove()
    ax_b.set_xlabel("")
    ax_b.set_ylabel("Z-mean score")
    ax_b.set_title("Peripheral severity module score", pad=10)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    vals = gene_stats["spearman_r_severity_all_groups"]
    colors = [OKABE["orange"] if f < 0.05 else "#BFC7D5" for f in gene_stats["FDR_severity_all_groups"]]
    ax_c.bar(np.arange(len(GENES)), vals, color=colors, edgecolor="#374151", lw=0.35)
    ax_c.axhline(0, color="black", lw=0.8)
    ax_c.set_ylim(min(-0.95, float(vals.min()) - 0.12), max(0.62, float(vals.max()) + 0.15))
    for i, row in gene_stats.iterrows():
        mark = sig_from_p(row["p_value_severity_all_groups"], row["FDR_severity_all_groups"])
        if mark:
            y = row["spearman_r_severity_all_groups"]
            offset = 0.07 if y >= 0 else -0.10
            ax_c.text(i, y + offset, mark, ha="center", va="bottom" if y >= 0 else "top", fontsize=9)
    ax_c.set_xticks(np.arange(len(GENES)))
    ax_c.set_xticklabels(GENES, rotation=45, ha="right")
    ax_c.set_ylabel("Spearman r, severity 0-3")
    ax_c.set_title("8-gene peripheral severity trend", pad=12)
    sns.despine(ax=ax_c)
    add_panel(ax_c, "C", -0.12, 1.08)

    ms = module_stats.set_index("module").reindex(["disulfidptosis_8gene", "transporter_2gene", "actin_6gene"]).reset_index()
    ax_d.bar(["8-gene", "Transporter", "Actin"], ms["spearman_r_severity_all_groups"], color=[OKABE["blue"], OKABE["red"], OKABE["green"]], edgecolor="#374151", lw=0.35)
    ax_d.axhline(0, color="black", lw=0.8)
    ax_d.set_ylabel("Spearman r, severity 0-3")
    ax_d.set_title("Module-level severity association", pad=10)
    sns.despine(ax=ax_d)
    add_panel(ax_d, "D", -0.12, 1.08)
    return export_figure(fig, "Fig7")


def make_fig8() -> Path:
    priority = read_csv("v3_integrated_8gene_priority_20260604.csv")
    corr = read_csv("v4_marker_priority_20260605.csv")
    fig = plt.figure(figsize=(12.2, 8.6))
    gs = fig.add_gridspec(2, 2, hspace=0.62, wspace=0.55, width_ratios=[1.0, 1.1])
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    metrics = priority.set_index("gene_symbol")[["human_FDR_comparisons", "human_nominal_comparisons", "mouse_nominal_units", "GSE209552_acute_logFC", "GSE223245_severity_r_all"]].reindex(GENES)
    metrics.columns = ["Human FDR", "Human nominal", "Mouse nominal", "Acute logFC", "Blood severity r"]
    sns.heatmap(metrics, cmap="viridis", ax=ax_a, linewidths=0.35, cbar_kws={"label": "Observed value"})
    ax_a.set_title("Integrated dry-evidence matrix", pad=10)
    ax_a.set_xlabel("")
    ax_a.set_ylabel("")
    ax_a.set_xticklabels(ax_a.get_xticklabels(), rotation=35, ha="right")
    add_panel(ax_a, "A", -0.12, 1.08)

    ranked = priority.sort_values("integrated_evidence_score")
    ax_b.barh(ranked["gene_symbol"], ranked["integrated_evidence_score"], color=OKABE["purple"], edgecolor="#374151", lw=0.35)
    ax_b.set_xlabel("Integrated evidence score")
    ax_b.set_title("Prioritized genes from public evidence", pad=10)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    ax_c.axis("off")
    steps = [
        ("1", "Fixed gene panel", "Entry: SLC3A2/SLC7A11; actin endpoints: WASF2/TLN1/ACTB/MYH9/MYL6/FLNA"),
        ("2", "Dataset-layer scoring", "Human FDR, human nominal, acute direction, mouse CCI and blood severity kept separate"),
        ("3", "Cell-context check", "Bulk marker-proxy and true snRNA localization interpreted as localization clues"),
        ("4", "Claim boundary", "RNA co-variation supports prioritization; it does not establish cell-death mechanism"),
    ]
    for i, (num, title, detail) in enumerate(steps):
        y = 0.84 - i * 0.22
        circ = mpl.patches.Circle((0.08, y), 0.045, fc=OKABE["purple"], ec="#374151", lw=0.6, alpha=0.85)
        ax_c.add_patch(circ)
        ax_c.text(0.08, y, num, ha="center", va="center", fontsize=8, color="white", fontweight="bold")
        ax_c.text(0.17, y + 0.035, title, ha="left", va="center", fontsize=8.4, fontweight="bold")
        ax_c.text(0.17, y - 0.035, detail, ha="left", va="center", fontsize=7.1, wrap=True)
        if i < len(steps) - 1:
            ax_c.plot([0.08, 0.08], [y - 0.055, y - 0.16], color="#9CA3AF", lw=1.0)
    ax_c.set_xlim(0, 1)
    ax_c.set_ylim(0, 1)
    ax_c.set_title("Evidence-to-claim workflow", pad=10)
    add_panel(ax_c, "C", -0.12, 1.08)

    top = corr.sort_values("abs_r", ascending=False).head(8).copy()
    mshort = {"score_disulfidptosis_8gene": "8-gene", "score_transporter_2gene": "Transporter", "score_actin_6gene": "Actin"}
    cshort = {"Endothelial": "Endo", "Neuron": "Neuron", "Microglia": "Micro", "OPC": "OPC", "Astrocyte": "Astro", "Oligodendrocyte": "Oligo"}
    top["label"] = top["module"].map(mshort).fillna(top["module"]) + " | " + top["celltype_proxy"].map(cshort).fillna(top["celltype_proxy"])
    ax_d.barh(np.arange(len(top)), top["spearman_r"], color=[OKABE["red"] if v > 0 else OKABE["blue"] for v in top["spearman_r"]], edgecolor="#374151", lw=0.35)
    ax_d.set_yticks(np.arange(len(top)))
    ax_d.set_yticklabels(top["label"], fontsize=7.2)
    ax_d.invert_yaxis()
    ax_d.axvline(0, color="black", lw=0.8)
    ax_d.set_xlabel("Spearman r")
    ax_d.set_title("Top marker-proxy correlations", pad=10)
    sns.despine(ax=ax_d)
    add_panel(ax_d, "D", -0.12, 1.08)
    return export_figure(fig, "Fig8")


def make_fig9() -> Path:
    human = read_csv("v4_human_comp_20260605.csv")
    mouse = read_csv("v4_mouse_units_20260605.csv").head(12)
    treatment = read_csv("v4_treatment_module_20260605.csv")
    fig = plt.figure(figsize=(12.0, 8.2))
    gs = fig.add_gridspec(2, 2, hspace=0.60, wspace=0.58)
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    human = human.sort_values("mean_effect", ascending=False)
    human["label"] = human["comparison_short"].str.replace("193407", "GSE193407", regex=False).str.replace("209552", "GSE209552", regex=False)
    ax_a.barh(np.arange(len(human)), human["mean_effect"], color=[OKABE["red"] if x > 0 else "#BFC7D5" for x in human["mean_effect"]], edgecolor="#374151", lw=0.25)
    ax_a.set_yticks(np.arange(len(human)))
    ax_a.set_yticklabels(wrapped(human["label"].tolist(), 26), fontsize=6.8)
    ax_a.invert_yaxis()
    ax_a.set_xlabel("Mean effect across 8 genes")
    ax_a.set_title("Human comparison-level directionality", pad=10)
    for i, row in human.reset_index(drop=True).iterrows():
        ax_a.text(row["mean_effect"] + 0.025, i, f"FDR {int(row['fdr_genes'])}; nom {int(row['nominal_genes'])}", va="center", fontsize=6.4)
    sns.despine(ax=ax_a)
    add_panel(ax_a, "A", -0.12, 1.08)

    m = mouse.iloc[::-1].copy()
    y = np.arange(len(m))
    ax_b.barh(y, m["nominal_genes"], color=OKABE["sky"], edgecolor="#374151", lw=0.25, label="Nominal genes")
    ax_b.barh(y, m["panel_FDR_genes"], color=OKABE["blue"], edgecolor="#374151", lw=0.25, label="Panel-FDR genes")
    ax_b.set_yticks(y)
    ax_b.set_yticklabels([f"{r.time} {r.region} {r.treatment}" for r in m.itertuples()], fontsize=6.8)
    ax_b.set_xlabel("Number of panel hits")
    ax_b.set_title("Mouse CCI support-window ranking", pad=10)
    ax_b.legend(frameon=False, loc="lower right")
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    tm = treatment[treatment["scope"].isin(["region_treatment", "treatment"])].copy()
    tm["unit"] = tm["time"] + " " + tm["region"]
    heat = tm.pivot_table(index="unit", columns="module", values="mean_drug_minus_vehicle", aggfunc="mean")
    sns.heatmap(heat, cmap="RdBu_r", center=0, annot=True, fmt=".2f", linewidths=0.35, cbar_kws={"label": "Drug - vehicle logFC"}, ax=ax_c)
    ax_c.set_title("Treatment-stratified modulation in GSE163415", pad=10)
    ax_c.set_xlabel("")
    ax_c.set_ylabel("")
    add_panel(ax_c, "C", -0.12, 1.08)

    ax_d.axis("off")
    rows = [
        ("Human brain", "FDR-supported CTE-stage association; acute severe TBI directionality"),
        ("Mouse CCI", "Supportive time-region evidence; cross-species transcriptomic signal only"),
        ("Peripheral blood", "Severity context; not a substitute for brain tissue"),
        ("Cell context", "snRNA localization improves precision; bulk proxies remain exploratory"),
        ("Mechanism", "RNA evidence alone cannot prove disulfidptotic cell death"),
    ]
    ax_d.set_title("Evidence boundary audit", pad=10)
    for i, (label, text) in enumerate(rows):
        y0 = 0.88 - i * 0.17
        ax_d.add_patch(mpl.patches.FancyBboxPatch((0.02, y0 - 0.06), 0.94, 0.105, boxstyle="round,pad=0.014,rounding_size=0.015", fc="#F9FAFB", ec="#D1D5DB", lw=0.7))
        ax_d.text(0.05, y0, label, ha="left", va="center", fontsize=7.8, fontweight="bold")
        ax_d.text(0.34, y0, text, ha="left", va="center", fontsize=7.2)
    add_panel(ax_d, "D", -0.12, 1.08)
    return export_figure(fig, "Fig9")


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def remove_block(doc: Document, start_text: str, end_text: str, include_end: bool = False) -> None:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if start is None and text == start_text:
            start = i
            continue
        if start is not None and text == end_text:
            end = i if include_end else i - 1
            break
    if start is None or end is None or end < start:
        return
    for p in list(doc.paragraphs[start : end + 1]):
        remove_paragraph(p)


def remove_reference_containing(doc: Document, needle: str) -> None:
    for p in list(doc.paragraphs):
        if needle.lower() in p.text.lower():
            remove_paragraph(p)


def remove_note_paragraphs(doc: Document) -> None:
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Note."):
            remove_paragraph(p)


def remove_display_table(doc: Document) -> None:
    for table in list(doc.tables):
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Display item" in text or "Primary biological question" in text:
            table._element.getparent().remove(table._element)


def set_doc_font(doc: Document) -> None:
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Times New Roman"
            doc.styles[style_name].font.size = Pt(10)
    heading_like = {
        "ABSTRACT", "Introduction", "Materials and Methods", "Results", "Discussion", "Conclusion",
        "References", "Statements and Declarations", "Funding", "Competing Interests", "Author Contributions",
        "Ethics Approval", "Consent to Participate", "Consent for Publication", "Data Availability", "Code Availability",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        if text in heading_like or re.match(r"^[A-Z][A-Za-z ,/-]+$", text) and len(text) < 95 and not text.startswith("Fig."):
            for run in p.runs:
                run.bold = True


def polish_text(text: str) -> str:
    replacements = {
        "The study is therefore not a simple candidate-gene screen. It is an evidence-layering exercise that uses public transcriptomes to localize where a plausible disulfidptosis-adjacent process should be tested experimentally. The analysis is organized around dataset-defined evidence layers, injury severity, cellular context and a transporter-to-cytoskeleton biological logic.": (
            "This study uses public transcriptomes to localize where a disulfidptosis-like transporter-actin stress signal is most plausible after TBI/CTE. The design fixes the eight-gene panel before analysis and keeps acute brain tissue, chronic CTE stage, remote regional TBI, mouse CCI, peripheral severity and single-nucleus localization as distinct evidence layers."
        ),
        "Fig. 1 is not merely a workflow diagram. It converts the recently proposed cell-death concept of disulfidptosis into proxy layers that public transcriptomes can reasonably address.": (
            "The tiered design translated the recently proposed cell-death concept of disulfidptosis into RNA-level proxy layers that public transcriptomes can reasonably address."
        ),
        "Fig. 1 is the logical starting point of the manuscript.": "The study design establishes the interpretive starting point.",
        "rather than": "instead of",
        "The correct inference is not cell-type assignment, but co-localization prioritization.": "The correct inference is cell-context prioritization with explicit uncertainty about cellular origin.",
        "The appropriate response is not to ignore these alternatives, but to define experimental criteria.": "These alternatives require explicit evidentiary criteria.",
        "This is not a weakness; it is the proper evidentiary boundary for a public-data manuscript.": "This boundary is essential for a public-data manuscript.",
        "wet-lab validation": "future tissue-level validation",
        "wet-lab": "tissue-level",
        "validation window": "support window",
        "validation windows": "support windows",
        "Prioritized validation genes": "Prioritized genes from public evidence",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("Fig. 9 and Fig. 10", "Fig. 4 and Fig. 5")
    text = text.replace("Fig. 9", "Fig. 4")
    text = text.replace("Fig. 10", "Fig. 5")
    text = text.replace("Fig. 4 turns", "Fig. 6 turns")
    text = text.replace("Fig. 4A", "Fig. 6A").replace("Fig. 4B", "Fig. 6B").replace("Fig. 4C", "Fig. 6C").replace("Fig. 4D", "Fig. 6D")
    text = text.replace("Fig. 4. Mouse", "Fig. 6. Mouse")
    text = text.replace("Fig. 5 places", "Fig. 7 places")
    text = text.replace("Fig. 5A", "Fig. 7A").replace("Fig. 5B", "Fig. 7B").replace("Fig. 5C", "Fig. 7C").replace("Fig. 5D", "Fig. 7D")
    text = text.replace("Fig. 5. Peripheral", "Fig. 7. Peripheral")
    text = text.replace("Fig. 6 should", "Fig. 8 should")
    text = text.replace("Fig. 6A", "Fig. 8A").replace("Fig. 6B", "Fig. 8B").replace("Fig. 6C", "Fig. 8C").replace("Fig. 6D", "Fig. 8D")
    text = text.replace("Fig. 6. Integrated", "Fig. 8. Integrated")
    text = text.replace("Fig. 7 addresses", "Fig. 9 addresses")
    text = text.replace("Fig. 7 connects", "Fig. 9 connects")
    text = text.replace("Fig. 7A", "Fig. 9A").replace("Fig. 7B", "Fig. 9B").replace("Fig. 7C", "Fig. 9C").replace("Fig. 7D", "Fig. 9D")
    text = text.replace("Fig. 7. Evidence", "Fig. 9. Evidence")
    text = text.replace("Fig. 8 therefore converts NeuN, GFAP, IBA1, OLIG2 and CD31 into a co-localization matrix.", "The marker-proxy and single-nucleus analyses together define the current cell-context boundary.")
    text = text.replace("Fig. 3D and Fig. 8B show", "Fig. 3D and Table 6 show")
    text = text.replace("Fig. 8 defines the upgrade criteria.", "The current evidence defines a strict interpretation ceiling.")
    text = re.sub(r"\s+", " ", text).strip() if text.strip() else text
    return text


def replace_paragraph_texts(doc: Document) -> None:
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        new_text = polish_text(p.text)
        if new_text != p.text:
            p.text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        p.text = polish_text(p.text)


def replace_images(doc: Document, fig_paths: dict[str, Path]) -> None:
    image_paras = [p for p in doc.paragraphs if "<w:drawing" in p._p.xml]
    ordered = ["Fig1", "Fig2", "Fig3", "Fig4", "Fig5", "Fig6", "Fig7", "Fig8"]
    for p, key in zip(image_paras, ordered):
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_paths[key]), width=Inches(6.7))


def insert_before(paragraph, text: str, bold: bool = False) -> None:
    p = paragraph.insert_paragraph_before(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = bold


def insert_figure9_before_discussion(doc: Document, fig_path: Path) -> None:
    discussion = next((p for p in doc.paragraphs if p.text.strip() == "Discussion"), None)
    if discussion is None:
        return
    caption = (
        "Fig. 9. Evidence audit across human, mouse and interpretation layers. Panel A audits human comparison-level directionality and separates FDR-supported chronic CTE-stage evidence from acute directional support. "
        "Panel B ranks mouse CCI support windows, highlighting 3DPI hippocampus as the most coherent animal layer. Panel C displays treatment-stratified module modulation within GSE163415 as contextual information. "
        "Panel D summarizes the evidence boundary used for interpretation; RNA-level association and localization do not establish disulfidptotic cell death."
    )
    body = (
        "The evidence audit separates statistical strength from biological interpretation. Human CTE-stage analyses provide the clearest FDR-supported brain association, acute severe TBI contributes directional cytoskeletal stress, mouse CCI provides time-region support, and peripheral blood remains a severity context. These layers are informative when kept separate and misleading if collapsed into a single proof claim."
    )
    insert_before(discussion, caption)
    fig_para = discussion.insert_paragraph_before("")
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run().add_picture(str(fig_path), width=Inches(6.7))
    insert_before(discussion, "Fig. 9", bold=True)
    insert_before(discussion, body)
    insert_before(discussion, "Evidence Audit Separated Statistical Support From Mechanistic Interpretation", bold=True)


def add_discussion_xct_paragraph(doc: Document) -> None:
    discussion = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Discussion"), None)
    if discussion is None:
        return
    insert_after = doc.paragraphs[discussion + 3] if discussion + 3 < len(doc.paragraphs) else doc.paragraphs[discussion]
    new_text = (
        "SLC7A11/xCT provides a focused biochemical interpretation of the transporter signal. Cystine import can support glutathione synthesis, yet it also increases the demand for intracellular reducing power because cystine must be reduced to cysteine. Under post-injury metabolic stress, this demand can become unfavorable when NADPH regeneration is limited. The same antiporter exports glutamate, so SLC7A11 upregulation after neural injury may sit between antioxidant adaptation, extracellular glutamate pressure and disulfide stress. This explains why the present analysis treats SLC3A2/SLC7A11 as transporter pressure within a disulfidptosis-like axis, not as an automatically protective or lethal marker."
    )
    p = insert_after.insert_paragraph_before(new_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def polish_captions(doc: Document) -> None:
    caption_replacements = {
        "Fig. 2. Human brain transcriptomic evidence across acute, chronic and regional comparisons. Panel A shows the 8-gene effect map across comparison contexts. Panel B focuses on GSE193407 continuous CTE stage and identifies SLC3A2, SLC7A11, WASF2 and TLN1 as FDR-supported stage-associated genes. Panel C summarizes recurrence across human comparisons by combining nominal support count and maximum absolute effect, distinguishing transporter-regulatory genes from downstream cytoskeletal endpoint genes.": (
            "Fig. 2. Human brain transcriptomic evidence across acute, chronic and regional comparisons. Panel A shows the 8-gene effect map across comparison contexts. Panel B focuses on the GSE193407 continuous CTE-stage trend and identifies SLC3A2, SLC7A11, WASF2 and TLN1 as FDR-supported stage-associated genes. Panel C summarizes recurrence across human comparisons by combining nominal support count and maximum absolute effect. In panels A and B, * denotes FDR < 0.05 and † denotes nominal P < 0.05 without FDR support."
        ),
        "Fig. 3. Acute severe TBI evidence from GSE209552. Panel A compares sample-level transporter/stress, actin-target and mitochondrial-synergy module scores between TBI and control samples. Panel B shows gene-level logFC for all eight genes, with ACTB, MYH9 and MYL6 carrying the strongest nominal cytoskeletal directionality. Panel C places the disulfidptosis gene set in a ranked pathway context. Panel D shows bulk marker-proxy correlations used for later co-localization prioritization, not for direct cell-type assignment.": (
            "Fig. 3. Acute severe TBI evidence from GSE209552. Panel A compares sample-level 8-gene, transporter and actin-cytoskeletal module scores between TBI and control samples. Panel B shows gene-level logFC for all eight genes, with ACTB, MYH9 and MYL6 carrying the strongest nominal cytoskeletal directionality. Panel C places the disulfidptosis gene set in a ranked pathway context. Panel D shows bulk marker-proxy correlations used for cell-context prioritization. In panel B, * denotes nominal P < 0.05 and † denotes 0.05 <= P < 0.10; in panel C, *** denotes the strongest ranked enrichment signal."
        ),
        "Fig. 7. Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data. Panel A shows PCA of the eight-gene expression matrix. Panel B compares module scores across control, mild, moderate and severe groups. Panel C shows gene-level Spearman correlations with severity score 0-3. Panel D summarizes severity association at the module level.": (
            "Fig. 7. Peripheral severity-focused analysis of GSE223245 whole blood/PBMC data. Panel A shows PCA of the eight-gene expression matrix. Panel B compares module scores across control, mild, moderate and severe groups. Panel C shows gene-level Spearman correlations with severity score 0-3. Panel D summarizes severity association at the module level. In panel C, *, ** and *** denote FDR < 0.05, FDR < 0.01 and FDR < 0.001, respectively; † denotes nominal support without FDR significance."
        ),
        "Fig. 8. Integrated prioritization of genes and validation readouts. Panel A combines human FDR support, human nominal support, mouse nominal support, acute logFC and blood severity association into a dry-evidence matrix. Panel B converts the same information into an integrated validation score. Panel C translates the ranking into a validation sequence. Panel D displays top marker-proxy correlations used to prioritize co-staining while preserving the limitation that bulk correlations do not prove cell-type origin.": (
            "Fig. 8. Integrated prioritization of genes and evidence boundaries. Panel A combines human FDR support, human nominal support, mouse nominal support, acute logFC and blood severity association into a dry-evidence matrix. Panel B converts the same information into an integrated public-evidence score. Panel C summarizes the evidence-to-claim workflow used to interpret the RNA-level findings. Panel D displays top marker-proxy correlations while preserving the limitation that bulk correlations do not prove cell-type origin."
        ),
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in caption_replacements:
            p.text = caption_replacements[t]


def clean_doc(fig_paths: dict[str, Path]) -> Path:
    src = FINAL_OUT / "01_Manuscript.docx"
    backup = CHECK_OUT / "01_Manuscript_before_language_figure_polish_20260623.docx"
    if src.exists() and not backup.exists():
        shutil.copy2(src, backup)
    doc = Document(src)
    remove_block(doc, "Literature-Driven Extension Genes and Cross-Compartment Modeling", "Display-Item Integration")
    remove_block(doc, "Display-Item Integration", "Results")
    remove_block(doc, "Reading Logic for the Results", "A Tiered Design Converted Disulfidptosis Into a Testable Transporter-Actin Axis")
    remove_block(doc, "Literature-Driven Extension Upgraded the Working Model to Transporter Entry, Redox Prerequisite, Inflammatory Amplification and Cytoskeletal Remodeling", "Mouse CCI Localized the First Validation Window to 3DPI Hippocampus")
    remove_block(doc, "Evidence Audit Prevented Collapsing Acute, Chronic and Mouse Layers Into One Claim", "Peripheral Blood Added Severity Context but Not Brain-Mechanism Evidence")
    remove_block(doc, "Mechanistic Upgrading Requires Matched Transporter, Cytoskeletal, Redox and Localization Evidence", "Discussion")
    remove_display_table(doc)
    remove_note_paragraphs(doc)
    remove_reference_containing(doc, "Dilated Cardiomyopathy")
    replace_paragraph_texts(doc)
    polish_captions(doc)
    image_paras = [p for p in doc.paragraphs if "<w:drawing" in p._p.xml]
    while len(image_paras) > 8:
        remove_paragraph(image_paras[-1])
        image_paras = [p for p in doc.paragraphs if "<w:drawing" in p._p.xml]
    replace_images(doc, fig_paths)
    insert_figure9_before_discussion(doc, fig_paths["Fig9"])
    add_discussion_xct_paragraph(doc)
    set_doc_font(doc)
    doc.save(src)
    adjusted = DOC_OUT / "Manuscript_NeurochemicalResearch_polished_20260623.docx"
    doc.save(adjusted)
    return src


def make_resolution_summary() -> Path:
    rows = []
    for tif in sorted(FINAL_TIFF.glob("Fig*.tif")):
        im = Image.open(tif)
        rows.append(
            {
                "file": tif.name,
                "pixel_width": im.width,
                "pixel_height": im.height,
                "dpi": im.info.get("dpi", ""),
                "width_at_600dpi_mm": round(im.width / 600 * 25.4, 1),
                "height_at_600dpi_mm": round(im.height / 600 * 25.4, 1),
            }
        )
    out = CHECK_OUT / "Figure_resolution_check_after_polish_20260623.csv"
    pd.DataFrame(rows).to_csv(out, index=False, encoding="utf-8-sig")
    return out


def main() -> None:
    ensure_dirs()
    configure_style()
    fig_paths = {
        "Fig1": make_fig1(),
        "Fig2": make_fig2(),
        "Fig3": make_fig3(),
        "Fig4": make_fig4(),
        "Fig5": make_fig5(),
        "Fig6": make_fig6(),
        "Fig7": make_fig7(),
        "Fig8": make_fig8(),
        "Fig9": make_fig9(),
    }
    doc_path = clean_doc(fig_paths)
    summary_path = make_resolution_summary()
    print(f"Updated manuscript: {doc_path}")
    print(f"Polished PNG figures: {POLISHED_PNG}")
    print(f"Updated TIFF figures: {FINAL_TIFF}")
    print(f"Resolution summary: {summary_path}")


if __name__ == "__main__":
    main()
