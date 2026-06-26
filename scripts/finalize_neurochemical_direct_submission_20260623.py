from __future__ import annotations

import re
import shutil
import textwrap
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw


ROOT = Path.cwd()
PROJECT_DIR = next(p.parent for p in ROOT.rglob("TBI_disulfidptosis_submission_package_20260611") if p.is_dir())
TABLE_DIR = PROJECT_DIR / "tables"
OUT_ROOT = next(p for p in ROOT.rglob("*Neurochemical_Research*") if p.is_dir() and (p / "06_Final_submission_ready").exists())
FINAL_OUT = OUT_ROOT / "06_Final_submission_ready"
FINAL_TIFF = FINAL_OUT / "06_Figures_600dpi_TIFF"
POLISHED_DIR = OUT_ROOT / "08_polished_figures_and_manuscript_20260623"
POLISHED_PNG = POLISHED_DIR / "png"
CHECK_OUT = FINAL_OUT / "07_Checks"
ROOT_TIFF = OUT_ROOT / "02_figures_600dpi_TIFF"

AUTHOR_LINE = ("Chenxu Zhang", "a", "Mingyang Zhang", "a,*")
AFFILIATION = (
    "a Department of Forensic Sciences, The Affiliated Guangji Hospital, School of Basic Medicine, "
    "Suzhou Medical College of Soochow University, Suzhou, China"
)
CORRESPONDING = (
    "Professor Mingyang Zhang, Department of Forensic Sciences, The Affiliated Guangji Hospital, "
    "School of Basic Medicine, Suzhou Medical College of Soochow University, Suzhou, China; "
    "Email: mingyangzhang@suda.edu.cn"
)
TITLE = (
    "Spatiotemporal, Severity-Linked and Cell-Context Prioritization of a "
    "Disulfidptosis-Like Transporter-Actin Stress Axis After Traumatic Brain Injury"
)
REPOSITORY = "https://github.com/deerlet5-nina/tbi-disulfidptosis"

OKABE = {
    "blue": "#0072B2",
    "sky": "#56B4E9",
    "green": "#009E73",
    "orange": "#E69F00",
    "red": "#D55E00",
    "purple": "#CC79A7",
    "gray": "#8A95A5",
    "black": "#222222",
}

GENES = ["SLC3A2", "SLC7A11", "WASF2", "TLN1", "ACTB", "MYH9", "MYL6", "FLNA"]


def configure_style() -> None:
    mpl.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial", "DejaVu Sans", "Microsoft YaHei"],
            "font.size": 8.8,
            "axes.titlesize": 10.2,
            "axes.labelsize": 8.8,
            "xtick.labelsize": 8.1,
            "ytick.labelsize": 8.1,
            "legend.fontsize": 7.4,
            "axes.unicode_minus": False,
            "figure.dpi": 160,
            "savefig.dpi": 320,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
        }
    )
    sns.set_theme(style="white", font="Arial")


def ensure_dirs() -> None:
    for path in [FINAL_TIFF, POLISHED_PNG, CHECK_OUT, ROOT_TIFF]:
        path.mkdir(parents=True, exist_ok=True)


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLE_DIR / name)


def add_panel(ax: mpl.axes.Axes, label: str, x: float = -0.10, y: float = 1.07) -> None:
    ax.text(x, y, label, transform=ax.transAxes, fontsize=12, fontweight="bold", va="top", ha="left")


def wrap_text(s: str, width: int = 32) -> str:
    return "\n".join(textwrap.wrap(str(s), width=width, break_long_words=False))


def export_figure(fig: plt.Figure, figure_name: str) -> Path:
    png_path = POLISHED_PNG / f"{figure_name}.png"
    tiff_path = FINAL_TIFF / f"{figure_name}.tif"
    fig.savefig(png_path, bbox_inches="tight", dpi=320)
    plt.close(fig)
    im = Image.open(png_path).convert("RGB")
    im.save(tiff_path, dpi=(600, 600), compression="tiff_lzw")
    shutil.copy2(tiff_path, ROOT_TIFF / tiff_path.name)
    return png_path


def make_fig1() -> Path:
    fig = plt.figure(figsize=(11.6, 7.3))
    gs = fig.add_gridspec(
        2, 2,
        height_ratios=[1.0, 1.18],
        width_ratios=[1.15, 1.0],
        hspace=0.48,
        wspace=0.36,
    )
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    ax_a.axis("off")
    boxes = [
        ("Human acute\nsevere TBI\nGSE209552", 0.06, 0.62, OKABE["red"]),
        ("Human chronic\nCTE stage\nGSE193407/GSE319253", 0.54, 0.62, OKABE["blue"]),
        ("Mouse CCI\ntime-region\nGSE163415", 0.06, 0.18, OKABE["green"]),
        ("Peripheral\nseverity\nGSE223245", 0.54, 0.18, OKABE["orange"]),
    ]
    for text, x, y, color in boxes:
        rect = mpl.patches.FancyBboxPatch(
            (x, y), 0.36, 0.24,
            boxstyle="round,pad=0.02,rounding_size=0.018",
            fc=color, ec="#374151", alpha=0.13, lw=0.8,
        )
        ax_a.add_patch(rect)
        ax_a.text(x + 0.18, y + 0.12, text, ha="center", va="center", fontsize=8.2, linespacing=1.10)
    ax_a.annotate("", xy=(0.54, 0.74), xytext=(0.42, 0.74), arrowprops=dict(arrowstyle="->", lw=1.0, color="#4B5563"))
    ax_a.annotate("", xy=(0.54, 0.30), xytext=(0.42, 0.30), arrowprops=dict(arrowstyle="->", lw=1.0, color="#4B5563"))
    ax_a.text(
        0.50, 0.50,
        "Fixed 8-gene panel\nSLC3A2/SLC7A11 + actin endpoints",
        ha="center", va="center", fontsize=9.2, fontweight="bold",
    )
    ax_a.set_xlim(0, 1)
    ax_a.set_ylim(0, 1)
    ax_a.set_title("Evidence-layer design", pad=8)
    add_panel(ax_a, "A", -0.08, 1.05)

    layers = ["acute", "remote region", "CTE course", "CTE external", "mouse CCI", "blood severity", "snRNA localization"]
    colors = [OKABE["red"], OKABE["gray"], OKABE["blue"], OKABE["sky"], OKABE["green"], OKABE["orange"], OKABE["purple"]]
    ax_b.barh(np.arange(len(layers)), np.ones(len(layers)), color=colors, edgecolor="#374151", linewidth=0.2)
    ax_b.set_yticks(np.arange(len(layers)))
    ax_b.set_yticklabels(layers)
    ax_b.set_xlim(0, 1.05)
    ax_b.set_xlabel("Evidence layer present")
    ax_b.set_title("Time, region, severity and cell axes", pad=8)
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.05)

    ax_c.set_title("Pre-fixed transporter-actin panel", pad=10)
    ax_c.set_xlim(0.65, 4.05)
    ax_c.set_ylim(-0.8, 7.8)
    ax_c.set_yticks([])
    ax_c.set_xticks([1, 2, 3])
    ax_c.set_xticklabels(["cystine import", "actin control", "cytoskeletal endpoint"])
    gene_rows = [
        ("SLC3A2", "Transporter", 1, 0.0, OKABE["red"]),
        ("SLC7A11", "Transporter", 1, 1.0, OKABE["red"]),
        ("WASF2", "WRC/remodeling", 2, 2.9, OKABE["blue"]),
        ("TLN1", "Adhesion/tension", 2, 4.0, OKABE["blue"]),
        ("ACTB", "Actin scaffold", 3, 5.0, OKABE["green"]),
        ("MYH9", "Myosin/tension", 3, 6.0, OKABE["green"]),
        ("MYL6", "Myosin light chain", 3, 7.0, OKABE["green"]),
        ("FLNA", "Filamin scaffold", 3, 7.55, OKABE["green"]),
    ]
    for gene, role, x, y, color in gene_rows:
        ax_c.scatter(x, y, s=132, c=color, edgecolor="black", lw=0.8, zorder=3)
        ax_c.text(x + 0.13, y, f"{gene}  {role}", va="center", ha="left", fontsize=7.8)
    sns.despine(ax=ax_c, left=True)
    add_panel(ax_c, "C", -0.08, 1.05)

    ax_d.axis("off")
    ax_d.set_title("Claim-boundary ladder", pad=8)
    rows = [
        ("FDR human brain", "association", "strongest RNA-level support"),
        ("Nominal direction", "candidate clue", "exploratory trend"),
        ("Mouse CCI", "support layer", "time-region context"),
        ("Peripheral blood", "severity context", "systemic signal"),
        ("Bulk marker proxy", "cell-context clue", "not cell origin"),
    ]
    for i, (label, claim, limit) in enumerate(rows):
        y = 0.88 - i * 0.17
        color = [OKABE["blue"], OKABE["gray"], OKABE["green"], OKABE["orange"], OKABE["purple"]][i]
        ax_d.add_patch(
            mpl.patches.FancyBboxPatch(
                (0.02, y - 0.062), 0.94, 0.115,
                boxstyle="round,pad=0.014,rounding_size=0.018",
                fc="#F9FAFB", ec="#D1D5DB", lw=0.7,
            )
        )
        ax_d.add_patch(mpl.patches.Rectangle((0.02, y - 0.062), 0.025, 0.115, fc=color, ec="none"))
        ax_d.text(0.07, y + 0.018, label, ha="left", va="center", fontsize=7.7, fontweight="bold", color="#111827")
        ax_d.text(0.52, y + 0.018, claim, ha="left", va="center", fontsize=7.7, color="#111827")
        ax_d.text(0.52, y - 0.024, limit, ha="left", va="center", fontsize=6.8, color="#6B7280")
    ax_d.text(
        0.02, 0.02,
        "The panel defines interpretable RNA-level evidence, while protein, redox and cytoskeletal morphology remain outside the present dataset.",
        ha="left", va="bottom", fontsize=7.0, color="#374151", wrap=True,
    )
    add_panel(ax_d, "D", -0.10, 1.05)
    return export_figure(fig, "Fig1")


def make_fig9() -> Path:
    human = read_csv("v4_human_comp_20260605.csv")
    mouse = read_csv("v4_mouse_units_20260605.csv").head(12)
    treatment = read_csv("v4_treatment_module_20260605.csv")

    fig = plt.figure(figsize=(12.6, 8.7))
    gs = fig.add_gridspec(2, 2, hspace=0.58, wspace=0.48, width_ratios=[1.06, 1.0])
    ax_a, ax_b, ax_c, ax_d = [fig.add_subplot(gs[i]) for i in [(0, 0), (0, 1), (1, 0), (1, 1)]]

    label_map = {
        "193407 CTE stage trend": "CTE stage trend",
        "193407 late CTE stage3 4 vs stage0": "Late CTE vs stage 0",
        "193407 any CTE vs stage0": "Any CTE vs stage 0",
        "GSE209552 human acute brain severe": "Acute severe TBI",
        "319253 CTE vs Control": "External CTE",
        "104687 HIP": "Remote HIP",
        "104687 PCx": "Remote PCx",
        "104687 TCx": "Remote TCx",
        "104687 FWM": "Remote FWM",
    }
    h = human.copy()
    h["label"] = h["comparison_short"].map(label_map).fillna(h["comparison_short"])
    h = h.sort_values("mean_effect", ascending=True)
    y = np.arange(len(h))
    colors = [OKABE["red"] if v > 0.25 else "#BFC7D5" for v in h["mean_effect"]]
    ax_a.barh(y, h["mean_effect"], color=colors, edgecolor="#374151", lw=0.25)
    ax_a.set_yticks(y)
    ax_a.set_yticklabels(h["label"], fontsize=7.5)
    ax_a.set_xlabel("Mean effect across 8 genes")
    ax_a.set_title("Human comparison directionality", pad=10)
    for i, row in h.reset_index(drop=True).iterrows():
        label = f"FDR {int(row.fdr_genes)} / nom {int(row.nominal_genes)}"
        xpos = row.mean_effect + 0.035 if row.mean_effect >= 0 else row.mean_effect - 0.035
        ha = "left" if row.mean_effect >= 0 else "right"
        ax_a.text(xpos, i, label, va="center", ha=ha, fontsize=6.7, color="#374151")
    ax_a.set_xlim(min(-0.03, h["mean_effect"].min() - 0.10), h["mean_effect"].max() + 0.42)
    sns.despine(ax=ax_a)
    add_panel(ax_a, "A", -0.12, 1.08)

    m = mouse.iloc[::-1].copy()
    m["label"] = [f"{r.time.replace('DPI','D')} {r.region} {r.treatment}" for r in m.itertuples()]
    y = np.arange(len(m))
    ax_b.barh(y, m["nominal_genes"], color="#B7D7F0", edgecolor="#374151", lw=0.25, label="Nominal genes")
    ax_b.barh(y, m["panel_FDR_genes"], color=OKABE["blue"], edgecolor="#374151", lw=0.25, label="Panel-FDR genes")
    ax_b.set_yticks(y)
    ax_b.set_yticklabels(m["label"], fontsize=7.0)
    ax_b.set_xlabel("Number of panel hits")
    ax_b.set_title("Mouse CCI support-window ranking", pad=10)
    ax_b.legend(frameon=False, loc="lower right")
    sns.despine(ax=ax_b)
    add_panel(ax_b, "B", -0.12, 1.08)

    tm = treatment[treatment["scope"].isin(["region_treatment", "treatment"])].copy()
    tm["module_short"] = tm["module"].replace(
        {
            "Actin cytoskeleton targets": "Actin",
            "Core transporter/stress": "Transporter",
        }
    )
    tm["unit"] = tm.apply(
        lambda r: f"{r['time'].replace('DPI','D')} {r['region'] if r['region'] != 'all' else 'all'}",
        axis=1,
    )
    unit_order = ["3D Hipp", "3D Hypo", "3D Thal", "3D all", "29D Hipp", "29D Hypo", "29D Thal", "29D all"]
    heat = (
        tm.pivot_table(index="unit", columns="module_short", values="mean_drug_minus_vehicle", aggfunc="mean")
        .reindex(unit_order)
        .reindex(columns=["Transporter", "Actin"])
    )
    sns.heatmap(
        heat,
        cmap="RdBu_r",
        center=0,
        annot=True,
        fmt=".2f",
        linewidths=0.35,
        linecolor="white",
        cbar_kws={"label": "Drug - vehicle logFC", "shrink": 0.78, "pad": 0.02},
        ax=ax_c,
    )
    ax_c.set_title("Treatment-stratified CCI modulation", pad=10)
    ax_c.set_xlabel("")
    ax_c.set_ylabel("")
    ax_c.set_xticklabels(ax_c.get_xticklabels(), rotation=0)
    ax_c.tick_params(axis="both", length=0)
    add_panel(ax_c, "C", -0.12, 1.08)

    ax_d.axis("off")
    ax_d.set_title("Evidence boundary ladder", pad=10)
    rows = [
        ("Human brain", "CTE-stage association; acute directionality"),
        ("Mouse CCI", "supportive time-region transcriptomic layer"),
        ("Blood severity", "systemic severity context"),
        ("Cell context", "snRNA localization plus bulk proxy clues"),
        ("Mechanism", "RNA-level candidate axis; direct proof outside dataset"),
    ]
    colors = [OKABE["blue"], OKABE["green"], OKABE["orange"], OKABE["purple"], OKABE["gray"]]
    for i, (label, text) in enumerate(rows):
        y = 0.88 - i * 0.17
        ax_d.add_patch(
            mpl.patches.FancyBboxPatch(
                (0.04, y - 0.06), 0.90, 0.11,
                boxstyle="round,pad=0.014,rounding_size=0.018",
                fc="#F9FAFB", ec="#D1D5DB", lw=0.7,
            )
        )
        ax_d.add_patch(mpl.patches.Circle((0.09, y), 0.028, fc=colors[i], ec="#374151", lw=0.4))
        ax_d.text(0.09, y, str(i + 1), ha="center", va="center", fontsize=7.4, color="white", fontweight="bold")
        ax_d.text(0.15, y + 0.019, label, ha="left", va="center", fontsize=7.8, fontweight="bold", color="#111827")
        ax_d.text(0.15, y - 0.028, wrap_text(text, 54), ha="left", va="center", fontsize=6.8, color="#4B5563")
    ax_d.text(
        0.04, 0.03,
        "The ladder is included as a visual audit of inference strength, not as a separate quantitative table.",
        ha="left", va="bottom", fontsize=7.0, color="#374151",
    )
    add_panel(ax_d, "D", -0.12, 1.08)
    return export_figure(fig, "Fig9")


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_format(paragraph, *, title: bool = False, heading: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(18 if title else 12)
        if title or heading:
            run.bold = True


def add_author_runs(paragraph) -> None:
    clear_paragraph(paragraph)
    first, first_sup, second, second_sup = AUTHOR_LINE
    paragraph.add_run(first + " ")
    r = paragraph.add_run(first_sup)
    r.font.superscript = True
    paragraph.add_run(", " + second + " ")
    r = paragraph.add_run(second_sup)
    r.font.superscript = True
    set_paragraph_format(paragraph)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_front_matter(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == TITLE:
            set_paragraph_format(p, title=True)
        elif "[Author" in t or t in {"Chenxu Zhang a, Mingyang Zhang a,*", "Chenxu Zhang a Mingyang Zhang a,*"}:
            add_author_runs(p)
        elif t.startswith("a ") and ("Department of" in t or "[Department]" in t):
            p.text = AFFILIATION
            set_paragraph_format(p)
        elif t.startswith("1 These authors contributed equally") or "Delete if not applicable" in t:
            remove_paragraph(p)
        elif t == "*Corresponding author:":
            set_paragraph_format(p, heading=True)
        elif "Corresponding author" in t and t != "*Corresponding author:":
            p.text = CORRESPONDING
            set_paragraph_format(p)
        elif "Professor/Dr." in t or "[active e-mail address]" in t:
            p.text = CORRESPONDING
            set_paragraph_format(p)


def paragraph_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def replace_text_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        t = p.text
        for old, new in replacements.items():
            if old in t:
                t = t.replace(old, new)
        if t != p.text:
            p.text = t
            set_paragraph_format(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    for old, new in replacements.items():
                        if old in t:
                            t = t.replace(old, new)
                    if t != p.text:
                        p.text = t
                        set_paragraph_format(p)


def update_declarations(doc: Document) -> None:
    section_bodies = {
        "Funding": "The authors received no specific funding for this work.",
        "Competing Interests": "The authors declare no competing interests.",
        "Author Contributions": (
            "Conceptualization: Mingyang Zhang, Chenxu Zhang; Methodology: Chenxu Zhang; "
            "Formal analysis and investigation: Chenxu Zhang; Data curation: Chenxu Zhang; "
            "Visualization: Chenxu Zhang; Writing - original draft: Chenxu Zhang; "
            "Writing - review and editing: Mingyang Zhang, Chenxu Zhang; Supervision: Mingyang Zhang. "
            "All authors read and approved the final manuscript."
        ),
        "Ethics Approval": (
            "This study used publicly available de-identified transcriptomic datasets and did not involve "
            "newly recruited human participants, new human biospecimen collection or new animal experiments."
        ),
        "Consent to Participate": "Not applicable.",
        "Consent for Publication": "Not applicable.",
        "Data Availability": (
            "The datasets analyzed in this study are publicly available from the Gene Expression Omnibus "
            "(GSE209552, GSE193407, GSE319253, GSE104687, GSE163415 and GSE223245). Processed data tables "
            f"and analysis scripts are available on GitHub at {REPOSITORY}."
        ),
        "Code Availability": f"Analysis scripts are available on GitHub at {REPOSITORY}.",
    }
    replacements = {
        "[To be confirmed with the corresponding author. If no specific funding supported this work, state: The authors received no specific funding for this work.]":
            "The authors received no specific funding for this work.",
        "The authors declare no competing interests. [Confirm before submission.]":
            "The authors declare no competing interests.",
        "Conceptualization: [corresponding author], [first author]; Methodology: [first author]; Formal analysis and investigation: [first author]; Data curation: [first author]; Visualization: [first author]; Writing - original draft: [first author]; Writing - review and editing: [corresponding author], [first author]; Supervision: [corresponding author]. All authors read and approved the final manuscript.":
            "Conceptualization: Mingyang Zhang, Chenxu Zhang; Methodology: Chenxu Zhang; Formal analysis and investigation: Chenxu Zhang; Data curation: Chenxu Zhang; Visualization: Chenxu Zhang; Writing - original draft: Chenxu Zhang; Writing - review and editing: Mingyang Zhang, Chenxu Zhang; Supervision: Mingyang Zhang. All authors read and approved the final manuscript.",
        "The datasets analyzed in this study are publicly available from the Gene Expression Omnibus, including GSE209552, GSE193407, GSE319253, GSE104687, GSE163415 and GSE223245. Processed tables and analysis scripts can be deposited in a public repository or made available by the corresponding author upon reasonable request.":
            f"The datasets analyzed in this study are publicly available from the Gene Expression Omnibus (GSE209552, GSE193407, GSE319253, GSE104687, GSE163415 and GSE223245). Processed data tables and analysis scripts are available on GitHub at {REPOSITORY}.",
        "[To be completed before submission: provide a GitHub/Zenodo repository link, or state that analysis code is available from the corresponding author upon reasonable request.]":
            f"Analysis scripts are available on GitHub at {REPOSITORY}.",
    }
    replace_text_in_doc(doc, replacements)
    headings = set(section_bodies) | {"Statements and Declarations", "References"}
    for i, p in enumerate(doc.paragraphs):
        heading = p.text.strip()
        if heading not in section_bodies:
            continue
        j = i + 1
        while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
            j += 1
        if j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() not in headings:
            doc.paragraphs[j].text = section_bodies[heading]
            set_paragraph_format(doc.paragraphs[j])


def move_declarations_before_references(doc: Document) -> None:
    ref_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References"), None)
    decl_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Statements and Declarations"), None)
    if ref_idx is None or decl_idx is None or decl_idx < ref_idx:
        return
    ref_el = doc.paragraphs[ref_idx]._element
    move_elements = [p._element for p in doc.paragraphs[decl_idx:]]
    for el in move_elements:
        ref_el.addprevious(el)


def add_reference_numbers(doc: Document) -> None:
    ref_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References"), None)
    decl_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Statements and Declarations"), None)
    if ref_idx is None:
        return
    end_idx = decl_idx if decl_idx is not None and decl_idx > ref_idx else len(doc.paragraphs)
    n = 1
    for p in doc.paragraphs[ref_idx + 1:end_idx]:
        t = p.text.strip()
        if not t:
            continue
        t = re.sub(r"^\[\d+\]\s*", "", t)
        p.text = f"[{n}] {t}"
        set_paragraph_format(p)
        n += 1


def add_body_citations(doc: Document) -> None:
    replacements = {
        "Disulfidptosis also requires conceptual unpacking.": "Disulfidptosis also requires conceptual unpacking [1, 2].",
        "Six GEO evidence layers were included.": "Six GEO evidence layers were included, with dataset-specific interpretation guided by the original public-data studies [4, 6-8].",
        "The model is biologically plausible because cytoskeletal stress after TBI is not restricted to one cell type.": "The model is biologically plausible because cytoskeletal stress after TBI is not restricted to one cell type [1, 2, 5].",
        "SLC7A11/xCT provides a focused biochemical interpretation of the transporter signal.": "SLC7A11/xCT provides a focused biochemical interpretation of the transporter signal, consistent with recent TBI evidence implicating Nrf2-related injury biology and SLC7A11-mediated disulfidptosis/neuroinflammation [3, 9].",
    }
    replace_text_in_doc(doc, replacements)


def update_captions_and_text(doc: Document) -> None:
    replacements = {
        "Fig. 1. Study design and inferential guardrails for the fixed transporter-actin panel. Panel A shows the evidence-layer design, assigning acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets to distinct roles. Panel B lists the time, region and severity axes addressed by the study. Panel C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints. Panel D states the wording rules used in the manuscript, separating FDR-supported associations, nominal candidate clues, animal support, peripheral severity context and bulk marker-proxy prioritization.":
            "Fig. 1. Study design and claim boundaries for the fixed transporter-actin panel. Panel A assigns acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets to distinct evidence layers. Panel B lists the time, region, severity and cell-context axes addressed by the study. Panel C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints. Panel D summarizes the claim-boundary ladder used throughout the manuscript, separating FDR-supported associations, exploratory directional clues, animal support, peripheral severity context and bulk marker-proxy prioritization.",
        "Fig. 9A shows partial but incomplete separation of control and severity groups in the eight-gene PCA space. Fig. 9B shows module-score variation across control, mild, moderate and severe groups without a perfectly monotonic pattern. Fig. 9C identifies gene-level severity trends, with the strongest result being FLNA: r=-0.788, P=0.000286, FDR=0.00229. Fig. 9D summarizes module-level associations, with disulfidptosis_8gene: r=-0.521, P=0.0383, FDR=0.115.":
            "Fig. 7A shows partial but incomplete separation of control and severity groups in the eight-gene PCA space. Fig. 7B shows module-score variation across control, mild, moderate and severe groups without a perfectly monotonic pattern. Fig. 7C identifies gene-level severity trends, with the strongest result being FLNA: r=-0.788, P=0.000286, FDR=0.00229. Fig. 7D summarizes module-level associations, with disulfidptosis_8gene: r=-0.521, P=0.0383, FDR=0.115.",
        "In panel C, *, ** and *** denote FDR < 0.05, FDR < 0.01 and FDR < 0.001, respectively; ? denotes nominal support without FDR significance.":
            "In panel C, *, ** and *** denote FDR < 0.05, FDR < 0.01 and FDR < 0.001, respectively; † denotes nominal support without FDR significance.",
        "Fig. 9. Evidence audit across human, mouse and interpretation layers. Panel A audits human comparison-level directionality and separates FDR-supported chronic CTE-stage evidence from acute directional support. Panel B ranks mouse CCI support windows, highlighting 3DPI hippocampus as the most coherent animal layer. Panel C displays treatment-stratified module modulation within GSE163415 as contextual information. Panel D summarizes the evidence boundary used for interpretation; RNA-level association and localization do not establish disulfidptotic cell death.":
            "Fig. 9. Evidence audit across human, mouse and interpretation layers. Panel A audits human comparison-level directionality and separates FDR-supported chronic CTE-stage evidence from acute directional support. Panel B ranks mouse CCI support windows, highlighting 3DPI hippocampus as the most coherent animal layer. Panel C displays treatment-stratified CCI module modulation with compact transporter and actin labels to preserve readability. Panel D visualizes the evidence-boundary ladder used to keep human association, animal support, peripheral severity, cell-context localization and mechanistic inference separate.",
    }
    replace_text_in_doc(doc, replacements)


def soften_binary_phrasing(doc: Document) -> None:
    replacements = {
        "The study does not prove disulfidptotic cell death, but provides a bounded map for subsequent protein, redox and cytoskeletal validation.":
            "The resulting RNA-level map prioritizes the disease contexts in which protein, redox and cytoskeletal validation would be most informative.",
        "TBI is therefore not a single molecular state; it is a temporally organized disease process.":
            "TBI is therefore a temporally organized disease process.",
        "A peripheral blood sample reflects systemic immune and trauma-severity responses instead of brain-region mechanism.":
            "A peripheral blood sample reflects systemic immune and trauma-severity responses, whereas brain-region mechanism requires tissue evidence.",
        "Under that condition, continued cystine uptake can become a burden instead of a benefit, increasing intracellular disulfide stress.":
            "Under that condition, continued cystine uptake can increase intracellular disulfide stress.",
        "The actin cytoskeleton is not a passive scaffold.":
            "The actin cytoskeleton is an active mechanical scaffold.",
        "Increased mRNA does not prove increased protein, protein abundance does not prove disulfide bonding, and disulfide stress does not by itself prove cell death.":
            "RNA abundance alone cannot establish protein abundance, disulfide bonding or cell-death morphology.",
        "This narrow panel is not intended to cover all redox biology. It is intended to make every result interpretable: does the signal support entry, regulatory coupling, structural endpoint stress, a time window, a severity context or a cell-localization priority?":
            "This narrow panel was designed for focused interpretation: each result can be read in terms of transporter entry, regulatory coupling, structural endpoint stress, time window, severity context or cell-localization priority.",
        "Acute severe TBI, chronic CTE stage, remote regions, mouse CCI and peripheral severity are not interchangeable measures and are therefore interpreted separately.":
            "Acute severe TBI, chronic CTE stage, remote regions, mouse CCI and peripheral severity represent distinct biological layers and are therefore interpreted separately.",
        "For this reason, marker-proxy analysis is used only to summarize candidate cell-context clues, not to assign disulfidptosis to a cell type.":
            "For this reason, marker-proxy analysis is used to summarize candidate cell-context clues, while cell-type assignment requires higher-resolution localization evidence.",
        "The primary purpose of this reanalysis was to establish true nucleus-level localization instead of to overstate donor-level TBI-control contrasts; because control donors include repeated brain regions, between-group cell-type differences are presented only as exploratory clues.":
            "The primary purpose of this reanalysis was to establish true nucleus-level localization. Because control donors include repeated brain regions, between-group cell-type differences are presented only as exploratory clues.",
        "Co-occurrence across these layers can nominate a transporter-redox-actin stress axis, but it cannot replace NADPH/GSH measurement, protein disulfidation, F-actin morphology or cell-death readouts. This guardrail is essential because much of the public evidence is exploratory instead of definitive.":
            "Co-occurrence across these layers can nominate a transporter-redox-actin stress axis. Direct mechanism still requires NADPH/GSH measurement, protein disulfidation, F-actin morphology or cell-death readouts, and much of the public evidence remains exploratory.",
        "Fig. 1A assigns acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets to different evidence layers instead of merging them into a single TBI dataset.":
            "Fig. 1A separates acute severe TBI, chronic CTE, mouse CCI and peripheral severity datasets into distinct evidence layers.",
        "Fig. 1C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints, clarifying why SLC3A2/SLC7A11, WASF2/TLN1 and ACTB/MYH9/MYL6/FLNA should not be interpreted as a homogeneous block.":
            "Fig. 1C maps the eight genes from cystine import to actin regulation and cytoskeletal endpoints, clarifying the internal logic of the transporter-actin sequence.",
        "The human-brain layer separates chronic CTE course, late-stage CTE contrast, external CTE data and remote multi-region TBI context instead of merging all human brain samples into one broad post-injury label (Fig. 2).":
            "The human-brain layer separates chronic CTE course, late-stage CTE contrast, external CTE data and remote multi-region TBI context (Fig. 2).",
        "Bulk transcriptomics cannot separate within-cell upregulation from changes in cell abundance, so Fig. 2 should be read as localization of a candidate disease-course axis instead of proof of cell-death morphology.":
            "Bulk transcriptomics leaves within-cell upregulation and cell-abundance shifts partly entangled, so Fig. 2 localizes a candidate disease-course axis within RNA-level evidence.",
        "These genes form a transporter-regulatory axis instead of a broad cytoskeletal endpoint signature.":
            "These genes form a transporter-regulatory axis.",
        "However, this is not proof of disulfidptosis. Sample size, tissue heterogeneity and the absence of protein or redox readouts require a cautious interpretation.":
            "Sample size, tissue heterogeneity and the absence of protein or redox readouts require a cautious interpretation.",
        "The marker-proxy analysis in Fig. 3 is retained to summarize candidate cell-context clues, not to assign cell origin from bulk data.":
            "The marker-proxy analysis in Fig. 3 summarizes candidate cell-context clues; cell origin requires localization data.",
        "These clues indicate where to look next, not which cell type has already undergone disulfidptosis.":
            "These clues indicate where higher-resolution localization is most informative.",
        "This provided true nucleus-level localization evidence instead of another layer of bulk proxy inference.":
            "This provided true nucleus-level localization evidence beyond bulk proxy inference.",
        "This pattern suggests that the acute severe TBI transporter-entry signal and the cytoskeletal-stress signal may be distributed across interacting compartments instead of converging in one single cell class.":
            "This pattern suggests that the acute severe TBI transporter-entry signal and the cytoskeletal-stress signal may be distributed across interacting compartments.",
        "The figure refines cell-context interpretation but does not by itself prove mechanism.":
            "The figure refines cell-context interpretation within the RNA-level evidence boundary.",
        "They do not substitute for brain tissue.":
            "Matched brain tissue evidence remains required for brain-mechanism inference.",
        "These results are useful because they remind future studies to include injury severity, but they remain peripheral blood/PBMC signals and cannot replace matched brain tissue evidence.":
            "These results support severity-aware interpretation while remaining peripheral blood/PBMC signals.",
        "Bulk marker-proxy correlations show sample-level co-variation, not co-expression within the same cell.":
            "Bulk marker-proxy correlations show sample-level co-variation, whereas co-expression within the same cell requires single-cell or single-nucleus evidence.",
        "The true snRNA atlas therefore functions as an important localization layer, not as proof that any cell type has undergone disulfidptosis.":
            "The true snRNA atlas therefore functions as an important localization layer within the present evidence boundary.",
        "They are labeled as proxy clues because bulk tissue cannot distinguish cell abundance from cell state.":
            "They are labeled as proxy clues because bulk tissue leaves cell abundance and cell state unresolved.",
        "The integrated ranking summarizes public-evidence prioritization and does not function as a diagnostic model (Fig. 8).":
            "The integrated ranking summarizes public-evidence prioritization (Fig. 8).",
        "The score ranks candidate genes by combining human, mouse, acute and peripheral evidence layers; it is not a diagnostic classifier or causal model.":
            "The score ranks candidate genes by combining human, mouse, acute and peripheral evidence layers and should be read as public-evidence prioritization.",
        "WASF2/TLN1 in chronic CTE and ACTB/MYH9/MYL6/FLNA in acute severe TBI and mouse CCI suggest that entry and endpoint signals are not entirely disconnected.":
            "WASF2/TLN1 in chronic CTE and ACTB/MYH9/MYL6/FLNA in acute severe TBI and mouse CCI suggest partial connection between entry and endpoint signals.",
        "The model is biologically plausible because cytoskeletal stress after TBI is not restricted to one cell type [1, 2, 5].":
            "The model is biologically plausible because cytoskeletal stress after TBI spans multiple cell types [1, 2, 5].",
        "This explains why the present analysis treats SLC3A2/SLC7A11 as transporter pressure within a disulfidptosis-like axis, not as an automatically protective or lethal marker.":
            "This explains why the present analysis treats SLC3A2/SLC7A11 as transporter pressure within a context-dependent disulfidptosis-like axis.",
        "These findings are not contradictory.":
            "This difference has a coherent interpretation.",
        "Severity is a multi-layer concept instead of one number that can be merged across datasets.":
            "Severity is a multi-layer concept across the included datasets.",
        "Peripheral FLNA and SLC3A2 trends can guide clinical stratification or graded-injury animal design, but they cannot be used as brain-mechanism evidence.":
            "Peripheral FLNA and SLC3A2 trends can guide clinical stratification or graded-injury animal design; brain-mechanism inference still requires brain tissue.",
        "Even so, localization is not equivalent to mechanistic proof: transcriptomes cannot show F-actin collapse, protein disulfidation or cell-death morphology, and the repeated-region structure of control donors still limits the strength of donor-level cell-type contrasts.":
            "Even so, transcriptomes cannot show F-actin collapse, protein disulfidation or cell-death morphology, and the repeated-region structure of control donors still limits the strength of donor-level cell-type contrasts.",
        "The eight-gene signal may reflect general cytoskeletal remodeling instead of disulfidptosis.":
            "The eight-gene signal may reflect general cytoskeletal remodeling.",
        "WASF2/TLN1 and endpoint gene changes may reflect structural repair after mechanical injury instead of disulfide stress.":
            "WASF2/TLN1 and endpoint gene changes may reflect structural repair after mechanical injury.",
        "Without NADPH/GSH alteration, protein disulfidation or non-reducing migration, F-actin abnormality and matched cell-context localization, the claim should remain broad cytoskeletal/redox stress instead of disulfidptosis.":
            "NADPH/GSH alteration, protein disulfidation or non-reducing migration, F-actin abnormality and matched cell-context localization would be required to move from broad cytoskeletal/redox stress toward disulfidptosis.",
        "Its value is to identify the most plausible disease stage, time window, genes and cellular backgrounds for the disulfidptosis-like transporter-actin signal without presenting unperformed protein, redox or histological experiments as part of the current evidence.":
            "Its value is to identify the most plausible disease stage, time window, genes and cellular backgrounds for the disulfidptosis-like transporter-actin signal while keeping protein, redox and histological evidence outside the present dataset.",
        "The appropriate conclusion is a candidate disulfidptosis-like transporter-actin stress axis, not proven disulfidptosis.":
            "The appropriate conclusion is a candidate disulfidptosis-like transporter-actin stress axis, with disulfidptotic cell death remaining unproven.",
        "This also explains why time and injury severity are central instead of secondary covariates.":
            "Time and injury severity are central biological axes in this analysis.",
        "Without timing, cytoskeletal endpoint stress in acute severe TBI may be confused with transporter-regulatory accumulation in chronic CTE. Without severity stratification, systemic peripheral trauma responses may be misread as brain mechanism. Without region and cell context, donor-level dependence and bulk cell-composition changes may be mistaken for independent within-cell expression changes. A manuscript intended for publication must make these boundaries part of the argument, not only a short limitation paragraph.":
            "Timing separates cytoskeletal endpoint stress in acute severe TBI from transporter-regulatory accumulation in chronic CTE. Severity stratification prevents systemic peripheral trauma responses from being treated as brain mechanism. Region and cell context reduce the risk of confusing donor-level dependence or bulk cell-composition changes with independent within-cell expression changes. These boundaries are part of the argument itself.",
        "GSE223245 is peripheral blood/PBMC and cannot replace brain data.":
            "GSE223245 is peripheral blood/PBMC and remains a severity-context dataset.",
        "The integrated score is a public-evidence prioritization tool instead of a prediction model.":
            "The integrated score is a public-evidence prioritization tool.",
        "Finally, mRNA-level evidence cannot substitute for direct protein, metabolic, morphological or localization evidence of disulfidptotic cell death.":
            "Finally, direct protein, metabolic, morphological and localization evidence remains necessary for disulfidptotic cell-death claims.",
        "Public transcriptomes support a candidate disulfidptosis-like transporter-actin cytoskeletal stress model after TBI/CTE, but they do not prove disulfidptosis.":
            "Public transcriptomes support a candidate disulfidptosis-like transporter-actin cytoskeletal stress model after TBI/CTE within RNA-level evidence boundaries.",
    }
    replace_text_in_doc(doc, replacements)


def replace_images(doc: Document, fig_paths: dict[str, Path]) -> None:
    image_paras = [p for p in doc.paragraphs if "<w:drawing" in p._p.xml]
    ordered = [f"Fig{i}" for i in range(1, 10)]
    for p, key in zip(image_paras, ordered):
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_paths[key]), width=Inches(6.7))


def apply_document_format(doc: Document) -> None:
    heading_texts = {
        "ABSTRACT", "Introduction", "Materials and Methods", "Results", "Discussion", "Conclusion",
        "References", "Statements and Declarations", "Funding", "Competing Interests",
        "Author Contributions", "Ethics Approval", "Consent to Participate", "Consent for Publication",
        "Data Availability", "Code Availability",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        set_paragraph_format(p, title=(text == TITLE), heading=(text in heading_texts or text.startswith("Table ") or text.startswith("Fig. ")))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_format(p)


def clean_placeholder_paragraphs(doc: Document) -> None:
    for p in list(doc.paragraphs):
        text = p.text
        if "These authors contributed equally" in text or "Delete if not applicable" in text:
            remove_paragraph(p)


def update_manuscript(fig_paths: dict[str, Path]) -> Path:
    path = FINAL_OUT / "01_Manuscript.docx"
    backup = CHECK_OUT / "01_Manuscript_before_direct_submission_finalize_20260623.docx"
    if path.exists() and not backup.exists():
        shutil.copy2(path, backup)
    doc = Document(str(path))
    replace_front_matter(doc)
    clean_placeholder_paragraphs(doc)
    add_body_citations(doc)
    update_captions_and_text(doc)
    soften_binary_phrasing(doc)
    update_declarations(doc)
    move_declarations_before_references(doc)
    add_reference_numbers(doc)
    replace_images(doc, fig_paths)
    apply_document_format(doc)
    doc.save(path)
    shutil.copy2(path, OUT_ROOT / "01_manuscript" / "Manuscript_NeurochemicalResearch_direct_submission_20260623.docx")
    return path


def update_title_page() -> Path:
    path = FINAL_OUT / "02_Title_page.docx"
    doc = Document(str(path))
    replace_front_matter(doc)
    clean_placeholder_paragraphs(doc)
    update_declarations(doc)
    apply_document_format(doc)
    doc.save(path)
    return path


def update_cover_letter() -> Path:
    path = FINAL_OUT / "03_Cover_letter.docx"
    doc = Document(str(path))
    replacements = {
        "[Corresponding author full name]": "Mingyang Zhang",
        "[active e-mail address]": "mingyangzhang@suda.edu.cn",
        "Dear editor:": "Dear Editor:",
    }
    replace_text_in_doc(doc, replacements)
    for p in doc.paragraphs:
        if p.text.strip() == "Sincerely,":
            set_paragraph_format(p)
        elif p.text.strip() == "Mingyang Zhang":
            set_paragraph_format(p)
        elif p.text.strip().startswith("E-mail:"):
            set_paragraph_format(p)
    doc.save(path)
    return path


def update_highlights() -> Path:
    path = FINAL_OUT / "04_Highlights.docx"
    doc = Document(str(path))
    for p in doc.paragraphs:
        set_paragraph_format(p)
    doc.save(path)
    return path


def make_contact_sheet() -> Path:
    imgs = []
    for f in sorted(POLISHED_PNG.glob("Fig*.png")):
        im = Image.open(f).convert("RGB")
        im.thumbnail((420, 305))
        imgs.append((f.name, im.copy()))
    w = 840
    h = ((len(imgs) + 1) // 2) * 360
    sheet = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(sheet)
    for idx, (name, im) in enumerate(imgs):
        x = (idx % 2) * 420
        y = (idx // 2) * 360
        draw.text((x + 8, y + 8), name, fill=(0, 0, 0))
        sheet.paste(im, (x, y + 32))
    out = CHECK_OUT / "figure_contact_sheet_after_direct_submission_finalize_20260623.png"
    sheet.save(out)
    return out


def make_resolution_summary() -> Path:
    rows = []
    for tif in sorted(FINAL_TIFF.glob("Fig*.tif")):
        im = Image.open(tif)
        rows.append(
            {
                "file": tif.name,
                "pixel_width": im.width,
                "pixel_height": im.height,
                "dpi": im.info.get("dpi", ""),
                "width_at_600dpi_mm": round(im.width / 600 * 25.4, 1),
                "height_at_600dpi_mm": round(im.height / 600 * 25.4, 1),
            }
        )
    out = CHECK_OUT / "Figure_resolution_check_direct_submission_20260623.csv"
    pd.DataFrame(rows).to_csv(out, index=False, encoding="utf-8-sig")
    return out


def package_submission() -> Path:
    base = ROOT / "TBI投稿关键文件包_20260623" / "Neurochemical_Research_直接投稿包_20260623"
    target = base
    counter = 2
    while target.exists():
        target = Path(f"{base}_v{counter}")
        counter += 1
    target.mkdir(parents=True, exist_ok=False)
    for name in ["01_Manuscript.docx", "02_Title_page.docx", "03_Cover_letter.docx", "04_Highlights.docx", "README_final_submission_package.txt"]:
        src = FINAL_OUT / name
        if src.exists():
            shutil.copy2(src, target / name)
    for dirname in ["05_Graphical_Abstract", "06_Figures_600dpi_TIFF", "07_Checks"]:
        src_dir = FINAL_OUT / dirname
        if src_dir.exists():
            shutil.copytree(src_dir, target / dirname)
    readme = target / "README_直接投稿包说明.txt"
    readme.write_text(
        "Neurochemical Research direct-submission package generated on 2026-06-23.\n"
        "Core files: manuscript, title page, cover letter, highlights, graphical abstract, and 600-dpi TIFF figures.\n"
        "Author information follows the confirmed project instruction: Chenxu Zhang and Mingyang Zhang; corresponding author: Professor Mingyang Zhang (mingyangzhang@suda.edu.cn).\n",
        encoding="utf-8",
    )
    return target


def qa_check(package_dir: Path, contact_sheet: Path, resolution_csv: Path) -> Path:
    doc = Document(str(FINAL_OUT / "01_Manuscript.docx"))
    text = paragraph_text(doc)
    needles = [
        "[Author", "[Corresponding", "[Department", "[Affiliated", "[Institution", "[City",
        "[active e-mail", "Delete if not applicable", "Professor/Dr.", "wet-lab",
        "Reading Logic", "Note.", "rather than", "Fig. 10", "Fig. 11",
    ]
    rows = [{"check": needle, "count": text.count(needle)} for needle in needles]
    rows.extend(
        [
            {"check": "inline_images", "count": sum(1 for p in doc.paragraphs if "<w:drawing" in p._p.xml)},
            {"check": "tables", "count": len(doc.tables)},
            {"check": "references_numbered", "count": sum(1 for p in doc.paragraphs if re.match(r"^\[\d+\]\s+", p.text.strip()))},
            {"check": "package_dir", "count": str(package_dir)},
            {"check": "contact_sheet", "count": str(contact_sheet)},
            {"check": "resolution_csv", "count": str(resolution_csv)},
        ]
    )
    out = CHECK_OUT / "direct_submission_QA_20260623.csv"
    pd.DataFrame(rows).to_csv(out, index=False, encoding="utf-8-sig")
    return out


def main() -> None:
    ensure_dirs()
    configure_style()
    fig1 = make_fig1()
    fig9 = make_fig9()
    fig_paths = {f"Fig{i}": POLISHED_PNG / f"Fig{i}.png" for i in range(1, 10)}
    fig_paths["Fig1"] = fig1
    fig_paths["Fig9"] = fig9
    update_manuscript(fig_paths)
    update_title_page()
    update_cover_letter()
    update_highlights()
    contact_sheet = make_contact_sheet()
    resolution_csv = make_resolution_summary()
    package_dir = package_submission()
    qa_csv = qa_check(package_dir, contact_sheet, resolution_csv)
    print("Updated manuscript:", FINAL_OUT / "01_Manuscript.docx")
    print("Updated figures:", FINAL_TIFF)
    print("Contact sheet:", contact_sheet)
    print("Submission package:", package_dir)
    print("QA:", qa_csv)


if __name__ == "__main__":
    main()
